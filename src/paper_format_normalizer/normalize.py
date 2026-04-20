from __future__ import annotations

import re
from collections import defaultdict
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from time import perf_counter
from typing import Literal

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.section import Section
from docx.shared import Cm, Inches, Length, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from paper_format_normalizer.classify import (
    _MATCH_KIND_ORDER,
    _numbering_rule_matches,
    _paragraph_rule_candidate,
    _table_rule_candidate,
    classify_document,
)
from paper_format_normalizer.model import (
    DocumentRule,
    NumberingRule,
    ParagraphRule,
    RuleSet,
    SpecialObjectRule,
    TableRule,
)
from paper_format_normalizer.parse import (
    ParsedBodyParagraph,
    ParsedBodyTable,
    _iter_body_elements,
    parse_docx,
)
from paper_format_normalizer.report import write_report

ReportStatus = Literal["modified", "unchanged", "unresolved"]
_RED = RGBColor(0xFF, 0x00, 0x00)
_TEXTUAL_PROPERTIES = frozenset(
    {
        "font_name",
        "font_size",
        "header_row_font_name",
        "header_row_font_size",
        "body_rows_font_name",
        "body_rows_font_size",
        "bold",
        "alignment",
        "vertical_alignment",
        "border",
        "header_row_bold",
        "header_row_alignment",
        "header_row_vertical_alignment",
        "header_row_border",
        "body_rows_bold",
        "body_rows_alignment",
        "body_rows_vertical_alignment",
        "body_rows_border",
        "label_font_name",
        "content_font_name",
        "label_font_size",
        "content_font_size",
        "label_bold",
        "content_bold",
    }
)
_TABLE_SELECTOR_PROPERTY_RE = re.compile(
    r"^(?P<selector_kind>column|cell)\[(?P<first>\d+)(?:,(?P<second>\d+))?\]_(?P<base>font_name|font_size|bold|alignment|vertical_alignment|border)$"
)
_TABLE_HEADER_SELECTOR_PROPERTY_RE = re.compile(
    r"^column_by_header\[(?P<header>[^\]]+)\]_(?P<base>font_name|font_size|bold|alignment|vertical_alignment|border)$"
)
_TABLE_COLUMN_RANGE_SELECTOR_PROPERTY_RE = re.compile(
    r"^column_range\[(?P<start>\d+):(?P<end>\d+)\]_(?P<base>font_name|font_size|bold|alignment|vertical_alignment|border)$"
)
_TABLE_ROW_SELECTOR_PROPERTY_RE = re.compile(
    r"^row\[(?P<row>\d+)\]_(?P<base>font_name|font_size|bold|alignment|vertical_alignment|border)$"
)
_TABLE_ROW_RANGE_SELECTOR_PROPERTY_RE = re.compile(
    r"^row_range\[(?P<start>\d+):(?P<end>\d+)\]_(?P<base>font_name|font_size|bold|alignment|vertical_alignment|border)$"
)
_TABLE_CELL_RANGE_SELECTOR_PROPERTY_RE = re.compile(
    r"^cell_range\[(?P<row_start>\d+):(?P<row_end>\d+),(?P<col_start>\d+):(?P<col_end>\d+)\]_(?P<base>font_name|font_size|bold|alignment|vertical_alignment|border)$"
)
_ANNOTATION_PROPERTIES = frozenset(
    {
        "first_line_indent",
        "hanging_indent",
        "space_before",
        "space_after",
        "line_spacing",
        "page_margin_top",
        "page_margin_bottom",
        "page_margin_left",
        "page_margin_right",
    }
)
_PROPERTY_LABELS = {
    "first_line_indent": "\u9996\u884c\u7f29\u8fdb",
    "hanging_indent": "\u60ac\u6302\u7f29\u8fdb",
    "space_before": "\u6bb5\u524d",
    "space_after": "\u6bb5\u540e",
    "line_spacing": "\u884c\u8ddd",
    "page_margin_top": "\u4e0a\u8fb9\u8ddd",
    "page_margin_bottom": "\u4e0b\u8fb9\u8ddd",
    "page_margin_left": "\u5de6\u8fb9\u8ddd",
    "page_margin_right": "\u53f3\u8fb9\u8ddd",
    "section_start_type": "\u5206\u8282\u8d77\u59cb\u65b9\u5f0f",
    "page_number_format": "\u9875\u7801\u683c\u5f0f",
    "page_number_start": "\u9875\u7801\u8d77\u59cb\u503c",
    "footer_page_number_alignment": "\u9875\u811a\u9875\u7801\u5bf9\u9f50",
    "different_first_page_header_footer": "\u9996\u9875\u4e0d\u540c\u9875\u7709\u9875\u811a",
    "odd_and_even_pages_header_footer": "\u5947\u5076\u9875\u4e0d\u540c\u9875\u7709\u9875\u811a",
    "label_font_name": "\u6807\u7b7e\u5b57\u4f53",
    "content_font_name": "\u5185\u5bb9\u5b57\u4f53",
    "label_font_size": "\u6807\u7b7e\u5b57\u53f7",
    "content_font_size": "\u5185\u5bb9\u5b57\u53f7",
    "label_bold": "\u6807\u7b7e\u52a0\u7c97",
    "content_bold": "\u5185\u5bb9\u52a0\u7c97",
    "header_row_font_name": "\u8868\u5934\u5b57\u4f53",
    "header_row_font_size": "\u8868\u5934\u5b57\u53f7",
    "body_rows_font_name": "\u6570\u636e\u884c\u5b57\u4f53",
    "body_rows_font_size": "\u6570\u636e\u884c\u5b57\u53f7",
    "vertical_alignment": "\u5782\u76f4\u5bf9\u9f50",
    "border": "\u8fb9\u6846",
    "header_row_vertical_alignment": "\u8868\u5934\u5782\u76f4\u5bf9\u9f50",
    "header_row_border": "\u8868\u5934\u8fb9\u6846",
    "body_rows_vertical_alignment": "\u6570\u636e\u884c\u5782\u76f4\u5bf9\u9f50",
    "body_rows_border": "\u6570\u636e\u884c\u8fb9\u6846",
}
_ANNOTATION_FONT_NAME = "SimSun"
_ANNOTATION_FONT_SIZE = Pt(10.5)
_WESTERN_FONT_NAME = "Times New Roman"
_TABLE_ANNOTATION_PROPERTIES = frozenset(
    {
        "vertical_alignment",
        "border",
        "header_row_vertical_alignment",
        "header_row_border",
        "body_rows_vertical_alignment",
        "body_rows_border",
    }
)
_TABLE_CELL_ANNOTATION_BASE_PROPERTIES = frozenset({"vertical_alignment", "border"})
_TABLE_BORDER_EDGES = ("top", "left", "bottom", "right")
_FONT_DISPLAY_NAMES = {
    "Arial": "Arial\uff08\u65e0\u886c\u7ebf\u4f53\uff09",
    "Calibri": "Calibri\uff08\u9ed8\u8ba4\u65e0\u886c\u7ebf\u4f53\uff09",
    "KaiTi": "\u6977\u4f53\uff08KaiTi\uff09",
    "FangSong": "\u4eff\u5b8b\uff08FangSong\uff09",
    "SimHei": "\u9ed1\u4f53\uff08SimHei\uff09",
    "SimSun": "\u5b8b\u4f53\uff08SimSun\uff09",
    "Times New Roman": "Times New Roman\uff08\u65b0\u7f57\u9a6c\u4f53\uff09",
}
_EAST_ASIAN_FONT_NAMES = frozenset({"SimHei", "SimSun", "FangSong", "KaiTi"})
_INLINE_PREFIXES = (
    "\u3010\u6458\u8981\u3011",
    "\u3010\u5173\u952e\u8bcd\u3011",
    "\u3010Abstract\u3011",
    "\u3010KeyWords\u3011",
)


@dataclass(frozen=True)
class _MatchedParagraphRule:
    rule_id: str
    source_family: str
    priority: int
    match_kind: str
    match_type: str
    match_value: str
    target_property: str
    target_value: str


@dataclass(frozen=True)
class _RunLengthState:
    values: tuple[Length | None, ...]


@dataclass(frozen=True)
class _MatchedTableRule:
    rule_id: str
    priority: int
    match_kind: str
    match_type: str
    match_value: str
    target_property: str
    target_value: str


@dataclass(frozen=True)
class _TablePropertyTarget:
    selector_kind: str
    indices: tuple[int, ...]
    base_property: str
    header_name: str | None = None


@dataclass(frozen=True)
class _DocumentRuntimeIndex:
    body_paragraphs: dict[str, Paragraph]
    body_tables: dict[str, Table]
    header_paragraphs: dict[str, Paragraph]
    footer_paragraphs: dict[str, Paragraph]


@dataclass(frozen=True)
class _ParsedLocationIndex:
    body_paragraphs: dict[str, ParsedBodyParagraph]
    body_tables: dict[str, ParsedBodyTable]
    header_paragraphs: dict[str, ParsedBodyParagraph]
    footer_paragraphs: dict[str, ParsedBodyParagraph]


@dataclass(frozen=True)
class _AnnotationPlan:
    section_rows: list[dict[str, str]]
    rows_by_location: dict[str, list[dict[str, str]]]


@dataclass(frozen=True)
class NormalizeBenchmarkRun:
    output_path: Path
    report_path: Path
    annotated_path: Path
    timings: dict[str, float]


def normalize_document(
    input_path: Path,
    rule_set: RuleSet,
    output_dir: Path,
) -> tuple[Path, Path, Path]:
    benchmark_run = benchmark_normalize_document(input_path, rule_set, output_dir)
    return benchmark_run.output_path, benchmark_run.report_path, benchmark_run.annotated_path


def benchmark_normalize_document(
    input_path: Path,
    rule_set: RuleSet,
    output_dir: Path,
) -> NormalizeBenchmarkRun:
    timings: dict[str, float] = {}
    total_started = perf_counter()

    stage_started = perf_counter()
    parsed = parse_docx(input_path)
    timings["parse_docx_seconds"] = perf_counter() - stage_started

    stage_started = perf_counter()
    classification = classify_document(parsed, rule_set)
    timings["classify_document_seconds"] = perf_counter() - stage_started

    stage_started = perf_counter()
    document = Document(input_path)
    timings["load_document_seconds"] = perf_counter() - stage_started

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{input_path.stem}_\u89c4\u8303\u5316.docx"
    report_path = output_dir / f"{input_path.stem}_\u89c4\u8303\u5316_\u4fee\u6539\u62a5\u544a.csv"
    annotated_path = output_dir / f"{input_path.stem}_\u89c4\u8303\u5316_\u7ea2\u5b57\u6807\u6ce8\u7248.docx"

    parsed_index = _build_parsed_location_index(parsed)
    runtime_index = _build_document_runtime_index(document)
    classification_by_location = {
        result.location: result for result in classification.object_results
    }

    stage_started = perf_counter()
    report_rows: list[dict[str, str]] = []
    report_rows.extend(_apply_document_rules(document, rule_set.document_rules))
    report_rows.extend(
        _apply_content_rules(
            parsed_index=parsed_index,
            runtime_index=runtime_index,
            classification_by_location=classification_by_location,
            classification_results=classification.object_results,
            rule_set=rule_set,
        )
    )
    timings["apply_rules_seconds"] = perf_counter() - stage_started

    stage_started = perf_counter()
    document.save(output_path)
    timings["save_normalized_docx_seconds"] = perf_counter() - stage_started

    stage_started = perf_counter()
    write_report(report_path, report_rows, rule_set)
    timings["write_report_seconds"] = perf_counter() - stage_started

    stage_started = perf_counter()
    annotation_plan = _build_annotation_plan(report_rows)
    timings["build_annotation_plan_seconds"] = perf_counter() - stage_started

    annotated_document = deepcopy(document)
    annotated_runtime_index = _build_document_runtime_index(annotated_document)

    stage_started = perf_counter()
    _write_annotated_document(
        annotated_path=annotated_path,
        annotation_plan=annotation_plan,
        document=annotated_document,
        runtime_index=annotated_runtime_index,
    )
    timings["write_annotated_docx_seconds"] = perf_counter() - stage_started

    timings["total_seconds"] = perf_counter() - total_started
    return NormalizeBenchmarkRun(
        output_path=output_path,
        report_path=report_path,
        annotated_path=annotated_path,
        timings=timings,
    )


def _apply_document_rules(document: DocumentObject, rules: list[DocumentRule]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    grouped_rules: dict[str, list[DocumentRule]] = defaultdict(list)
    for rule in rules:
        grouped_rules[rule.property_name].append(rule)

    for property_name, property_rules in grouped_rules.items():
        property_accessor = _document_property_accessor(property_name)
        sorted_rules = sorted(property_rules, key=lambda rule: (rule.priority, rule.rule_id))
        winner = sorted_rules[0]
        conflict_rules = [
            rule for rule in sorted_rules if rule.priority == winner.priority
        ]
        if len(conflict_rules) > 1:
            conflict_ids = ", ".join(rule.rule_id for rule in conflict_rules)
            if property_accessor is not None and property_accessor.target_kind == "document":
                rows.append(
                    _report_row(
                        object_id="document-settings",
                        object_type_before="document",
                        object_type_after="document",
                        location="document_settings",
                        text_preview="",
                        property_name=property_name,
                        before="",
                        after="",
                        rule_id="",
                        status="unresolved",
                        reason=f"conflicting document rules at same priority: {conflict_ids}",
                    )
                )
            else:
                for section_index, _ in enumerate(document.sections):
                    rows.append(
                        _report_row(
                            object_id=f"section-{section_index}",
                            object_type_before="document",
                            object_type_after="document",
                            location=f"sections[{section_index}]",
                            text_preview="",
                            property_name=property_name,
                            before="",
                            after="",
                            rule_id="",
                            status="unresolved",
                            reason=f"conflicting document rules at same priority: {conflict_ids}",
                        )
                    )
            continue

        if property_accessor is not None and property_accessor.target_kind == "document":
            rows.append(_apply_document_rule(document=document, section=None, section_index=None, rule=winner))
            continue

        for section_index, section in enumerate(document.sections):
            rows.append(_apply_document_rule(document=document, section=section, section_index=section_index, rule=winner))
    return rows


def _apply_content_rules(
    *,
    parsed_index: _ParsedLocationIndex,
    runtime_index: _DocumentRuntimeIndex,
    classification_by_location: dict[str, object],
    classification_results,
    rule_set: RuleSet,
) -> list[dict[str, str]]:
    report_rows: list[dict[str, str]] = []
    handled_locations: set[str] = set()

    for location, paragraph in runtime_index.body_paragraphs.items():
        parsed_item = parsed_index.body_paragraphs.get(location)
        if parsed_item is None:
            continue
        handled_locations.add(location)
        classification_result = classification_by_location[location]
        report_rows.extend(
            _apply_paragraph_rules(
                paragraph=paragraph,
                parsed_paragraph=parsed_item,
                classification_result=classification_result,
                rule_set=rule_set,
            )
        )

    for location, paragraph in runtime_index.header_paragraphs.items():
        parsed_header_item = parsed_index.header_paragraphs.get(location)
        if parsed_header_item is None:
            continue
        handled_locations.add(location)
        classification_result = classification_by_location[location]
        report_rows.extend(
            _apply_header_paragraph_rules(
                paragraph=paragraph,
                classification_result=classification_result,
                rule_set=rule_set,
            )
        )

    for location, paragraph in runtime_index.footer_paragraphs.items():
        if location not in classification_by_location:
            continue
        parsed_footer_item = parsed_index.footer_paragraphs.get(location)
        if parsed_footer_item is None:
            continue
        handled_locations.add(location)
        classification_result = classification_by_location[location]
        report_rows.extend(
            _apply_header_paragraph_rules(
                paragraph=paragraph,
                classification_result=classification_result,
                rule_set=rule_set,
            )
        )

    for location, table in runtime_index.body_tables.items():
        parsed_item = parsed_index.body_tables.get(location)
        if parsed_item is None:
            continue
        handled_locations.add(location)
        classification_result = classification_by_location[location]
        report_rows.extend(
            _apply_table_rules(
                table=table,
                parsed_table=parsed_item,
                classification_result=classification_result,
                rule_set=rule_set,
            )
        )

    for result in classification_results:
        if result.location in handled_locations:
            continue
        report_rows.append(
            _report_row(
                object_id=result.object_id,
                object_type_before=result.object_type,
                object_type_after=result.object_type,
                location=result.location,
                text_preview=result.original_text,
                property_name="classification",
                before="",
                after="",
                rule_id=result.matched_rule_id or "",
                status="unresolved",
                reason=(
                    result.reason
                    if result.status == "unresolved"
                    else f"phase-1 normalization does not support {result.object_type} objects"
                ),
            )
        )

    return report_rows


def _apply_document_rule(
    *,
    document: DocumentObject,
    section: Section | None,
    section_index: int | None,
    rule: DocumentRule,
) -> dict[str, str]:
    property_accessor = _document_property_accessor(rule.property_name)
    if property_accessor is not None and property_accessor.target_kind == "document":
        object_id = "document-settings"
        location = "document_settings"
        target_object = document
    else:
        object_id = f"section-{section_index}"
        location = f"sections[{section_index}]"
        target_object = section

    if rule.scope != "document":
        return _report_row(
            object_id=object_id,
            object_type_before="document",
            object_type_after="document",
            location=location,
            text_preview="",
            property_name=rule.property_name,
            before="",
            after="",
            rule_id=rule.rule_id,
            status="unresolved",
            reason=f"unsupported document rule scope: {rule.scope}",
        )

    if property_accessor is None:
        return _report_row(
            object_id=object_id,
            object_type_before="document",
            object_type_after="document",
            location=location,
            text_preview="",
            property_name=rule.property_name,
            before="",
            after="",
            rule_id=rule.rule_id,
            status="unresolved",
            reason=f"unsupported document property: {rule.property_name}",
        )

    before_value = property_accessor.get(target_object)
    try:
        target_value = property_accessor.parse(rule.value)
    except ValueError as exc:
        return _report_row(
            object_id=object_id,
            object_type_before="document",
            object_type_after="document",
            location=location,
            text_preview="",
            property_name=rule.property_name,
            before=property_accessor.format(before_value, rule.value),
            after=rule.value,
            rule_id=rule.rule_id,
            status="unresolved",
            reason=str(exc),
        )

    property_accessor.set(target_object, target_value)
    status: ReportStatus = "modified" if before_value != target_value else "unchanged"
    return _report_row(
        object_id=object_id,
        object_type_before="document",
        object_type_after="document",
        location=location,
        text_preview="",
        property_name=rule.property_name,
        before=property_accessor.format(before_value, rule.value),
        after=property_accessor.format(target_value, rule.value),
        rule_id=rule.rule_id,
        status=status,
        reason="",
    )


def _apply_paragraph_rules(
    *,
    paragraph: Paragraph,
    parsed_paragraph: ParsedBodyParagraph,
    classification_result,
    rule_set: RuleSet,
) -> list[dict[str, str]]:
    if classification_result.status == "unresolved":
        return [
            _report_row(
                object_id=classification_result.object_id,
                object_type_before=classification_result.object_type,
                object_type_after=classification_result.object_type,
                location=classification_result.location,
                text_preview=classification_result.original_text,
                property_name="classification",
                before="",
                after="",
                rule_id="",
                status="unresolved",
                reason=classification_result.reason or "no matching classification rule",
            )
        ]

    rows: list[dict[str, str]] = []
    matched_rules = _matched_paragraph_rules(parsed_paragraph, rule_set)
    compatible_rules = _compatible_paragraph_rules(
        matched_rules=matched_rules,
        classification_result=classification_result,
    )
    rules_by_property: dict[str, list[_MatchedParagraphRule]] = defaultdict(list)
    for rule in compatible_rules:
        rules_by_property[rule.target_property].append(rule)

    if not rules_by_property:
        rows.append(
            _report_row(
                object_id=classification_result.object_id,
                object_type_before=classification_result.object_type,
                object_type_after=classification_result.object_type,
                location=classification_result.location,
                text_preview=classification_result.original_text,
                property_name="classification",
                before="",
                after="",
                rule_id=classification_result.matched_rule_id or "",
                status="unresolved",
                reason=(
                    classification_result.reason
                    or "no normalization rule compatible with resolved classification winner"
                ),
            )
        )
        return rows

    for property_name, property_rules in sorted(rules_by_property.items()):
        sorted_rules = sorted(
            property_rules,
            key=lambda rule: (
                _MATCH_KIND_ORDER[rule.match_kind],
                rule.priority,
                rule.rule_id,
            ),
        )
        winner = sorted_rules[0]
        conflict_rules = [
            rule
            for rule in sorted_rules
            if rule.match_kind == winner.match_kind and rule.priority == winner.priority
        ]
        if len(conflict_rules) > 1:
            conflict_ids = ", ".join(rule.rule_id for rule in conflict_rules)
            rows.append(
                _report_row(
                    object_id=classification_result.object_id,
                    object_type_before=classification_result.object_type,
                    object_type_after=classification_result.object_type,
                    location=classification_result.location,
                    text_preview=classification_result.original_text,
                    property_name=property_name,
                    before="",
                    after="",
                    rule_id="",
                    status="unresolved",
                    reason=(
                        "conflicting rules at same priority: "
                        f"{conflict_ids} "
                        f"(match_kind={winner.match_kind}, priority={winner.priority})"
                    ),
                )
            )
            continue

        rows.append(
            _apply_single_paragraph_rule(
                paragraph=paragraph,
                classification_result=classification_result,
                rule=winner,
            )
        )

    return rows


def _apply_header_paragraph_rules(
    *,
    paragraph: Paragraph,
    classification_result,
    rule_set: RuleSet,
) -> list[dict[str, str]]:
    if classification_result.status == "unresolved":
        return [
            _report_row(
                object_id=classification_result.object_id,
                object_type_before=classification_result.object_type,
                object_type_after=classification_result.object_type,
                location=classification_result.location,
                text_preview=classification_result.original_text,
                property_name="classification",
                before="",
                after="",
                rule_id="",
                status="unresolved",
                reason=classification_result.reason or "no matching classification rule",
            )
        ]

    object_class = _special_object_target_object_type(
        rule_set.special_object_rules,
        classification_result.matched_rule_id,
    )
    if object_class is None:
        return [
            _report_row(
                object_id=classification_result.object_id,
                object_type_before=classification_result.object_type,
                object_type_after=classification_result.object_type,
                location=classification_result.location,
                text_preview=classification_result.original_text,
                property_name="classification",
                before="",
                after="",
                rule_id=classification_result.matched_rule_id or "",
                status="unresolved",
                reason="no target_object_type found for resolved header classification rule",
            )
        ]

    matched_rules = _matched_header_paragraph_rules(object_class, rule_set)
    if not matched_rules:
        return [
            _report_row(
                object_id=classification_result.object_id,
                object_type_before=classification_result.object_type,
                object_type_after=classification_result.object_type,
                location=classification_result.location,
                text_preview=classification_result.original_text,
                property_name="classification",
                before="",
                after="",
                rule_id=classification_result.matched_rule_id or "",
                status="unresolved",
                reason="no normalization rule compatible with resolved header classification winner",
            )
        ]

    rows: list[dict[str, str]] = []
    rules_by_property: dict[str, list[_MatchedParagraphRule]] = defaultdict(list)
    for rule in matched_rules:
        rules_by_property[rule.target_property].append(rule)

    for property_name, property_rules in sorted(rules_by_property.items()):
        sorted_rules = sorted(
            property_rules,
            key=lambda rule: (
                _MATCH_KIND_ORDER[rule.match_kind],
                rule.priority,
                rule.rule_id,
            ),
        )
        winner = sorted_rules[0]
        conflict_rules = [
            rule
            for rule in sorted_rules
            if rule.match_kind == winner.match_kind and rule.priority == winner.priority
        ]
        if len(conflict_rules) > 1:
            conflict_ids = ", ".join(rule.rule_id for rule in conflict_rules)
            rows.append(
                _report_row(
                    object_id=classification_result.object_id,
                    object_type_before=classification_result.object_type,
                    object_type_after=classification_result.object_type,
                    location=classification_result.location,
                    text_preview=classification_result.original_text,
                    property_name=property_name,
                    before="",
                    after="",
                    rule_id="",
                    status="unresolved",
                    reason=(
                        "conflicting rules at same priority: "
                        f"{conflict_ids} "
                        f"(match_kind={winner.match_kind}, priority={winner.priority})"
                    ),
                )
            )
            continue

        rows.append(
            _apply_single_paragraph_rule(
                paragraph=paragraph,
                classification_result=classification_result,
                rule=winner,
            )
        )

    return rows


def _apply_table_rules(
    *,
    table: Table,
    parsed_table: ParsedBodyTable,
    classification_result,
    rule_set: RuleSet,
) -> list[dict[str, str]]:
    if classification_result.status == "unresolved":
        return [
            _report_row(
                object_id=classification_result.object_id,
                object_type_before=classification_result.object_type,
                object_type_after=classification_result.object_type,
                location=classification_result.location,
                text_preview=classification_result.original_text,
                property_name="classification",
                before="",
                after="",
                rule_id="",
                status="unresolved",
                reason=classification_result.reason or "no matching classification rule",
            )
        ]

    matched_rules = _matched_table_rules(parsed_table, rule_set)
    compatible_rules = _compatible_table_rules(
        matched_rules=matched_rules,
        classification_result=classification_result,
    )
    if not compatible_rules:
        return [
            _report_row(
                object_id=classification_result.object_id,
                object_type_before=classification_result.object_type,
                object_type_after=classification_result.object_type,
                location=classification_result.location,
                text_preview=classification_result.original_text,
                property_name="classification",
                before="",
                after="",
                rule_id=classification_result.matched_rule_id or "",
                status="unresolved",
                reason="no normalization rule compatible with resolved table classification winner",
            )
        ]

    rows: list[dict[str, str]] = []
    rules_by_property: dict[str, list[_MatchedTableRule]] = defaultdict(list)
    for rule in compatible_rules:
        rules_by_property[rule.target_property].append(rule)

    for property_name, property_rules in sorted(rules_by_property.items()):
        sorted_rules = sorted(
            property_rules,
            key=lambda rule: (
                _MATCH_KIND_ORDER[rule.match_kind],
                rule.priority,
                rule.rule_id,
            ),
        )
        winner = sorted_rules[0]
        conflict_rules = [
            rule
            for rule in sorted_rules
            if rule.match_kind == winner.match_kind and rule.priority == winner.priority
        ]
        if len(conflict_rules) > 1:
            conflict_ids = ", ".join(rule.rule_id for rule in conflict_rules)
            rows.append(
                _report_row(
                    object_id=classification_result.object_id,
                    object_type_before=classification_result.object_type,
                    object_type_after=classification_result.object_type,
                    location=classification_result.location,
                    text_preview=classification_result.original_text,
                    property_name=property_name,
                    before="",
                    after="",
                    rule_id="",
                    status="unresolved",
                    reason=(
                        "conflicting rules at same priority: "
                        f"{conflict_ids} "
                        f"(match_kind={winner.match_kind}, priority={winner.priority})"
                    ),
                )
            )
            continue

        rows.append(
            _apply_single_table_rule(
                table=table,
                classification_result=classification_result,
                rule=winner,
            )
        )

    return rows


def _apply_single_paragraph_rule(
    *,
    paragraph: Paragraph,
    classification_result,
    rule: _MatchedParagraphRule,
) -> dict[str, str]:
    property_accessor = _paragraph_property_accessor(rule.target_property)
    if property_accessor is None:
        return _report_row(
            object_id=classification_result.object_id,
            object_type_before=classification_result.object_type,
            object_type_after=classification_result.object_type,
            location=classification_result.location,
            text_preview=classification_result.original_text,
            property_name=rule.target_property,
            before="",
            after="",
            rule_id=rule.rule_id,
            status="unresolved",
            reason=f"unsupported paragraph property: {rule.target_property}",
        )

    before_value = property_accessor.get(paragraph)
    try:
        target_value = property_accessor.parse(rule.target_value)
    except ValueError as exc:
        return _report_row(
            object_id=classification_result.object_id,
            object_type_before=classification_result.object_type,
            object_type_after=classification_result.object_type,
            location=classification_result.location,
            text_preview=classification_result.original_text,
            property_name=rule.target_property,
            before=property_accessor.format(before_value, rule.target_value),
            after=rule.target_value,
            rule_id=rule.rule_id,
            status="unresolved",
            reason=str(exc),
        )

    property_accessor.set(paragraph, target_value)
    after_value = property_accessor.get(paragraph)
    status: ReportStatus = "modified" if before_value != after_value else "unchanged"
    return _report_row(
        object_id=classification_result.object_id,
        object_type_before=classification_result.object_type,
        object_type_after=classification_result.object_type,
        location=classification_result.location,
        text_preview=classification_result.original_text,
        property_name=rule.target_property,
        before=property_accessor.format(before_value, rule.target_value),
        after=property_accessor.format(after_value, rule.target_value),
        rule_id=rule.rule_id,
        status=status,
        reason="",
    )


def _apply_single_table_rule(
    *,
    table: Table,
    classification_result,
    rule: _MatchedTableRule,
) -> dict[str, str]:
    property_accessor = _table_property_accessor(rule.target_property)
    if property_accessor is None:
        return _report_row(
            object_id=classification_result.object_id,
            object_type_before=classification_result.object_type,
            object_type_after=classification_result.object_type,
            location=classification_result.location,
            text_preview=classification_result.original_text,
            property_name=rule.target_property,
            before="",
            after="",
            rule_id=rule.rule_id,
            status="unresolved",
            reason=f"unsupported table property: {rule.target_property}",
        )

    before_value = property_accessor.get(table)
    try:
        target_value = property_accessor.parse(rule.target_value)
    except ValueError as exc:
        return _report_row(
            object_id=classification_result.object_id,
            object_type_before=classification_result.object_type,
            object_type_after=classification_result.object_type,
            location=classification_result.location,
            text_preview=classification_result.original_text,
            property_name=rule.target_property,
            before=property_accessor.format(before_value, rule.target_value),
            after=rule.target_value,
            rule_id=rule.rule_id,
            status="unresolved",
            reason=str(exc),
        )

    property_accessor.set(table, target_value)
    after_value = property_accessor.get(table)
    status: ReportStatus = "modified" if before_value != after_value else "unchanged"
    return _report_row(
        object_id=classification_result.object_id,
        object_type_before=classification_result.object_type,
        object_type_after=classification_result.object_type,
        location=classification_result.location,
        text_preview=classification_result.original_text,
        property_name=rule.target_property,
        before=property_accessor.format(before_value, rule.target_value),
        after=property_accessor.format(after_value, rule.target_value),
        rule_id=rule.rule_id,
        status=status,
        reason="",
    )


def _matched_paragraph_rules(
    paragraph: ParsedBodyParagraph,
    rule_set: RuleSet,
) -> list[_MatchedParagraphRule]:
    matched_rules: list[_MatchedParagraphRule] = []
    for rule in rule_set.paragraph_rules:
        candidate = _paragraph_rule_candidate(rule, paragraph)
        if candidate is None:
            continue
        matched_rules.append(
            _MatchedParagraphRule(
                rule_id=rule.rule_id,
                source_family="paragraph",
                priority=rule.priority,
                match_kind=candidate.match_kind,
                match_type=rule.match_type,
                match_value=rule.match_value,
                target_property=rule.target_property,
                target_value=rule.target_value,
            )
        )

    for rule in rule_set.numbering_rules:
        if not _numbering_rule_matches(rule, paragraph):
            continue
        matched_rules.append(
            _MatchedParagraphRule(
                rule_id=rule.rule_id,
                source_family="numbering",
                priority=rule.priority,
                match_kind="numbering",
                match_type=rule.match_type,
                match_value=rule.match_value,
                target_property=rule.target_property,
                target_value=rule.target_value,
            )
        )

    return matched_rules


def _matched_header_paragraph_rules(
    object_class: str,
    rule_set: RuleSet,
) -> list[_MatchedParagraphRule]:
    matched_rules: list[_MatchedParagraphRule] = []
    for rule in rule_set.paragraph_rules:
        if rule.match_type != "class" or rule.match_value != object_class:
            continue
        matched_rules.append(
            _MatchedParagraphRule(
                rule_id=rule.rule_id,
                source_family="header_class",
                priority=rule.priority,
                match_kind="structural",
                match_type=rule.match_type,
                match_value=rule.match_value,
                target_property=rule.target_property,
                target_value=rule.target_value,
            )
        )
    return matched_rules


def _matched_table_rules(
    table: ParsedBodyTable,
    rule_set: RuleSet,
) -> list[_MatchedTableRule]:
    flattened_text = "\n".join("\t".join(cell for cell in row) for row in table.rows)
    matched_rules: list[_MatchedTableRule] = []
    for rule in rule_set.table_rules:
        candidate = _table_rule_candidate(rule, flattened_text)
        if candidate is None:
            continue
        matched_rules.append(
            _MatchedTableRule(
                rule_id=rule.rule_id,
                priority=rule.priority,
                match_kind=candidate.match_kind,
                match_type=rule.match_type,
                match_value=rule.match_value,
                target_property=rule.target_property,
                target_value=rule.target_value,
            )
        )
    return matched_rules


def _compatible_paragraph_rules(
    *,
    matched_rules: list[_MatchedParagraphRule],
    classification_result,
) -> list[_MatchedParagraphRule]:
    if classification_result.status != "matched" or classification_result.matched_rule_id is None:
        return []

    winner_rule = next(
        (rule for rule in matched_rules if rule.rule_id == classification_result.matched_rule_id),
        None,
    )
    if winner_rule is None:
        return []

    return [
        rule
        for rule in matched_rules
        if rule.source_family == winner_rule.source_family
        and rule.match_kind == winner_rule.match_kind
        and rule.match_type == winner_rule.match_type
        and rule.match_value == winner_rule.match_value
    ]


def _compatible_table_rules(
    *,
    matched_rules: list[_MatchedTableRule],
    classification_result,
) -> list[_MatchedTableRule]:
    if classification_result.status != "matched" or classification_result.matched_rule_id is None:
        return []

    winner_rule = next(
        (rule for rule in matched_rules if rule.rule_id == classification_result.matched_rule_id),
        None,
    )
    if winner_rule is None:
        return []

    return [
        rule
        for rule in matched_rules
        if rule.match_kind == winner_rule.match_kind
        and rule.match_type == winner_rule.match_type
        and rule.match_value == winner_rule.match_value
    ]


def _body_content_maps(document: DocumentObject) -> tuple[dict[str, Paragraph], dict[str, Table]]:
    body_paragraphs: dict[str, Paragraph] = {}
    body_tables: dict[str, Table] = {}
    body_index = 0
    for element in _iter_body_elements(document):
        if isinstance(element, Paragraph):
            if not element.text.strip():
                continue
            body_paragraphs[f"body_items[{body_index}]"] = element
            body_index += 1
            continue
        body_tables[f"body_items[{body_index}]"] = element
        body_index += 1
    return body_paragraphs, body_tables


def _body_paragraph_map(document: DocumentObject) -> dict[str, Paragraph]:
    body_paragraphs, _ = _body_content_maps(document)
    return body_paragraphs


def _body_table_map(document: DocumentObject) -> dict[str, Table]:
    _, body_tables = _body_content_maps(document)
    return body_tables


def _header_paragraph_map(document: DocumentObject) -> dict[str, Paragraph]:
    mapped: dict[str, Paragraph] = {}
    seen_parts: set[int] = set()
    header_index = 0
    for section in document.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            part_key = id(header.part)
            if part_key in seen_parts:
                continue
            seen_parts.add(part_key)
            header_items: list[Paragraph] = []
            item_index = 0
            for element in header.iter_inner_content():
                if isinstance(element, Paragraph):
                    if not element.text.strip():
                        continue
                    header_items.append(element)
                    item_index += 1
                    continue
                item_index += 1
            if not header_items:
                continue
            for mapped_item_index, paragraph in enumerate(header_items):
                mapped[f"headers[{header_index}].items[{mapped_item_index}]"] = paragraph
            header_index += 1
    return mapped


def _header_location_indexes(location: str) -> tuple[int, int]:
    match = re.fullmatch(r"headers\[(\d+)\]\.items\[(\d+)\]", location)
    if match is None:
        raise ValueError(f"unsupported header location format: {location}")
    return int(match.group(1)), int(match.group(2))


def _footer_paragraph_map(document: DocumentObject) -> dict[str, Paragraph]:
    mapped: dict[str, Paragraph] = {}
    seen_parts: set[int] = set()
    footer_index = 0
    for section in document.sections:
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            part_key = id(footer.part)
            if part_key in seen_parts:
                continue
            seen_parts.add(part_key)
            footer_items: list[Paragraph] = []
            for element in footer.iter_inner_content():
                if not isinstance(element, Paragraph):
                    continue
                if not element.text.strip():
                    continue
                footer_items.append(element)
            if not footer_items:
                continue
            for mapped_item_index, paragraph in enumerate(footer_items):
                mapped[f"footers[{footer_index}].items[{mapped_item_index}]"] = paragraph
            footer_index += 1
    return mapped


def _footer_location_indexes(location: str) -> tuple[int, int]:
    match = re.fullmatch(r"footers\[(\d+)\]\.items\[(\d+)\]", location)
    if match is None:
        raise ValueError(f"unsupported footer location format: {location}")
    return int(match.group(1)), int(match.group(2))


def _build_parsed_location_index(parsed) -> _ParsedLocationIndex:
    body_paragraphs: dict[str, ParsedBodyParagraph] = {}
    body_tables: dict[str, ParsedBodyTable] = {}
    for body_index, item in enumerate(parsed.body_items):
        location = f"body_items[{body_index}]"
        if isinstance(item, ParsedBodyParagraph):
            body_paragraphs[location] = item
        elif isinstance(item, ParsedBodyTable):
            body_tables[location] = item

    header_paragraphs: dict[str, ParsedBodyParagraph] = {}
    for header_index, header in enumerate(parsed.headers):
        for item_index, item in enumerate(header.items):
            if isinstance(item, ParsedBodyParagraph):
                header_paragraphs[f"headers[{header_index}].items[{item_index}]"] = item

    footer_paragraphs: dict[str, ParsedBodyParagraph] = {}
    for footer_index, footer in enumerate(parsed.footers):
        for item_index, item in enumerate(footer.items):
            if isinstance(item, ParsedBodyParagraph):
                footer_paragraphs[f"footers[{footer_index}].items[{item_index}]"] = item

    return _ParsedLocationIndex(
        body_paragraphs=body_paragraphs,
        body_tables=body_tables,
        header_paragraphs=header_paragraphs,
        footer_paragraphs=footer_paragraphs,
    )


def _build_document_runtime_index(document: DocumentObject) -> _DocumentRuntimeIndex:
    body_paragraphs, body_tables = _body_content_maps(document)
    return _DocumentRuntimeIndex(
        body_paragraphs=body_paragraphs,
        body_tables=body_tables,
        header_paragraphs=_header_paragraph_map(document),
        footer_paragraphs=_footer_paragraph_map(document),
    )


def _build_annotation_plan(report_rows: list[dict[str, str]]) -> _AnnotationPlan:
    section_rows: list[dict[str, str]] = []
    rows_by_location: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in report_rows:
        if row["status"] != "modified":
            continue
        if row["location"].startswith("sections[") or row["location"] == "document_settings":
            section_rows.append(row)
            continue
        rows_by_location[row["location"]].append(row)
    return _AnnotationPlan(
        section_rows=section_rows,
        rows_by_location=dict(rows_by_location),
    )


def _write_annotated_document(
    *,
    annotated_path: Path,
    annotation_plan: _AnnotationPlan,
    document: DocumentObject | None = None,
    runtime_index: _DocumentRuntimeIndex | None = None,
) -> None:
    if document is None:
        raise ValueError("document is required for annotated document writing")
    if runtime_index is None:
        runtime_index = _build_document_runtime_index(document)
    body_paragraphs = runtime_index.body_paragraphs
    body_tables = runtime_index.body_tables
    header_paragraphs = runtime_index.header_paragraphs
    footer_paragraphs = runtime_index.footer_paragraphs

    if annotation_plan.section_rows:
        _insert_document_annotation(
            document,
            _build_section_annotation(annotation_plan.section_rows),
        )

    for location, rows in annotation_plan.rows_by_location.items():
        textual_rows = [row for row in rows if _is_textual_property(row["property"])]
        annotation_rows = [row for row in rows if _is_annotation_property(row["property"])]

        if location in body_paragraphs:
            paragraph = body_paragraphs[location]
            if textual_rows:
                _annotate_paragraph_textual_changes(paragraph, textual_rows)
            if annotation_rows:
                _insert_note_after_paragraph(paragraph, _build_annotation_text(annotation_rows))
            continue

        if location in header_paragraphs:
            paragraph = header_paragraphs[location]
            if textual_rows:
                _annotate_paragraph_textual_changes(paragraph, textual_rows)
            if annotation_rows:
                _insert_note_after_paragraph(paragraph, _build_annotation_text(annotation_rows))
            continue

        if location in footer_paragraphs:
            paragraph = footer_paragraphs[location]
            if textual_rows:
                _annotate_paragraph_textual_changes(paragraph, textual_rows)
            if annotation_rows:
                _insert_note_after_paragraph(paragraph, _build_annotation_text(annotation_rows))
            continue

        if location in body_tables:
            table = body_tables[location]
            if textual_rows:
                _annotate_table_textual_changes(table, textual_rows)
            if annotation_rows:
                _insert_note_after_table(table, _build_annotation_text(annotation_rows))

    document.save(annotated_path)


def _mark_paragraph_red(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = _RED


def _mark_table_red(table: Table) -> None:
    for paragraph in _iter_table_paragraphs(table):
        _mark_paragraph_red(paragraph)


def _annotate_paragraph_textual_changes(
    paragraph: Paragraph,
    rows: list[dict[str, str]],
) -> None:
    inline_rows = [row for row in rows if _is_inline_textual_property(row["property"])]
    general_rows = [row for row in rows if not _is_inline_textual_property(row["property"])]
    if inline_rows:
        _mark_inline_textual_segments_red(paragraph, inline_rows)
    if general_rows:
        _mark_paragraph_red(paragraph)
    _append_annotation_run(paragraph, _build_textual_annotation_text(rows))


def _annotate_table_textual_changes(
    table: Table,
    rows: list[dict[str, str]],
) -> None:
    rows_by_paragraph: dict[int, tuple[Paragraph, list[dict[str, str]]]] = {}
    for row in rows:
        for paragraph in _table_paragraphs_for_property(table, row["property"]):
            if not paragraph.text.strip():
                continue
            record = rows_by_paragraph.setdefault(id(paragraph), (paragraph, []))
            record[1].append(row)

    for paragraph, paragraph_rows in rows_by_paragraph.values():
        _annotate_paragraph_textual_changes(paragraph, paragraph_rows)


def _insert_note_after_paragraph(paragraph: Paragraph, text: str) -> Paragraph:
    note = paragraph._parent.add_paragraph()
    _style_annotation_paragraph(note, text)
    paragraph._p.addnext(note._p)
    return note


def _insert_note_after_table(table: Table, text: str) -> Paragraph:
    note = table._parent.add_paragraph()
    _style_annotation_paragraph(note, text)
    table._tbl.addnext(note._p)
    return note


def _insert_document_annotation(document: DocumentObject, text: str) -> Paragraph:
    note = document.add_paragraph()
    _style_annotation_paragraph(note, text)
    first_paragraph = next((paragraph for paragraph in document.paragraphs), None)
    if first_paragraph is not None:
        first_paragraph._p.addprevious(note._p)
    return note


def _style_annotation_paragraph(paragraph: Paragraph, text: str) -> None:
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.color.rgb = _RED
    run.font.name = _ANNOTATION_FONT_NAME
    run.font.size = _ANNOTATION_FONT_SIZE
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), _ANNOTATION_FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), _ANNOTATION_FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), _ANNOTATION_FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), _ANNOTATION_FONT_NAME)


def _append_annotation_run(paragraph: Paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.font.color.rgb = _RED
    run.font.name = _ANNOTATION_FONT_NAME
    run.font.size = _ANNOTATION_FONT_SIZE
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), _ANNOTATION_FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), _ANNOTATION_FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), _ANNOTATION_FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), _ANNOTATION_FONT_NAME)


def _build_textual_annotation_text(rows: list[dict[str, str]]) -> str:
    ordered_rows = sorted(
        rows,
        key=lambda row: (
            0 if row["property"] == "font_name" else 1,
            row["rule_id"],
        ),
    )
    fragments = [_describe_textual_change(row) for row in ordered_rows]
    return "\u3010\u89c4\u8303\u5316\u8bf4\u660e\uff1a" + "\uff1b".join(fragments) + "\u3011"


def _build_textual_annotation_text(rows: list[dict[str, str]]) -> str:
    ordered_rows = sorted(
        rows,
        key=lambda row: (
            _textual_property_order(row["property"]),
            row["rule_id"],
        ),
    )
    fragments = [_describe_textual_change(row) for row in ordered_rows]
    return "\u3010\u89c4\u8303\u5316\u8bf4\u660e\uff1a" + "\uff1b".join(fragments) + "\u3011"


def _is_inline_textual_property(property_name: str) -> bool:
    return property_name.startswith("label_") or property_name.startswith("content_")


def _mark_inline_textual_segments_red(paragraph: Paragraph, rows: list[dict[str, str]]) -> None:
    label_run, content_run = _ensure_inline_segment_runs(paragraph)
    for row in rows:
        if row["property"].startswith("label_"):
            label_run.font.color.rgb = _RED
            continue
        if row["property"].startswith("content_") and content_run is not None:
            content_run.font.color.rgb = _RED


def _textual_property_order(property_name: str) -> int:
    order = [
        "header_row_font_name",
        "header_row_font_size",
        "header_row_bold",
        "header_row_alignment",
        "body_rows_font_name",
        "body_rows_font_size",
        "body_rows_bold",
        "body_rows_alignment",
        "label_font_name",
        "label_font_size",
        "label_bold",
        "content_font_name",
        "content_font_size",
        "content_bold",
        "font_name",
        "font_size",
        "bold",
        "alignment",
    ]
    if property_name in order:
        return order.index(property_name)
    return len(order)


def _is_textual_property(property_name: str) -> bool:
    if property_name in _TABLE_ANNOTATION_PROPERTIES:
        return False
    table_target = _parse_table_property_target(property_name)
    if table_target is not None:
        return table_target.base_property not in _TABLE_CELL_ANNOTATION_BASE_PROPERTIES
    return property_name in _TEXTUAL_PROPERTIES


def _is_annotation_property(property_name: str) -> bool:
    if property_name in _ANNOTATION_PROPERTIES or property_name in _TABLE_ANNOTATION_PROPERTIES:
        return True
    table_target = _parse_table_property_target(property_name)
    return table_target is not None and table_target.base_property in _TABLE_CELL_ANNOTATION_BASE_PROPERTIES


def _describe_table_target_position(target: _TablePropertyTarget) -> str:
    if target.selector_kind == "cell_range":
        return (
            f"\u7b2c{target.indices[0] + 1}\u884c\u5230\u7b2c{target.indices[1]}\u884c\u3001"
            f"\u7b2c{target.indices[2] + 1}\u5217\u5230\u7b2c{target.indices[3]}\u5217"
        )
    if target.selector_kind == "row":
        return f"\u7b2c{target.indices[0] + 1}\u884c"
    if target.selector_kind == "row_range":
        return f"\u7b2c{target.indices[0] + 1}\u884c\u5230\u7b2c{target.indices[1]}\u884c"
    if target.selector_kind == "column_by_header":
        return f"\u8868\u5934\u201c{target.header_name}\u201d\u5217"
    if target.selector_kind == "column_range":
        return f"\u7b2c{target.indices[0] + 1}\u5217\u5230\u7b2c{target.indices[1]}\u5217"
    if target.selector_kind == "column":
        return f"\u7b2c{target.indices[0] + 1}\u5217"
    return f"\u7b2c{target.indices[0] + 1}\u884c\u7b2c{target.indices[1] + 1}\u5217"


def _describe_annotation_change(row: dict[str, str]) -> str:
    property_name = row["property"]
    before = row["before"]
    after = row["after"]
    table_target = _parse_table_property_target(property_name)
    if table_target is not None and table_target.base_property in _TABLE_CELL_ANNOTATION_BASE_PROPERTIES:
        position_label = _describe_table_target_position(table_target)
        if table_target.base_property == "vertical_alignment":
            return (
                f"{position_label}\u5782\u76f4\u5bf9\u9f50\u539f\u4e3a {_prettify_annotation_value(_format_vertical_alignment(before, ''))}"
                f"\uff0c\u89c4\u8303\u4e3a {_prettify_annotation_value(_format_vertical_alignment(after, ''))}"
            )
        return (
            f"{position_label}\u8fb9\u6846\u539f\u4e3a {_prettify_annotation_value(_format_border_style(before, ''))}"
            f"\uff0c\u89c4\u8303\u4e3a {_prettify_annotation_value(_format_border_style(after, ''))}"
        )

    if property_name in {"vertical_alignment", "header_row_vertical_alignment", "body_rows_vertical_alignment"}:
        return (
            f"{_PROPERTY_LABELS.get(property_name, property_name)}\u539f\u4e3a {_prettify_annotation_value(_format_vertical_alignment(before, ''))}"
            f"\uff0c\u89c4\u8303\u4e3a {_prettify_annotation_value(_format_vertical_alignment(after, ''))}"
        )
    if property_name in {"border", "header_row_border", "body_rows_border"}:
        return (
            f"{_PROPERTY_LABELS.get(property_name, property_name)}\u539f\u4e3a {_prettify_annotation_value(_format_border_style(before, ''))}"
            f"\uff0c\u89c4\u8303\u4e3a {_prettify_annotation_value(_format_border_style(after, ''))}"
        )
    return (
        f"{_PROPERTY_LABELS.get(property_name, property_name)}\u539f\u4e3a {_prettify_annotation_value(before)}"
        f"\uff0c\u89c4\u8303\u4e3a {_prettify_annotation_value(after)}"
    )


def _build_annotation_text(rows: list[dict[str, str]]) -> str:
    ordered_rows = sorted(
        rows,
        key=lambda row: (
            _annotation_property_order(row["property"]),
            row["rule_id"],
        ),
    )
    fragments = [
        f"{_PROPERTY_LABELS.get(row['property'], row['property'])}鍘熶负 {_prettify_annotation_value(row['before'])}锛岃鑼冧负 {_prettify_annotation_value(row['after'])}"
        for row in ordered_rows
    ]
    return "[\u89c4\u8303\u5316\u6279\u6ce8] " + "\uff1b".join(fragments)


def _build_annotation_text(rows: list[dict[str, str]]) -> str:
    ordered_rows = sorted(
        rows,
        key=lambda row: (
            _annotation_property_order(row["property"]),
            row["rule_id"],
        ),
    )
    fragments = [_describe_annotation_change(row) for row in ordered_rows]
    return "[\u89c4\u8303\u5316\u6279\u6ce8] " + "\uff1b".join(fragments)


def _build_section_annotation(rows: list[dict[str, str]]) -> str:
    deduped: dict[str, dict[str, str]] = {}
    for row in rows:
        deduped.setdefault(row["property"], row)
    ordered_rows = sorted(
        deduped.values(),
        key=lambda row: _annotation_property_order(row["property"]),
    )
    fragments = [
        f"{_PROPERTY_LABELS.get(row['property'], row['property'])}\u5df2\u6309\u89c4\u8303\u8bbe\u4e3a {_prettify_annotation_value(row['after'])}"
        for row in ordered_rows
    ]
    return "[\u89c4\u8303\u5316\u6279\u6ce8] \u9875\u9762\u8bbe\u7f6e\uff1a" + "\uff1b".join(fragments)


def _describe_textual_change(row: dict[str, str]) -> str:
    property_name = row["property"]
    before = row["before"]
    after = row["after"]
    table_target = _parse_table_property_target(property_name)
    if table_target is not None:
        return _describe_targeted_table_change(table_target, before, after)
    if property_name == "font_name":
        return f"瀛椾綋鍘熶负 {_prettify_font_name_value(before)}锛岃皟鏁翠负 {_prettify_font_name_value(after)}"
    if property_name == "font_size":
        return f"瀛楀彿鍘熶负 {_prettify_font_size_value(before)}锛岃皟鏁翠负 {_prettify_font_size_value(after)}"
    return f"{_PROPERTY_LABELS.get(property_name, property_name)}鍘熶负 {_prettify_annotation_value(before)}锛岀幇璋冩暣涓?{_prettify_annotation_value(after)}"


def _describe_textual_change(row: dict[str, str]) -> str:
    property_name = row["property"]
    before = row["before"]
    after = row["after"]
    if property_name == "font_name":
        return f"瀛椾綋鍘熶负 {_prettify_font_name_value(before)}锛岃皟鏁翠负 {_prettify_font_name_value(after)}"
    if property_name == "font_size":
        return f"瀛楀彿鍘熶负 {_prettify_font_size_value(before)}锛岃皟鏁翠负 {_prettify_font_size_value(after)}"
    if property_name == "header_row_font_name":
        return f"琛ㄥご瀛椾綋鍘熶负 {_prettify_font_name_value(before)}锛岃皟鏁翠负 {_prettify_font_name_value(after)}"
    if property_name == "header_row_font_size":
        return f"琛ㄥご瀛楀彿鍘熶负 {_prettify_font_size_value(before)}锛岃皟鏁翠负 {_prettify_font_size_value(after)}"
    if property_name == "body_rows_font_name":
        return f"鏁版嵁琛屽瓧浣撳師涓?{_prettify_font_name_value(before)}锛岃皟鏁翠负 {_prettify_font_name_value(after)}"
    if property_name == "body_rows_font_size":
        return f"鏁版嵁琛屽瓧鍙峰師涓?{_prettify_font_size_value(before)}锛岃皟鏁翠负 {_prettify_font_size_value(after)}"
    if property_name == "label_font_name":
        return f"鏍囩瀛椾綋鍘熶负 {_prettify_font_name_value(before)}锛岃皟鏁翠负 {_prettify_font_name_value(after)}"
    if property_name == "label_font_size":
        return f"鏍囩瀛楀彿鍘熶负 {_prettify_font_size_value(before)}锛岃皟鏁翠负 {_prettify_font_size_value(after)}"
    if property_name == "content_font_name":
        return f"鍐呭瀛椾綋鍘熶负 {_prettify_font_name_value(before)}锛岃皟鏁翠负 {_prettify_font_name_value(after)}"
    if property_name == "content_font_size":
        return f"鍐呭瀛楀彿鍘熶负 {_prettify_font_size_value(before)}锛岃皟鏁翠负 {_prettify_font_size_value(after)}"
    if property_name == "label_bold":
        return f"鏍囩鍔犵矖鍘熶负 {_format_bool_display(before)}锛岃皟鏁翠负 {_format_bool_display(after)}"
    if property_name == "content_bold":
        return f"鍐呭鍔犵矖鍘熶负 {_format_bool_display(before)}锛岃皟鏁翠负 {_format_bool_display(after)}"
    return f"{_PROPERTY_LABELS.get(property_name, property_name)}鍘熶负 {_prettify_annotation_value(before)}锛岃皟鏁翠负 {_prettify_annotation_value(after)}"


def _annotation_property_order(property_name: str) -> int:
    order = [
        "first_line_indent",
        "hanging_indent",
        "space_before",
        "space_after",
        "line_spacing",
        "page_margin_top",
        "page_margin_bottom",
        "page_margin_left",
        "page_margin_right",
        "section_start_type",
        "page_number_format",
        "page_number_start",
        "footer_page_number_alignment",
        "different_first_page_header_footer",
        "odd_and_even_pages_header_footer",
    ]
    if property_name in order:
        return order.index(property_name)
    return len(order)


def _prettify_annotation_value(value: str) -> str:
    if value == "":
        return "\u672a\u663e\u5f0f\u8bbe\u7f6e\uff08\u7ee7\u627f\u6837\u5f0f\uff09"
    cm_match = re.fullmatch(r"(-?\d+(?:\.\d+)?)cm", value)
    if cm_match is not None:
        amount = float(cm_match.group(1))
        return f"{amount:.2f}".rstrip("0").rstrip(".") + "cm"

    pt_match = re.fullmatch(r"(-?\d+(?:\.\d+)?)pt", value)
    if pt_match is not None:
        amount = float(pt_match.group(1))
        return f"{amount:g}pt"

    return value


def _prettify_font_size_value(value: str) -> str:
    if value.startswith("mixed[") and value.endswith("]"):
        items = value[len("mixed[") : -1].split("|")
        return "娣峰悎[" + "|".join(_prettify_font_size_value(item) for item in items) + "]"

    pt_match = re.fullmatch(r"(-?\d+(?:\.\d+)?)pt", value)
    if pt_match is None:
        return value

    amount = float(pt_match.group(1))
    size_name = _font_size_name(amount)
    if size_name is None:
        return f"{amount:g}pt"
    return f"{size_name}\uff08{amount:g}pt\uff09"


def _prettify_font_name_value(value: str) -> str:
    normalized = value.strip()
    if not normalized:
        return "\u672a\u663e\u5f0f\u8bbe\u7f6e\uff08\u7ee7\u627f\u6837\u5f0f\uff09"
    if "|" in normalized:
        return "\u3001".join(_prettify_font_name_value(part) for part in normalized.split("|"))
    return _FONT_DISPLAY_NAMES.get(normalized, normalized)


def _parse_bool(value: str) -> bool:
    normalized = value.strip().lower()
    if normalized in {"true", "1", "yes", "\u662f"}:
        return True
    if normalized in {"false", "0", "no", "\u5426"}:
        return False
    raise ValueError(f"unsupported boolean value: {value}")


def _format_bool_display(value) -> str:
    if isinstance(value, str):
        normalized = value.strip().lower()
        if normalized in {"true", "1", "yes", "\u662f"}:
            return "\u662f"
        if normalized in {"false", "0", "no", "\u5426", ""}:
            return "\u5426"
    return "\u662f" if bool(value) else "\u5426"


def _parse_alignment(value: str) -> WD_ALIGN_PARAGRAPH:
    normalized = value.strip().lower()
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    if normalized not in mapping:
        raise ValueError(f"unsupported alignment value: {value}")
    return mapping[normalized]


def _format_alignment(value, _raw_target: str) -> str:
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        None: "left",
    }
    return mapping.get(value, str(value).lower())


def _parse_vertical_alignment(value: str) -> WD_CELL_VERTICAL_ALIGNMENT:
    normalized = value.strip().lower()
    mapping = {
        "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
        "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
        "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
        "both": WD_CELL_VERTICAL_ALIGNMENT.BOTH,
    }
    if normalized not in mapping:
        raise ValueError(f"unsupported vertical alignment value: {value}")
    return mapping[normalized]


def _format_vertical_alignment(value, _raw_target: str) -> str:
    mapping = {
        WD_CELL_VERTICAL_ALIGNMENT.TOP: "top",
        WD_CELL_VERTICAL_ALIGNMENT.CENTER: "center",
        WD_CELL_VERTICAL_ALIGNMENT.BOTTOM: "bottom",
        WD_CELL_VERTICAL_ALIGNMENT.BOTH: "both",
        None: "",
    }
    return mapping.get(value, str(value).lower())


def _parse_border_style(value: str) -> str:
    normalized = value.strip().lower()
    if normalized not in {"single", "none"}:
        raise ValueError(f"unsupported border value: {value}")
    return normalized


def _format_border_style(value, _raw_target: str) -> str:
    if value in {None, ""}:
        return ""
    normalized = str(value).strip().lower()
    if normalized in {"nil", "none"}:
        return "none"
    return normalized


def _parse_page_number_format(value: str) -> str:
    normalized = value.strip()
    allowed = {"decimal", "lowerRoman", "upperRoman", "lowerLetter", "upperLetter"}
    if normalized not in allowed:
        raise ValueError(f"unsupported page number format: {value}")
    return normalized


def _format_page_number_format(value, _raw_target: str) -> str:
    return "" if value in {None, ""} else str(value)


def _parse_page_number_start(value: str) -> int:
    try:
        parsed = int(value.strip())
    except ValueError as exc:
        raise ValueError(f"unsupported page number start value: {value}") from exc
    if parsed < 1:
        raise ValueError(f"page number start must be >= 1: {value}")
    return parsed


def _format_page_number_start(value, _raw_target: str) -> str:
    return "" if value in {None, ""} else str(value)


def _parse_section_start_type(value: str) -> WD_SECTION_START:
    normalized = value.strip().lower()
    mapping = {
        "new_page": WD_SECTION_START.NEW_PAGE,
        "odd_page": WD_SECTION_START.ODD_PAGE,
        "even_page": WD_SECTION_START.EVEN_PAGE,
        "continuous": WD_SECTION_START.CONTINUOUS,
        "new_column": WD_SECTION_START.NEW_COLUMN,
    }
    if normalized not in mapping:
        raise ValueError(f"unsupported section start type: {value}")
    return mapping[normalized]


def _format_section_start_type(value, _raw_target: str) -> str:
    mapping = {
        WD_SECTION_START.NEW_PAGE: "new_page",
        WD_SECTION_START.ODD_PAGE: "odd_page",
        WD_SECTION_START.EVEN_PAGE: "even_page",
        WD_SECTION_START.CONTINUOUS: "continuous",
        WD_SECTION_START.NEW_COLUMN: "new_column",
        None: "",
    }
    return mapping.get(value, str(value))


def _font_size_name(value_pt: float) -> str | None:
    named_sizes = (
        (26.0, "\u4e00\u53f7"),
        (24.0, "\u5c0f\u4e00"),
        (22.0, "\u4e8c\u53f7"),
        (18.0, "\u5c0f\u4e8c"),
        (16.0, "\u4e09\u53f7"),
        (15.0, "\u5c0f\u4e09"),
        (14.0, "\u56db\u53f7"),
        (12.0, "\u5c0f\u56db"),
        (10.5, "\u4e94\u53f7"),
        (9.0, "\u5c0f\u4e94"),
        (7.5, "\u516d\u53f7"),
        (6.5, "\u5c0f\u516d"),
        (5.5, "\u4e03\u53f7"),
        (5.0, "鍏彿"),
    )
    for amount, label in named_sizes:
        if abs(value_pt - amount) < 0.05:
            return label
    return None


@dataclass(frozen=True)
class _DocumentPropertyAccessor:
    target_kind: str
    get: callable
    set: callable
    parse: callable
    format: callable


def _document_property_accessor(property_name: str) -> _DocumentPropertyAccessor | None:
    accessors: dict[str, _DocumentPropertyAccessor] = {
        "page_margin_top": _DocumentPropertyAccessor(
            target_kind="section",
            get=lambda section: section.top_margin,
            set=lambda section, value: setattr(section, "top_margin", value),
            parse=_parse_measurement,
            format=lambda value, raw: _format_length(value, raw),
        ),
        "page_margin_bottom": _DocumentPropertyAccessor(
            target_kind="section",
            get=lambda section: section.bottom_margin,
            set=lambda section, value: setattr(section, "bottom_margin", value),
            parse=_parse_measurement,
            format=lambda value, raw: _format_length(value, raw),
        ),
        "page_margin_left": _DocumentPropertyAccessor(
            target_kind="section",
            get=lambda section: section.left_margin,
            set=lambda section, value: setattr(section, "left_margin", value),
            parse=_parse_measurement,
            format=lambda value, raw: _format_length(value, raw),
        ),
        "page_margin_right": _DocumentPropertyAccessor(
            target_kind="section",
            get=lambda section: section.right_margin,
            set=lambda section, value: setattr(section, "right_margin", value),
            parse=_parse_measurement,
            format=lambda value, raw: _format_length(value, raw),
        ),
        "section_start_type": _DocumentPropertyAccessor(
            target_kind="section",
            get=lambda section: section.start_type,
            set=lambda section, value: setattr(section, "start_type", value),
            parse=_parse_section_start_type,
            format=_format_section_start_type,
        ),
        "page_number_format": _DocumentPropertyAccessor(
            target_kind="section",
            get=_get_section_page_number_format,
            set=_set_section_page_number_format,
            parse=_parse_page_number_format,
            format=lambda value, _: _format_page_number_format(value, ""),
        ),
        "page_number_start": _DocumentPropertyAccessor(
            target_kind="section",
            get=_get_section_page_number_start,
            set=_set_section_page_number_start,
            parse=_parse_page_number_start,
            format=lambda value, _: _format_page_number_start(value, ""),
        ),
        "footer_page_number_alignment": _DocumentPropertyAccessor(
            target_kind="section",
            get=_get_footer_page_number_alignment,
            set=_set_footer_page_number_alignment,
            parse=_parse_alignment,
            format=_format_alignment,
        ),
        "different_first_page_header_footer": _DocumentPropertyAccessor(
            target_kind="section",
            get=lambda section: section.different_first_page_header_footer,
            set=lambda section, value: setattr(section, "different_first_page_header_footer", value),
            parse=_parse_bool,
            format=lambda value, _: _format_bool_display(value),
        ),
        "odd_and_even_pages_header_footer": _DocumentPropertyAccessor(
            target_kind="document",
            get=lambda document: document.settings.odd_and_even_pages_header_footer,
            set=lambda document, value: setattr(document.settings, "odd_and_even_pages_header_footer", value),
            parse=_parse_bool,
            format=lambda value, _: _format_bool_display(value),
        ),
    }
    return accessors.get(property_name)


@dataclass(frozen=True)
class _ParagraphPropertyAccessor:
    get: callable
    set: callable
    parse: callable
    format: callable


@dataclass(frozen=True)
class _TablePropertyAccessor:
    get: callable
    set: callable
    parse: callable
    format: callable


def _paragraph_property_accessor(property_name: str) -> _ParagraphPropertyAccessor | None:
    accessors: dict[str, _ParagraphPropertyAccessor] = {
        "font_name": _ParagraphPropertyAccessor(
            get=_get_paragraph_font_name,
            set=_set_paragraph_font_name,
            parse=lambda value: value,
            format=lambda value, _: value,
        ),
        "font_size": _ParagraphPropertyAccessor(
            get=lambda paragraph: _get_run_length_state(paragraph, "size"),
            set=lambda paragraph, value: _set_run_length(paragraph, "size", value),
            parse=_parse_measurement,
            format=_format_run_length_state,
        ),
        "label_font_name": _ParagraphPropertyAccessor(
            get=_get_label_font_name,
            set=_set_label_font_name,
            parse=lambda value: value,
            format=lambda value, _: value,
        ),
        "content_font_name": _ParagraphPropertyAccessor(
            get=_get_content_font_name,
            set=_set_content_font_name,
            parse=lambda value: value,
            format=lambda value, _: value,
        ),
        "label_font_size": _ParagraphPropertyAccessor(
            get=_get_label_font_size,
            set=_set_label_font_size,
            parse=_parse_measurement,
            format=_format_length,
        ),
        "content_font_size": _ParagraphPropertyAccessor(
            get=_get_content_font_size,
            set=_set_content_font_size,
            parse=_parse_measurement,
            format=_format_length,
        ),
        "label_bold": _ParagraphPropertyAccessor(
            get=_get_label_bold,
            set=_set_label_bold,
            parse=_parse_bool,
            format=lambda value, _: _format_bool_display(value),
        ),
        "content_bold": _ParagraphPropertyAccessor(
            get=_get_content_bold,
            set=_set_content_bold,
            parse=_parse_bool,
            format=lambda value, _: _format_bool_display(value),
        ),
        "first_line_indent": _ParagraphPropertyAccessor(
            get=lambda paragraph: paragraph.paragraph_format.first_line_indent,
            set=lambda paragraph, value: setattr(paragraph.paragraph_format, "first_line_indent", value),
            parse=_parse_measurement,
            format=_format_length,
        ),
        "hanging_indent": _ParagraphPropertyAccessor(
            get=_get_paragraph_hanging_indent,
            set=_set_paragraph_hanging_indent,
            parse=_parse_measurement,
            format=_format_length,
        ),
        "space_before": _ParagraphPropertyAccessor(
            get=lambda paragraph: paragraph.paragraph_format.space_before,
            set=lambda paragraph, value: setattr(paragraph.paragraph_format, "space_before", value),
            parse=_parse_measurement,
            format=_format_length,
        ),
        "space_after": _ParagraphPropertyAccessor(
            get=lambda paragraph: paragraph.paragraph_format.space_after,
            set=lambda paragraph, value: setattr(paragraph.paragraph_format, "space_after", value),
            parse=_parse_measurement,
            format=_format_length,
        ),
        "line_spacing": _ParagraphPropertyAccessor(
            get=lambda paragraph: paragraph.paragraph_format.line_spacing,
            set=lambda paragraph, value: setattr(paragraph.paragraph_format, "line_spacing", value),
            parse=_parse_line_spacing,
            format=_format_line_spacing,
        ),
    }
    return accessors.get(property_name)


def _table_property_accessor(property_name: str) -> _TablePropertyAccessor | None:
    accessors: dict[str, _TablePropertyAccessor] = {
        "font_name": _TablePropertyAccessor(
            get=_get_table_font_name,
            set=_set_table_font_name,
            parse=lambda value: value,
            format=lambda value, _: value,
        ),
        "font_size": _TablePropertyAccessor(
            get=lambda table: _get_table_run_length_state(table, "size"),
            set=lambda table, value: _set_table_run_length(table, "size", value),
            parse=_parse_measurement,
            format=_format_run_length_state,
        ),
        "header_row_font_name": _TablePropertyAccessor(
            get=lambda table: _get_table_font_name_in_rows(table, row_selector="header"),
            set=lambda table, value: _set_table_font_name_in_rows(table, value, row_selector="header"),
            parse=lambda value: value,
            format=lambda value, _: value,
        ),
        "header_row_font_size": _TablePropertyAccessor(
            get=lambda table: _get_table_run_length_state_in_rows(table, "size", row_selector="header"),
            set=lambda table, value: _set_table_run_length_in_rows(table, "size", value, row_selector="header"),
            parse=_parse_measurement,
            format=_format_run_length_state,
        ),
        "body_rows_font_name": _TablePropertyAccessor(
            get=lambda table: _get_table_font_name_in_rows(table, row_selector="body"),
            set=lambda table, value: _set_table_font_name_in_rows(table, value, row_selector="body"),
            parse=lambda value: value,
            format=lambda value, _: value,
        ),
        "body_rows_font_size": _TablePropertyAccessor(
            get=lambda table: _get_table_run_length_state_in_rows(table, "size", row_selector="body"),
            set=lambda table, value: _set_table_run_length_in_rows(table, "size", value, row_selector="body"),
            parse=_parse_measurement,
            format=_format_run_length_state,
        ),
    }
    return accessors.get(property_name)


def _get_paragraph_font_name(paragraph: Paragraph) -> str:
    fonts: list[str] = []
    seen_fonts: set[str] = set()
    for run in paragraph.runs:
        for font_name in _used_run_fonts(run):
            if font_name in seen_fonts:
                continue
            seen_fonts.add(font_name)
            fonts.append(font_name)
    return "|".join(fonts)


def _get_paragraph_hanging_indent(paragraph: Paragraph) -> Length:
    first_line_indent = paragraph.paragraph_format.first_line_indent
    if first_line_indent is None or int(first_line_indent) >= 0:
        return Cm(0)
    return Length(abs(int(first_line_indent)))


def _set_paragraph_hanging_indent(paragraph: Paragraph, value: Length) -> None:
    paragraph.paragraph_format.left_indent = value
    paragraph.paragraph_format.first_line_indent = Length(-int(value))


def _get_label_font_name(paragraph: Paragraph) -> str:
    segmented_runs = _inline_segment_runs_if_present(paragraph)
    if segmented_runs is not None:
        label_run, _ = segmented_runs
    else:
        label_run = next((run for run in paragraph.runs if run.text), None)
        if label_run is None:
            return ""
    return "|".join(_used_run_fonts(label_run))


def _set_label_font_name(paragraph: Paragraph, font_name: str) -> None:
    label_run, _ = _ensure_inline_segment_runs(paragraph)
    _set_run_font_name(label_run, font_name)


def _get_content_font_name(paragraph: Paragraph) -> str:
    segmented_runs = _inline_segment_runs_if_present(paragraph)
    if segmented_runs is not None:
        _, content_run = segmented_runs
    else:
        content_run = next((run for run in paragraph.runs if run.text), None)
    if content_run is None:
        return ""
    return "|".join(_used_run_fonts(content_run))


def _get_label_font_size(paragraph: Paragraph):
    segmented_runs = _inline_segment_runs_if_present(paragraph)
    if segmented_runs is not None:
        label_run, _ = segmented_runs
    else:
        label_run = next((run for run in paragraph.runs if run.text), None)
        if label_run is None:
            return None
    return label_run.font.size


def _set_label_font_size(paragraph: Paragraph, value: Length) -> None:
    label_run, _ = _ensure_inline_segment_runs(paragraph)
    label_run.font.size = value


def _get_content_font_size(paragraph: Paragraph):
    segmented_runs = _inline_segment_runs_if_present(paragraph)
    if segmented_runs is not None:
        _, content_run = segmented_runs
    else:
        content_run = next((run for run in paragraph.runs if run.text), None)
    if content_run is None:
        return None
    return content_run.font.size


def _set_content_font_size(paragraph: Paragraph, value: Length) -> None:
    _, content_run = _ensure_inline_segment_runs(paragraph)
    if content_run is None:
        return
    content_run.font.size = value


def _set_content_font_name(paragraph: Paragraph, font_name: str) -> None:
    _, content_run = _ensure_inline_segment_runs(paragraph)
    if content_run is None:
        return
    _set_run_font_name(content_run, font_name)


def _get_label_bold(paragraph: Paragraph) -> bool:
    segmented_runs = _inline_segment_runs_if_present(paragraph)
    if segmented_runs is not None:
        label_run, _ = segmented_runs
    else:
        label_run = next((run for run in paragraph.runs if run.text), None)
        if label_run is None:
            return False
    return bool(label_run.font.bold)


def _set_label_bold(paragraph: Paragraph, value: bool) -> None:
    label_run, _ = _ensure_inline_segment_runs(paragraph)
    label_run.font.bold = value


def _get_content_bold(paragraph: Paragraph) -> bool:
    segmented_runs = _inline_segment_runs_if_present(paragraph)
    if segmented_runs is not None:
        _, content_run = segmented_runs
    else:
        content_run = next((run for run in paragraph.runs if run.text), None)
    if content_run is None:
        return False
    return bool(content_run.font.bold)


def _set_content_bold(paragraph: Paragraph, value: bool) -> None:
    _, content_run = _ensure_inline_segment_runs(paragraph)
    if content_run is None:
        return
    content_run.font.bold = value


def _get_paragraph_bold_state(paragraph: Paragraph) -> bool:
    values = [bool(run.font.bold) for run in paragraph.runs if run.text]
    return any(values)


def _set_paragraph_bold(paragraph: Paragraph, value: bool) -> None:
    for run in paragraph.runs:
        if run.text:
            run.font.bold = value


def _get_paragraph_alignment(paragraph: Paragraph):
    return paragraph.alignment


def _set_paragraph_alignment(paragraph: Paragraph, value: WD_ALIGN_PARAGRAPH) -> None:
    paragraph.alignment = value


def _set_paragraph_font_name(paragraph: Paragraph, font_name: str) -> None:
    if font_name in _EAST_ASIAN_FONT_NAMES:
        for run in paragraph.runs:
            _set_run_script_aware_font_name(
                run,
                east_asian_font_name=font_name,
                western_font_name=_WESTERN_FONT_NAME,
            )
        return

    for run in paragraph.runs:
        _set_run_uniform_font_name(run, font_name)


def _get_table_font_name(table: Table) -> str:
    fonts: list[str] = []
    seen_fonts: set[str] = set()
    for paragraph in _iter_table_paragraphs(table):
        for run in paragraph.runs:
            for font_name in _used_run_fonts(run):
                if font_name in seen_fonts:
                    continue
                seen_fonts.add(font_name)
                fonts.append(font_name)
    return "|".join(fonts)


def _set_table_font_name(table: Table, font_name: str) -> None:
    for paragraph in _iter_table_paragraphs(table):
        _set_paragraph_font_name(paragraph, font_name)


def _get_table_font_name_in_rows(table: Table, *, row_selector: str) -> str:
    fonts: list[str] = []
    seen_fonts: set[str] = set()
    for paragraph in _iter_table_paragraphs_in_rows(table, row_selector=row_selector):
        for run in paragraph.runs:
            for font_name in _used_run_fonts(run):
                if font_name in seen_fonts:
                    continue
                seen_fonts.add(font_name)
                fonts.append(font_name)
    return "|".join(fonts)


def _set_table_font_name_in_rows(table: Table, font_name: str, *, row_selector: str) -> None:
    for paragraph in _iter_table_paragraphs_in_rows(table, row_selector=row_selector):
        _set_paragraph_font_name(paragraph, font_name)


def _ensure_inline_segment_runs(paragraph: Paragraph) -> tuple:
    inline_parts = _split_inline_paragraph_text(paragraph.text)
    if inline_parts is None:
        raise ValueError(f"paragraph does not match a supported inline-prefix rule: {paragraph.text!r}")

    label_text, content_text = inline_parts
    if len(paragraph.runs) == 2 and paragraph.runs[0].text == label_text and paragraph.runs[1].text == content_text:
        return paragraph.runs[0], paragraph.runs[1]
    if len(paragraph.runs) == 1 and paragraph.runs[0].text == label_text and content_text == "":
        return paragraph.runs[0], None

    _rebuild_inline_segment_runs(paragraph, label_text, content_text)
    if content_text == "":
        return paragraph.runs[0], None
    return paragraph.runs[0], paragraph.runs[1]


def _inline_segment_runs_if_present(paragraph: Paragraph) -> tuple | None:
    inline_parts = _split_inline_paragraph_text(paragraph.text)
    if inline_parts is None:
        return None
    label_text, content_text = inline_parts
    if len(paragraph.runs) == 2 and paragraph.runs[0].text == label_text and paragraph.runs[1].text == content_text:
        return paragraph.runs[0], paragraph.runs[1]
    if len(paragraph.runs) == 1 and paragraph.runs[0].text == label_text and content_text == "":
        return paragraph.runs[0], None
    return None


def _split_inline_paragraph_text(text: str) -> tuple[str, str] | None:
    for prefix in _INLINE_PREFIXES:
        if text.startswith(prefix):
            return prefix, text[len(prefix) :]
    return None


def _rebuild_inline_segment_runs(paragraph: Paragraph, label_text: str, content_text: str) -> None:
    source_run = next((run for run in paragraph.runs if run.text), None)
    source_rpr = None
    if source_run is not None and source_run._element.rPr is not None:
        source_rpr = deepcopy(source_run._element.rPr)

    for child in list(paragraph._p):
        if child.tag == qn("w:pPr"):
            continue
        paragraph._p.remove(child)

    if label_text:
        label_run = paragraph.add_run(label_text)
        if source_rpr is not None:
            label_run._element.insert(0, deepcopy(source_rpr))
    if content_text:
        content_run = paragraph.add_run(content_text)
        if source_rpr is not None:
            content_run._element.insert(0, deepcopy(source_rpr))


def _used_run_fonts(run) -> tuple[str, ...]:
    fonts: list[str] = []
    seen_fonts: set[str] = set()
    for bucket in _ordered_script_buckets(run.text):
        font_name = _run_font_for_bucket(run, bucket)
        if font_name in seen_fonts:
            continue
        seen_fonts.add(font_name)
        fonts.append(font_name)
    return tuple(fonts)


def _ordered_script_buckets(text: str) -> tuple[str, ...]:
    buckets: list[str] = []
    previous_bucket: str | None = None
    for character in text:
        bucket = _script_bucket(character)
        if bucket == previous_bucket:
            continue
        buckets.append(bucket)
        previous_bucket = bucket
    return tuple(buckets)


def _script_bucket(character: str) -> str:
    if _is_east_asian_character(character):
        return "east_asia"
    return "western"


def _is_east_asian_character(character: str) -> bool:
    code_point = ord(character)
    east_asian_ranges = (
        (0x2E80, 0x2EFF),
        (0x2F00, 0x2FDF),
        (0x3000, 0x303F),
        (0x3040, 0x30FF),
        (0x3100, 0x312F),
        (0x31A0, 0x31BF),
        (0x3400, 0x4DBF),
        (0x4E00, 0x9FFF),
        (0xF900, 0xFAFF),
        (0xFE30, 0xFE4F),
        (0xFF00, 0xFFEF),
    )
    return any(start <= code_point <= end for start, end in east_asian_ranges)


def _run_font_for_bucket(run, bucket: str) -> str:
    r_pr = run._element.rPr
    r_fonts = None if r_pr is None else r_pr.rFonts
    if bucket == "east_asia":
        if r_fonts is not None:
            east_asian = r_fonts.get(qn("w:eastAsia"))
            if east_asian:
                return east_asian
        return run.font.name or ""

    if r_fonts is not None:
        for slot in ("w:ascii", "w:hAnsi", "w:cs"):
            font_name = r_fonts.get(qn(slot))
            if font_name:
                return font_name
    return run.font.name or ""


def _set_run_uniform_font_name(run, font_name: str) -> None:
    run.font.name = font_name
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def _set_run_script_aware_font_name(
    run,
    *,
    east_asian_font_name: str,
    western_font_name: str,
) -> None:
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), western_font_name)
    r_fonts.set(qn("w:hAnsi"), western_font_name)
    r_fonts.set(qn("w:eastAsia"), east_asian_font_name)
    r_fonts.set(qn("w:cs"), western_font_name)


def _set_run_font_name(run, font_name: str) -> None:
    if font_name in _EAST_ASIAN_FONT_NAMES:
        _set_run_script_aware_font_name(
            run,
            east_asian_font_name=font_name,
            western_font_name=_WESTERN_FONT_NAME,
        )
        return
    _set_run_uniform_font_name(run, font_name)


def _get_run_length_state(paragraph: Paragraph, property_name: str) -> _RunLengthState:
    return _RunLengthState(
        values=tuple(getattr(run.font, property_name) for run in paragraph.runs)
    )


def _set_run_length(paragraph: Paragraph, property_name: str, value: Length) -> None:
    for run in paragraph.runs:
        setattr(run.font, property_name, value)


def _get_table_run_length_state(table: Table, property_name: str) -> _RunLengthState:
    values: list[Length | None] = []
    for paragraph in _iter_table_paragraphs(table):
        values.extend(getattr(run.font, property_name) for run in paragraph.runs)
    return _RunLengthState(values=tuple(values))


def _set_table_run_length(table: Table, property_name: str, value: Length) -> None:
    for paragraph in _iter_table_paragraphs(table):
        _set_run_length(paragraph, property_name, value)


def _get_table_run_length_state_in_rows(
    table: Table,
    property_name: str,
    *,
    row_selector: str,
) -> _RunLengthState:
    values: list[Length | None] = []
    for paragraph in _iter_table_paragraphs_in_rows(table, row_selector=row_selector):
        values.extend(getattr(run.font, property_name) for run in paragraph.runs)
    return _RunLengthState(values=tuple(values))


def _set_table_run_length_in_rows(
    table: Table,
    property_name: str,
    value: Length,
    *,
    row_selector: str,
) -> None:
    for paragraph in _iter_table_paragraphs_in_rows(table, row_selector=row_selector):
        _set_run_length(paragraph, property_name, value)


def _parse_line_spacing(value: str) -> float | Length:
    stripped = value.strip()
    if re.fullmatch(r"-?\d+(?:\.\d+)?", stripped):
        return float(stripped)
    return _parse_measurement(stripped)


def _format_line_spacing(value: object, raw_target: str) -> str:
    if value is None:
        return ""
    if isinstance(value, Length):
        return _format_length(value, raw_target)
    if isinstance(value, (int, float)):
        return f"{value:g}"
    return _format_length(value, raw_target)


def _parse_measurement(value: str) -> Length:
    match = re.fullmatch(r"\s*(-?\d+(?:\.\d+)?)\s*(cm|pt|in)\s*", value)
    if match is None:
        raise ValueError(f"unsupported measurement value: {value}")
    amount = float(match.group(1))
    unit = match.group(2)
    if unit == "cm":
        return Cm(amount)
    if unit == "pt":
        return Pt(amount)
    return Inches(amount)


def _format_length(value: object, raw_target: str) -> str:
    if value is None:
        return ""
    if not isinstance(value, Length):
        return str(value)
    unit_match = re.fullmatch(r"\s*-?\d+(?:\.\d+)?\s*(cm|pt|in)\s*", raw_target)
    unit = unit_match.group(1) if unit_match is not None else "pt"
    if unit == "cm":
        return f"{value.cm:g}cm"
    if unit == "in":
        return f"{value.inches:g}in"
    return f"{value.pt:g}pt"


def _format_run_length_state(value: object, raw_target: str) -> str:
    if not isinstance(value, _RunLengthState):
        return _format_length(value, raw_target)
    if not value.values:
        return ""

    formatted_values = [_format_length(item, raw_target) for item in value.values]
    if len(set(formatted_values)) == 1:
        return formatted_values[0]
    return f"mixed[{'|'.join(formatted_values)}]"


def _iter_table_paragraphs(table: Table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph


def _iter_table_cells(table: Table):
    for row in table.rows:
        for cell in row.cells:
            yield cell


def _iter_table_paragraphs_in_rows(table: Table, *, row_selector: str):
    if row_selector == "header":
        row_indexes = range(0, min(1, len(table.rows)))
    elif row_selector == "body":
        row_indexes = range(1, len(table.rows))
    else:
        raise ValueError(f"unsupported table row selector: {row_selector}")

    for row_index in row_indexes:
        for cell in table.rows[row_index].cells:
            for paragraph in cell.paragraphs:
                yield paragraph


def _iter_table_cells_in_rows(table: Table, *, row_selector: str):
    if row_selector == "header":
        row_indexes = range(0, min(1, len(table.rows)))
    elif row_selector == "body":
        row_indexes = range(1, len(table.rows))
    else:
        raise ValueError(f"unsupported table row selector: {row_selector}")

    for row_index in row_indexes:
        for cell in table.rows[row_index].cells:
            yield cell


def _table_paragraphs_for_property(table: Table, property_name: str):
    if property_name.startswith("header_row_"):
        yield from _iter_table_paragraphs_in_rows(table, row_selector="header")
        return
    if property_name.startswith("body_rows_"):
        yield from _iter_table_paragraphs_in_rows(table, row_selector="body")
        return
    yield from _iter_table_paragraphs(table)


def _parse_table_property_target(property_name: str) -> _TablePropertyTarget | None:
    match = _TABLE_SELECTOR_PROPERTY_RE.fullmatch(property_name)
    if match is None:
        return None

    selector_kind = match.group("selector_kind")
    first = int(match.group("first"))
    second = match.group("second")
    indices = (first,) if second is None else (first, int(second))
    if selector_kind == "column" and len(indices) != 1:
        return None
    if selector_kind == "cell" and len(indices) != 2:
        return None
    return _TablePropertyTarget(
        selector_kind=selector_kind,
        indices=indices,
        base_property=match.group("base"),
    )


def _iter_table_paragraphs_in_column(table: Table, column_index: int):
    for row in table.rows:
        if column_index >= len(row.cells):
            continue
        for paragraph in row.cells[column_index].paragraphs:
            yield paragraph


def _iter_table_cells_in_column(table: Table, column_index: int):
    for row in table.rows:
        if column_index >= len(row.cells):
            continue
        yield row.cells[column_index]


def _iter_table_paragraphs_in_row(table: Table, row_index: int):
    if row_index >= len(table.rows):
        return
    for cell in table.rows[row_index].cells:
        for paragraph in cell.paragraphs:
            yield paragraph


def _iter_table_cells_in_row(table: Table, row_index: int):
    if row_index >= len(table.rows):
        return
    for cell in table.rows[row_index].cells:
        yield cell


def _iter_table_paragraphs_in_cell(table: Table, row_index: int, column_index: int):
    if row_index >= len(table.rows):
        return
    row = table.rows[row_index]
    if column_index >= len(row.cells):
        return
    for paragraph in row.cells[column_index].paragraphs:
        yield paragraph


def _iter_table_cells_in_cell(table: Table, row_index: int, column_index: int):
    if row_index >= len(table.rows):
        return
    row = table.rows[row_index]
    if column_index >= len(row.cells):
        return
    yield row.cells[column_index]


def _iter_table_paragraphs_in_cell_range(
    table: Table,
    row_start: int,
    row_end: int,
    col_start: int,
    col_end: int,
):
    for row_index in range(row_start, row_end):
        if row_index >= len(table.rows):
            return
        row = table.rows[row_index]
        for column_index in range(col_start, col_end):
            if column_index >= len(row.cells):
                continue
            for paragraph in row.cells[column_index].paragraphs:
                yield paragraph


def _iter_table_cells_in_cell_range(
    table: Table,
    row_start: int,
    row_end: int,
    col_start: int,
    col_end: int,
):
    for row_index in range(row_start, row_end):
        if row_index >= len(table.rows):
            return
        row = table.rows[row_index]
        for column_index in range(col_start, col_end):
            if column_index >= len(row.cells):
                continue
            yield row.cells[column_index]


def _iter_table_paragraphs_for_target(table: Table, target: _TablePropertyTarget):
    if target.selector_kind == "column":
        yield from _iter_table_paragraphs_in_column(table, target.indices[0])
        return
    if target.selector_kind == "cell":
        yield from _iter_table_paragraphs_in_cell(table, target.indices[0], target.indices[1])
        return
    raise ValueError(f"unsupported table selector kind: {target.selector_kind}")


def _iter_table_cells_for_target(table: Table, target: _TablePropertyTarget):
    if target.selector_kind == "cell_range":
        row_start, row_end, col_start, col_end = target.indices
        yield from _iter_table_cells_in_cell_range(table, row_start, row_end, col_start, col_end)
        return
    if target.selector_kind == "row_range":
        start, end = target.indices
        for row_index in range(start, end):
            yield from _iter_table_cells_in_row(table, row_index)
        return
    if target.selector_kind == "row":
        yield from _iter_table_cells_in_row(table, target.indices[0])
        return
    if target.selector_kind == "column_range":
        start, end = target.indices
        for column_index in range(start, end):
            yield from _iter_table_cells_in_column(table, column_index)
        return
    if target.selector_kind == "column_by_header":
        column_index = _resolve_header_named_column_index(table, target.header_name or "")
        yield from _iter_table_cells_in_column(table, column_index)
        return
    if target.selector_kind == "column":
        yield from _iter_table_cells_in_column(table, target.indices[0])
        return
    if target.selector_kind == "cell":
        yield from _iter_table_cells_in_cell(table, target.indices[0], target.indices[1])
        return
    raise ValueError(f"unsupported table selector kind: {target.selector_kind}")


def _validate_table_target_exists(table: Table, target: _TablePropertyTarget) -> None:
    if target.selector_kind == "column":
        column_index = target.indices[0]
        if not table.rows:
            raise ValueError("table has no rows")
        if any(column_index >= len(row.cells) for row in table.rows):
            raise ValueError(f"column selector out of range: {column_index}")
        return

    if target.selector_kind == "cell":
        row_index, column_index = target.indices
        if row_index >= len(table.rows):
            raise ValueError(f"cell row selector out of range: {row_index}")
        if column_index >= len(table.rows[row_index].cells):
            raise ValueError(
                f"cell column selector out of range: row={row_index}, column={column_index}"
            )
        return

    raise ValueError(f"unsupported table selector kind: {target.selector_kind}")


def _get_targeted_table_font_name(table: Table, target: _TablePropertyTarget) -> str:
    _validate_table_target_exists(table, target)
    fonts: list[str] = []
    seen_fonts: set[str] = set()
    for paragraph in _iter_table_paragraphs_for_target(table, target):
        for run in paragraph.runs:
            for font_name in _used_run_fonts(run):
                if font_name in seen_fonts:
                    continue
                seen_fonts.add(font_name)
                fonts.append(font_name)
    return "|".join(fonts)


def _set_targeted_table_font_name(table: Table, target: _TablePropertyTarget, font_name: str) -> None:
    _validate_table_target_exists(table, target)
    for paragraph in _iter_table_paragraphs_for_target(table, target):
        _set_paragraph_font_name(paragraph, font_name)


def _get_targeted_table_run_length_state(
    table: Table,
    target: _TablePropertyTarget,
    property_name: str,
) -> _RunLengthState:
    _validate_table_target_exists(table, target)
    values: list[Length | None] = []
    for paragraph in _iter_table_paragraphs_for_target(table, target):
        values.extend(getattr(run.font, property_name) for run in paragraph.runs)
    return _RunLengthState(values=tuple(values))


def _set_targeted_table_run_length(
    table: Table,
    target: _TablePropertyTarget,
    property_name: str,
    value: Length,
) -> None:
    _validate_table_target_exists(table, target)
    for paragraph in _iter_table_paragraphs_for_target(table, target):
        _set_run_length(paragraph, property_name, value)


def _get_table_paragraph_bool_state(table: Table) -> bool:
    return any(_get_paragraph_bold_state(paragraph) for paragraph in _iter_table_paragraphs(table))


def _set_table_paragraph_bool_state(table: Table, value: bool) -> None:
    for paragraph in _iter_table_paragraphs(table):
        _set_paragraph_bold(paragraph, value)


def _get_table_paragraph_bool_state_in_rows(table: Table, *, row_selector: str) -> bool:
    return any(
        _get_paragraph_bold_state(paragraph)
        for paragraph in _iter_table_paragraphs_in_rows(table, row_selector=row_selector)
    )


def _set_table_paragraph_bool_state_in_rows(table: Table, value: bool, *, row_selector: str) -> None:
    for paragraph in _iter_table_paragraphs_in_rows(table, row_selector=row_selector):
        _set_paragraph_bold(paragraph, value)


def _get_targeted_table_paragraph_bool_state(table: Table, target: _TablePropertyTarget) -> bool:
    _validate_table_target_exists(table, target)
    return any(
        _get_paragraph_bold_state(paragraph)
        for paragraph in _iter_table_paragraphs_for_target(table, target)
    )


def _set_targeted_table_paragraph_bool_state(
    table: Table,
    target: _TablePropertyTarget,
    value: bool,
) -> None:
    _validate_table_target_exists(table, target)
    for paragraph in _iter_table_paragraphs_for_target(table, target):
        _set_paragraph_bold(paragraph, value)


def _get_table_alignment_state(table: Table):
    values = [paragraph.alignment for paragraph in _iter_table_paragraphs(table)]
    return values[0] if values and len(set(values)) == 1 else (values[0] if values else None)


def _set_table_alignment(table: Table, value: WD_ALIGN_PARAGRAPH) -> None:
    for paragraph in _iter_table_paragraphs(table):
        _set_paragraph_alignment(paragraph, value)


def _get_table_alignment_state_in_rows(table: Table, *, row_selector: str):
    values = [paragraph.alignment for paragraph in _iter_table_paragraphs_in_rows(table, row_selector=row_selector)]
    return values[0] if values and len(set(values)) == 1 else (values[0] if values else None)


def _set_table_alignment_in_rows(table: Table, value: WD_ALIGN_PARAGRAPH, *, row_selector: str) -> None:
    for paragraph in _iter_table_paragraphs_in_rows(table, row_selector=row_selector):
        _set_paragraph_alignment(paragraph, value)


def _get_targeted_table_alignment_state(table: Table, target: _TablePropertyTarget):
    _validate_table_target_exists(table, target)
    values = [paragraph.alignment for paragraph in _iter_table_paragraphs_for_target(table, target)]
    return values[0] if values and len(set(values)) == 1 else (values[0] if values else None)


def _set_targeted_table_alignment_state(
    table: Table,
    target: _TablePropertyTarget,
    value: WD_ALIGN_PARAGRAPH,
) -> None:
    _validate_table_target_exists(table, target)
    for paragraph in _iter_table_paragraphs_for_target(table, target):
        _set_paragraph_alignment(paragraph, value)


def _get_cell_vertical_alignment(cell: _Cell):
    return cell.vertical_alignment


def _set_cell_vertical_alignment(cell: _Cell, value: WD_CELL_VERTICAL_ALIGNMENT) -> None:
    cell.vertical_alignment = value


def _get_table_vertical_alignment_state(table: Table):
    values = [cell.vertical_alignment for cell in _iter_table_cells(table)]
    return values[0] if values and len(set(values)) == 1 else (values[0] if values else None)


def _set_table_vertical_alignment(table: Table, value: WD_CELL_VERTICAL_ALIGNMENT) -> None:
    for cell in _iter_table_cells(table):
        _set_cell_vertical_alignment(cell, value)


def _get_table_vertical_alignment_state_in_rows(table: Table, *, row_selector: str):
    values = [cell.vertical_alignment for cell in _iter_table_cells_in_rows(table, row_selector=row_selector)]
    return values[0] if values and len(set(values)) == 1 else (values[0] if values else None)


def _set_table_vertical_alignment_in_rows(
    table: Table,
    value: WD_CELL_VERTICAL_ALIGNMENT,
    *,
    row_selector: str,
) -> None:
    for cell in _iter_table_cells_in_rows(table, row_selector=row_selector):
        _set_cell_vertical_alignment(cell, value)


def _get_targeted_table_vertical_alignment_state(table: Table, target: _TablePropertyTarget):
    _validate_table_target_exists(table, target)
    values = [cell.vertical_alignment for cell in _iter_table_cells_for_target(table, target)]
    return values[0] if values and len(set(values)) == 1 else (values[0] if values else None)


def _set_targeted_table_vertical_alignment_state(
    table: Table,
    target: _TablePropertyTarget,
    value: WD_CELL_VERTICAL_ALIGNMENT,
) -> None:
    _validate_table_target_exists(table, target)
    for cell in _iter_table_cells_for_target(table, target):
        _set_cell_vertical_alignment(cell, value)


def _get_or_add_tc_borders(cell: _Cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    return borders


def _normalize_border_style(value: str | None) -> str:
    if value in {None, ""}:
        return ""
    normalized = str(value).strip().lower()
    if normalized in {"nil", "none"}:
        return "none"
    return normalized


def _get_cell_border_style(cell: _Cell) -> str:
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return ""
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        return ""
    values = []
    for edge_name in _TABLE_BORDER_EDGES:
        edge = borders.find(qn(f"w:{edge_name}"))
        values.append(_normalize_border_style("" if edge is None else edge.get(qn("w:val"))))
    unique_values = set(values)
    if len(unique_values) == 1:
        return values[0]
    return values[0] if values else ""


def _set_cell_border_style(cell: _Cell, value: str) -> None:
    borders = _get_or_add_tc_borders(cell)
    xml_value = "nil" if value == "none" else value
    for edge_name in _TABLE_BORDER_EDGES:
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), xml_value)
        if xml_value == "nil":
            for attr in ("w:sz", "w:space", "w:color"):
                if edge.get(qn(attr)) is not None:
                    del edge.attrib[qn(attr)]
        else:
            edge.set(qn("w:sz"), "4")
            edge.set(qn("w:space"), "0")
            edge.set(qn("w:color"), "auto")


def _get_table_border_state(table: Table) -> str:
    values = [_get_cell_border_style(cell) for cell in _iter_table_cells(table)]
    return values[0] if values and len(set(values)) == 1 else (values[0] if values else "")


def _set_table_border(table: Table, value: str) -> None:
    for cell in _iter_table_cells(table):
        _set_cell_border_style(cell, value)


def _get_table_border_state_in_rows(table: Table, *, row_selector: str) -> str:
    values = [_get_cell_border_style(cell) for cell in _iter_table_cells_in_rows(table, row_selector=row_selector)]
    return values[0] if values and len(set(values)) == 1 else (values[0] if values else "")


def _set_table_border_in_rows(table: Table, value: str, *, row_selector: str) -> None:
    for cell in _iter_table_cells_in_rows(table, row_selector=row_selector):
        _set_cell_border_style(cell, value)


def _get_targeted_table_border_state(table: Table, target: _TablePropertyTarget) -> str:
    _validate_table_target_exists(table, target)
    values = [_get_cell_border_style(cell) for cell in _iter_table_cells_for_target(table, target)]
    return values[0] if values and len(set(values)) == 1 else (values[0] if values else "")


def _set_targeted_table_border_state(table: Table, target: _TablePropertyTarget, value: str) -> None:
    _validate_table_target_exists(table, target)
    for cell in _iter_table_cells_for_target(table, target):
        _set_cell_border_style(cell, value)


def _get_or_add_pg_num_type(section: Section):
    pg_num_type = section._sectPr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        section._sectPr.append(pg_num_type)
    return pg_num_type


def _get_section_page_number_format(section: Section) -> str:
    pg_num_type = section._sectPr.find(qn("w:pgNumType"))
    return "" if pg_num_type is None else (pg_num_type.get(qn("w:fmt")) or "")


def _set_section_page_number_format(section: Section, value: str) -> None:
    _ensure_section_footer_page_number(section)
    pg_num_type = _get_or_add_pg_num_type(section)
    pg_num_type.set(qn("w:fmt"), value)


def _get_section_page_number_start(section: Section):
    pg_num_type = section._sectPr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        return None
    start = pg_num_type.get(qn("w:start"))
    return None if start in {None, ""} else int(start)


def _set_section_page_number_start(section: Section, value: int) -> None:
    _ensure_section_footer_page_number(section)
    pg_num_type = _get_or_add_pg_num_type(section)
    pg_num_type.set(qn("w:start"), str(value))


def _footer_contains_page_field(section: Section) -> bool:
    for paragraph in section.footer.paragraphs:
        if " PAGE " in paragraph._p.xml or "PAGE" in paragraph._p.xml:
            return True
    return False


def _append_page_field(paragraph: Paragraph) -> None:
    begin = paragraph.add_run()
    begin_fld = OxmlElement("w:fldChar")
    begin_fld.set(qn("w:fldCharType"), "begin")
    begin._r.append(begin_fld)

    instr = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = " PAGE "
    instr._r.append(instr_text)

    separate = paragraph.add_run()
    separate_fld = OxmlElement("w:fldChar")
    separate_fld.set(qn("w:fldCharType"), "separate")
    separate._r.append(separate_fld)

    text_run = paragraph.add_run("1")
    text_run.font.name = _WESTERN_FONT_NAME

    end = paragraph.add_run()
    end_fld = OxmlElement("w:fldChar")
    end_fld.set(qn("w:fldCharType"), "end")
    end._r.append(end_fld)


def _ensure_section_footer_page_number(section: Section) -> Paragraph:
    section.footer.is_linked_to_previous = False
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    if not _footer_contains_page_field(section):
        if paragraph.text.strip():
            paragraph = footer.add_paragraph()
        _append_page_field(paragraph)
    return paragraph


def _get_footer_page_number_alignment(section: Section):
    paragraph = section.footer.paragraphs[0] if section.footer.paragraphs else None
    return None if paragraph is None else paragraph.alignment


def _set_footer_page_number_alignment(section: Section, value: WD_ALIGN_PARAGRAPH) -> None:
    paragraph = _ensure_section_footer_page_number(section)
    paragraph.alignment = value


def _table_property_accessor(property_name: str) -> _TablePropertyAccessor | None:
    target = _parse_table_property_target(property_name)
    if target is not None:
        if target.base_property == "font_name":
            return _TablePropertyAccessor(
                get=lambda table: _get_targeted_table_font_name(table, target),
                set=lambda table, value: _set_targeted_table_font_name(table, target, value),
                parse=lambda value: value,
                format=lambda value, _: value,
            )
        if target.base_property == "font_size":
            return _TablePropertyAccessor(
                get=lambda table: _get_targeted_table_run_length_state(table, target, "size"),
                set=lambda table, value: _set_targeted_table_run_length(table, target, "size", value),
                parse=_parse_measurement,
                format=_format_run_length_state,
            )
        if target.base_property == "bold":
            return _TablePropertyAccessor(
                get=lambda table: _get_targeted_table_paragraph_bool_state(table, target),
                set=lambda table, value: _set_targeted_table_paragraph_bool_state(table, target, value),
                parse=_parse_bool,
                format=lambda value, _: _format_bool_display(value),
            )
        if target.base_property == "alignment":
            return _TablePropertyAccessor(
                get=lambda table: _get_targeted_table_alignment_state(table, target),
                set=lambda table, value: _set_targeted_table_alignment_state(table, target, value),
                parse=_parse_alignment,
                format=_format_alignment,
            )
        if target.base_property == "vertical_alignment":
            return _TablePropertyAccessor(
                get=lambda table: _get_targeted_table_vertical_alignment_state(table, target),
                set=lambda table, value: _set_targeted_table_vertical_alignment_state(table, target, value),
                parse=_parse_vertical_alignment,
                format=_format_vertical_alignment,
            )
        if target.base_property == "border":
            return _TablePropertyAccessor(
                get=lambda table: _get_targeted_table_border_state(table, target),
                set=lambda table, value: _set_targeted_table_border_state(table, target, value),
                parse=_parse_border_style,
                format=_format_border_style,
            )
        return None

    accessors: dict[str, _TablePropertyAccessor] = {
        "font_name": _TablePropertyAccessor(
            get=_get_table_font_name,
            set=_set_table_font_name,
            parse=lambda value: value,
            format=lambda value, _: value,
        ),
        "font_size": _TablePropertyAccessor(
            get=lambda table: _get_table_run_length_state(table, "size"),
            set=lambda table, value: _set_table_run_length(table, "size", value),
            parse=_parse_measurement,
            format=_format_run_length_state,
        ),
        "bold": _TablePropertyAccessor(
            get=_get_table_paragraph_bool_state,
            set=_set_table_paragraph_bool_state,
            parse=_parse_bool,
            format=lambda value, _: _format_bool_display(value),
        ),
        "alignment": _TablePropertyAccessor(
            get=_get_table_alignment_state,
            set=_set_table_alignment,
            parse=_parse_alignment,
            format=_format_alignment,
        ),
        "vertical_alignment": _TablePropertyAccessor(
            get=_get_table_vertical_alignment_state,
            set=_set_table_vertical_alignment,
            parse=_parse_vertical_alignment,
            format=_format_vertical_alignment,
        ),
        "border": _TablePropertyAccessor(
            get=_get_table_border_state,
            set=_set_table_border,
            parse=_parse_border_style,
            format=_format_border_style,
        ),
        "header_row_font_name": _TablePropertyAccessor(
            get=lambda table: _get_table_font_name_in_rows(table, row_selector="header"),
            set=lambda table, value: _set_table_font_name_in_rows(table, value, row_selector="header"),
            parse=lambda value: value,
            format=lambda value, _: value,
        ),
        "header_row_font_size": _TablePropertyAccessor(
            get=lambda table: _get_table_run_length_state_in_rows(table, "size", row_selector="header"),
            set=lambda table, value: _set_table_run_length_in_rows(table, "size", value, row_selector="header"),
            parse=_parse_measurement,
            format=_format_run_length_state,
        ),
        "header_row_bold": _TablePropertyAccessor(
            get=lambda table: _get_table_paragraph_bool_state_in_rows(table, row_selector="header"),
            set=lambda table, value: _set_table_paragraph_bool_state_in_rows(table, value, row_selector="header"),
            parse=_parse_bool,
            format=lambda value, _: _format_bool_display(value),
        ),
        "header_row_alignment": _TablePropertyAccessor(
            get=lambda table: _get_table_alignment_state_in_rows(table, row_selector="header"),
            set=lambda table, value: _set_table_alignment_in_rows(table, value, row_selector="header"),
            parse=_parse_alignment,
            format=_format_alignment,
        ),
        "header_row_vertical_alignment": _TablePropertyAccessor(
            get=lambda table: _get_table_vertical_alignment_state_in_rows(table, row_selector="header"),
            set=lambda table, value: _set_table_vertical_alignment_in_rows(table, value, row_selector="header"),
            parse=_parse_vertical_alignment,
            format=_format_vertical_alignment,
        ),
        "header_row_border": _TablePropertyAccessor(
            get=lambda table: _get_table_border_state_in_rows(table, row_selector="header"),
            set=lambda table, value: _set_table_border_in_rows(table, value, row_selector="header"),
            parse=_parse_border_style,
            format=_format_border_style,
        ),
        "body_rows_font_name": _TablePropertyAccessor(
            get=lambda table: _get_table_font_name_in_rows(table, row_selector="body"),
            set=lambda table, value: _set_table_font_name_in_rows(table, value, row_selector="body"),
            parse=lambda value: value,
            format=lambda value, _: value,
        ),
        "body_rows_font_size": _TablePropertyAccessor(
            get=lambda table: _get_table_run_length_state_in_rows(table, "size", row_selector="body"),
            set=lambda table, value: _set_table_run_length_in_rows(table, "size", value, row_selector="body"),
            parse=_parse_measurement,
            format=_format_run_length_state,
        ),
        "body_rows_bold": _TablePropertyAccessor(
            get=lambda table: _get_table_paragraph_bool_state_in_rows(table, row_selector="body"),
            set=lambda table, value: _set_table_paragraph_bool_state_in_rows(table, value, row_selector="body"),
            parse=_parse_bool,
            format=lambda value, _: _format_bool_display(value),
        ),
        "body_rows_alignment": _TablePropertyAccessor(
            get=lambda table: _get_table_alignment_state_in_rows(table, row_selector="body"),
            set=lambda table, value: _set_table_alignment_in_rows(table, value, row_selector="body"),
            parse=_parse_alignment,
            format=_format_alignment,
        ),
        "body_rows_vertical_alignment": _TablePropertyAccessor(
            get=lambda table: _get_table_vertical_alignment_state_in_rows(table, row_selector="body"),
            set=lambda table, value: _set_table_vertical_alignment_in_rows(table, value, row_selector="body"),
            parse=_parse_vertical_alignment,
            format=_format_vertical_alignment,
        ),
        "body_rows_border": _TablePropertyAccessor(
            get=lambda table: _get_table_border_state_in_rows(table, row_selector="body"),
            set=lambda table, value: _set_table_border_in_rows(table, value, row_selector="body"),
            parse=_parse_border_style,
            format=_format_border_style,
        ),
    }
    return accessors.get(property_name)


def _table_paragraphs_for_property(table: Table, property_name: str):
    target = _parse_table_property_target(property_name)
    if target is not None:
        yield from _iter_table_paragraphs_for_target(table, target)
        return
    if property_name.startswith("header_row_"):
        yield from _iter_table_paragraphs_in_rows(table, row_selector="header")
        return
    if property_name.startswith("body_rows_"):
        yield from _iter_table_paragraphs_in_rows(table, row_selector="body")
        return
    yield from _iter_table_paragraphs(table)


def _apply_single_table_rule(
    *,
    table: Table,
    classification_result,
    rule: _MatchedTableRule,
) -> dict[str, str]:
    property_accessor = _table_property_accessor(rule.target_property)
    if property_accessor is None:
        return _report_row(
            object_id=classification_result.object_id,
            object_type_before=classification_result.object_type,
            object_type_after=classification_result.object_type,
            location=classification_result.location,
            text_preview=classification_result.original_text,
            property_name=rule.target_property,
            before="",
            after="",
            rule_id=rule.rule_id,
            status="unresolved",
            reason=f"unsupported table property: {rule.target_property}",
        )

    try:
        before_value = property_accessor.get(table)
        target_value = property_accessor.parse(rule.target_value)
        property_accessor.set(table, target_value)
        after_value = property_accessor.get(table)
    except ValueError as exc:
        return _report_row(
            object_id=classification_result.object_id,
            object_type_before=classification_result.object_type,
            object_type_after=classification_result.object_type,
            location=classification_result.location,
            text_preview=classification_result.original_text,
            property_name=rule.target_property,
            before="",
            after="",
            rule_id=rule.rule_id,
            status="unresolved",
            reason=str(exc),
        )

    status: ReportStatus = "modified" if before_value != after_value else "unchanged"
    return _report_row(
        object_id=classification_result.object_id,
        object_type_before=classification_result.object_type,
        object_type_after=classification_result.object_type,
        location=classification_result.location,
        text_preview=classification_result.original_text,
        property_name=rule.target_property,
        before=property_accessor.format(before_value, rule.target_value),
        after=property_accessor.format(after_value, rule.target_value),
        rule_id=rule.rule_id,
        status=status,
        reason="applied normalization rule" if status == "modified" else "already normalized",
    )


def _describe_textual_change(row: dict[str, str]) -> str:
    property_name = row["property"]
    before = row["before"]
    after = row["after"]
    table_target = _parse_table_property_target(property_name)
    if table_target is not None:
        if table_target.selector_kind == "column":
            position_label = f"\u7b2c{table_target.indices[0] + 1}\u5217"
        else:
            position_label = (
                f"\u7b2c{table_target.indices[0] + 1}\u884c\u7b2c{table_target.indices[1] + 1}\u5217"
            )
        if table_target.base_property == "font_name":
            return (
                f"{position_label}\u5b57\u4f53\u539f\u4e3a {_prettify_font_name_value(before)}"
                f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_name_value(after)}"
            )
        if table_target.base_property == "font_size":
            return (
                f"{position_label}\u5b57\u53f7\u539f\u4e3a {_prettify_font_size_value(before)}"
                f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_size_value(after)}"
            )
        if table_target.base_property == "bold":
            return (
                f"{position_label}\u52a0\u7c97\u539f\u4e3a {_format_bool_display(before)}"
                f"\uff0c\u8c03\u6574\u4e3a {_format_bool_display(after)}"
            )
        return (
            f"{position_label}\u5bf9\u9f50\u539f\u4e3a {_format_alignment(before, '')}"
            f"\uff0c\u8c03\u6574\u4e3a {_format_alignment(after, '')}"
        )

    if property_name == "font_name":
        return f"鐎涙ぞ缍嬮崢鐔惰礋 {_prettify_font_name_value(before)}閿涘矁鐨熼弫缈犺礋 {_prettify_font_name_value(after)}"
    if property_name == "font_size":
        return f"鐎涙褰块崢鐔惰礋 {_prettify_font_size_value(before)}閿涘矁鐨熼弫缈犺礋 {_prettify_font_size_value(after)}"
    if property_name == "header_row_font_name":
        return f"鐞涖劌銇旂€涙ぞ缍嬮崢鐔惰礋 {_prettify_font_name_value(before)}閿涘矁鐨熼弫缈犺礋 {_prettify_font_name_value(after)}"
    if property_name == "header_row_font_size":
        return f"鐞涖劌銇旂€涙褰块崢鐔惰礋 {_prettify_font_size_value(before)}閿涘矁鐨熼弫缈犺礋 {_prettify_font_size_value(after)}"
    if property_name == "body_rows_font_name":
        return f"閺佺増宓佺悰灞界摟娴ｆ挸甯稉?{_prettify_font_name_value(before)}閿涘矁鐨熼弫缈犺礋 {_prettify_font_name_value(after)}"
    if property_name == "body_rows_font_size":
        return f"閺佺増宓佺悰灞界摟閸欏嘲甯稉?{_prettify_font_size_value(before)}閿涘矁鐨熼弫缈犺礋 {_prettify_font_size_value(after)}"
    if property_name == "label_font_name":
        return f"閺嶅洨顒风€涙ぞ缍嬮崢鐔惰礋 {_prettify_font_name_value(before)}閿涘矁鐨熼弫缈犺礋 {_prettify_font_name_value(after)}"
    if property_name == "label_font_size":
        return f"閺嶅洨顒风€涙褰块崢鐔惰礋 {_prettify_font_size_value(before)}閿涘矁鐨熼弫缈犺礋 {_prettify_font_size_value(after)}"
    if property_name == "content_font_name":
        return f"閸愬懎顔愮€涙ぞ缍嬮崢鐔惰礋 {_prettify_font_name_value(before)}閿涘矁鐨熼弫缈犺礋 {_prettify_font_name_value(after)}"
    if property_name == "content_font_size":
        return f"閸愬懎顔愮€涙褰块崢鐔惰礋 {_prettify_font_size_value(before)}閿涘矁鐨熼弫缈犺礋 {_prettify_font_size_value(after)}"
    if property_name == "label_bold":
        return f"閺嶅洨顒烽崝鐘电煐閸樼喍璐?{_format_bool_display(before)}閿涘矁鐨熼弫缈犺礋 {_format_bool_display(after)}"
    if property_name == "content_bold":
        return f"閸愬懎顔愰崝鐘电煐閸樼喍璐?{_format_bool_display(before)}閿涘矁鐨熼弫缈犺礋 {_format_bool_display(after)}"
    return f"{_PROPERTY_LABELS.get(property_name, property_name)}閸樼喍璐?{_prettify_annotation_value(before)}閿涘矁鐨熼弫缈犺礋 {_prettify_annotation_value(after)}"


def _describe_textual_change(row: dict[str, str]) -> str:
    property_name = row["property"]
    before = row["before"]
    after = row["after"]
    table_target = _parse_table_property_target(property_name)
    if table_target is not None:
        if table_target.selector_kind == "cell_range":
            position_label = (
                f"\u7b2c{table_target.indices[0] + 1}\u884c\u5230\u7b2c{table_target.indices[1]}\u884c\u3001"
                f"\u7b2c{table_target.indices[2] + 1}\u5217\u5230\u7b2c{table_target.indices[3]}\u5217"
            )
        elif table_target.selector_kind == "row":
            position_label = f"\u7b2c{table_target.indices[0] + 1}\u884c"
        elif table_target.selector_kind == "row_range":
            position_label = (
                f"\u7b2c{table_target.indices[0] + 1}\u884c\u5230\u7b2c{table_target.indices[1]}\u884c"
            )
        elif table_target.selector_kind == "column_by_header":
            position_label = f"\u8868\u5934\u201c{table_target.header_name}\u201d\u5217"
        elif table_target.selector_kind == "column_range":
            position_label = (
                f"\u7b2c{table_target.indices[0] + 1}\u5217\u5230\u7b2c{table_target.indices[1]}\u5217"
            )
        elif table_target.selector_kind == "column":
            position_label = f"\u7b2c{table_target.indices[0] + 1}\u5217"
        else:
            position_label = (
                f"\u7b2c{table_target.indices[0] + 1}\u884c\u7b2c{table_target.indices[1] + 1}\u5217"
            )
        if table_target.base_property == "font_name":
            return (
                f"{position_label}\u5b57\u4f53\u539f\u4e3a {_prettify_font_name_value(before)}"
                f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_name_value(after)}"
            )
        if table_target.base_property == "font_size":
            return (
                f"{position_label}\u5b57\u53f7\u539f\u4e3a {_prettify_font_size_value(before)}"
                f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_size_value(after)}"
            )
        if table_target.base_property == "bold":
            return (
                f"{position_label}\u52a0\u7c97\u539f\u4e3a {_format_bool_display(before)}"
                f"\uff0c\u8c03\u6574\u4e3a {_format_bool_display(after)}"
            )
        return (
            f"{position_label}\u5bf9\u9f50\u539f\u4e3a {_format_alignment(before, '')}"
            f"\uff0c\u8c03\u6574\u4e3a {_format_alignment(after, '')}"
        )

    if property_name == "font_name":
        return (
            f"\u5b57\u4f53\u539f\u4e3a {_prettify_font_name_value(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_name_value(after)}"
        )
    if property_name == "font_size":
        return (
            f"\u5b57\u53f7\u539f\u4e3a {_prettify_font_size_value(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_size_value(after)}"
        )
    if property_name == "bold":
        return (
            f"\u52a0\u7c97\u539f\u4e3a {_format_bool_display(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_format_bool_display(after)}"
        )
    if property_name == "alignment":
        return (
            f"\u5bf9\u9f50\u539f\u4e3a {_format_alignment(before, '')}"
            f"\uff0c\u8c03\u6574\u4e3a {_format_alignment(after, '')}"
        )
    if property_name == "header_row_font_name":
        return (
            f"\u8868\u5934\u5b57\u4f53\u539f\u4e3a {_prettify_font_name_value(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_name_value(after)}"
        )
    if property_name == "header_row_font_size":
        return (
            f"\u8868\u5934\u5b57\u53f7\u539f\u4e3a {_prettify_font_size_value(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_size_value(after)}"
        )
    if property_name == "header_row_bold":
        return (
            f"\u8868\u5934\u52a0\u7c97\u539f\u4e3a {_format_bool_display(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_format_bool_display(after)}"
        )
    if property_name == "header_row_alignment":
        return (
            f"\u8868\u5934\u5bf9\u9f50\u539f\u4e3a {_format_alignment(before, '')}"
            f"\uff0c\u8c03\u6574\u4e3a {_format_alignment(after, '')}"
        )
    if property_name == "body_rows_font_name":
        return (
            f"\u6570\u636e\u884c\u5b57\u4f53\u539f\u4e3a {_prettify_font_name_value(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_name_value(after)}"
        )
    if property_name == "body_rows_font_size":
        return (
            f"\u6570\u636e\u884c\u5b57\u53f7\u539f\u4e3a {_prettify_font_size_value(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_size_value(after)}"
        )
    if property_name == "body_rows_bold":
        return (
            f"\u6570\u636e\u884c\u52a0\u7c97\u539f\u4e3a {_format_bool_display(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_format_bool_display(after)}"
        )
    if property_name == "body_rows_alignment":
        return (
            f"\u6570\u636e\u884c\u5bf9\u9f50\u539f\u4e3a {_format_alignment(before, '')}"
            f"\uff0c\u8c03\u6574\u4e3a {_format_alignment(after, '')}"
        )
    if property_name == "label_font_name":
        return (
            f"\u6807\u7b7e\u5b57\u4f53\u539f\u4e3a {_prettify_font_name_value(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_name_value(after)}"
        )
    if property_name == "label_font_size":
        return (
            f"\u6807\u7b7e\u5b57\u53f7\u539f\u4e3a {_prettify_font_size_value(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_size_value(after)}"
        )
    if property_name == "content_font_name":
        return (
            f"\u5185\u5bb9\u5b57\u4f53\u539f\u4e3a {_prettify_font_name_value(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_name_value(after)}"
        )
    if property_name == "content_font_size":
        return (
            f"\u5185\u5bb9\u5b57\u53f7\u539f\u4e3a {_prettify_font_size_value(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_size_value(after)}"
        )
    if property_name == "label_bold":
        return (
            f"\u6807\u7b7e\u52a0\u7c97\u539f\u4e3a {_format_bool_display(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_format_bool_display(after)}"
        )
    if property_name == "content_bold":
        return (
            f"\u5185\u5bb9\u52a0\u7c97\u539f\u4e3a {_format_bool_display(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_format_bool_display(after)}"
        )
    return (
        f"{_PROPERTY_LABELS.get(property_name, property_name)}"
        f"\u539f\u4e3a {_prettify_annotation_value(before)}"
        f"\uff0c\u8c03\u6574\u4e3a {_prettify_annotation_value(after)}"
    )


def _parse_table_property_target(property_name: str) -> _TablePropertyTarget | None:
    cell_range_match = _TABLE_CELL_RANGE_SELECTOR_PROPERTY_RE.fullmatch(property_name)
    if cell_range_match is not None:
        row_start = int(cell_range_match.group("row_start"))
        row_end = int(cell_range_match.group("row_end"))
        col_start = int(cell_range_match.group("col_start"))
        col_end = int(cell_range_match.group("col_end"))
        if row_end <= row_start or col_end <= col_start:
            return None
        return _TablePropertyTarget(
            selector_kind="cell_range",
            indices=(row_start, row_end, col_start, col_end),
            base_property=cell_range_match.group("base"),
        )

    row_range_match = _TABLE_ROW_RANGE_SELECTOR_PROPERTY_RE.fullmatch(property_name)
    if row_range_match is not None:
        start = int(row_range_match.group("start"))
        end = int(row_range_match.group("end"))
        if end <= start:
            return None
        return _TablePropertyTarget(
            selector_kind="row_range",
            indices=(start, end),
            base_property=row_range_match.group("base"),
        )

    row_match = _TABLE_ROW_SELECTOR_PROPERTY_RE.fullmatch(property_name)
    if row_match is not None:
        return _TablePropertyTarget(
            selector_kind="row",
            indices=(int(row_match.group("row")),),
            base_property=row_match.group("base"),
        )

    range_match = _TABLE_COLUMN_RANGE_SELECTOR_PROPERTY_RE.fullmatch(property_name)
    if range_match is not None:
        start = int(range_match.group("start"))
        end = int(range_match.group("end"))
        if end <= start:
            return None
        return _TablePropertyTarget(
            selector_kind="column_range",
            indices=(start, end),
            base_property=range_match.group("base"),
        )

    header_match = _TABLE_HEADER_SELECTOR_PROPERTY_RE.fullmatch(property_name)
    if header_match is not None:
        return _TablePropertyTarget(
            selector_kind="column_by_header",
            indices=(),
            base_property=header_match.group("base"),
            header_name=header_match.group("header"),
        )

    match = _TABLE_SELECTOR_PROPERTY_RE.fullmatch(property_name)
    if match is None:
        return None

    selector_kind = match.group("selector_kind")
    first = int(match.group("first"))
    second = match.group("second")
    indices = (first,) if second is None else (first, int(second))
    if selector_kind == "column" and len(indices) != 1:
        return None
    if selector_kind == "cell" and len(indices) != 2:
        return None
    return _TablePropertyTarget(
        selector_kind=selector_kind,
        indices=indices,
        base_property=match.group("base"),
    )


def _resolve_header_named_column_index(table: Table, header_name: str) -> int:
    if not table.rows:
        raise ValueError("table has no rows")
    header_row = table.rows[0]
    for column_index, cell in enumerate(header_row.cells):
        text = "\n".join(paragraph.text.strip() for paragraph in cell.paragraphs).strip()
        if text == header_name:
            return column_index
    raise ValueError(f"header selector not found: {header_name}")


def _iter_table_paragraphs_for_target(table: Table, target: _TablePropertyTarget):
    if target.selector_kind == "cell_range":
        row_start, row_end, col_start, col_end = target.indices
        yield from _iter_table_paragraphs_in_cell_range(
            table,
            row_start,
            row_end,
            col_start,
            col_end,
        )
        return
    if target.selector_kind == "row_range":
        start, end = target.indices
        for row_index in range(start, end):
            yield from _iter_table_paragraphs_in_row(table, row_index)
        return
    if target.selector_kind == "row":
        yield from _iter_table_paragraphs_in_row(table, target.indices[0])
        return
    if target.selector_kind == "column_range":
        start, end = target.indices
        for column_index in range(start, end):
            yield from _iter_table_paragraphs_in_column(table, column_index)
        return
    if target.selector_kind == "column_by_header":
        column_index = _resolve_header_named_column_index(table, target.header_name or "")
        yield from _iter_table_paragraphs_in_column(table, column_index)
        return
    if target.selector_kind == "column":
        yield from _iter_table_paragraphs_in_column(table, target.indices[0])
        return
    if target.selector_kind == "cell":
        yield from _iter_table_paragraphs_in_cell(table, target.indices[0], target.indices[1])
        return
    raise ValueError(f"unsupported table selector kind: {target.selector_kind}")


def _validate_table_target_exists(table: Table, target: _TablePropertyTarget) -> None:
    if target.selector_kind == "cell_range":
        row_start, row_end, col_start, col_end = target.indices
        if not table.rows:
            raise ValueError("table has no rows")
        if row_start < 0 or row_end <= row_start or col_start < 0 or col_end <= col_start:
            raise ValueError(
                f"invalid cell range selector: rows={row_start}:{row_end}, cols={col_start}:{col_end}"
            )
        if row_end > len(table.rows):
            raise ValueError(
                f"cell range row selector out of range: rows={row_start}:{row_end}"
            )
        if any(col_end > len(row.cells) for row in table.rows[row_start:row_end]):
            raise ValueError(
                f"cell range column selector out of range: cols={col_start}:{col_end}"
            )
        return
    if target.selector_kind == "row_range":
        start, end = target.indices
        if not table.rows:
            raise ValueError("table has no rows")
        if start < 0 or end <= start:
            raise ValueError(f"invalid row range selector: {start}:{end}")
        if end > len(table.rows):
            raise ValueError(f"row range selector out of range: {start}:{end}")
        return
    if target.selector_kind == "row":
        row_index = target.indices[0]
        if row_index >= len(table.rows):
            raise ValueError(f"row selector out of range: {row_index}")
        return
    if target.selector_kind == "column_range":
        start, end = target.indices
        if not table.rows:
            raise ValueError("table has no rows")
        if start < 0 or end <= start:
            raise ValueError(f"invalid column range selector: {start}:{end}")
        if any(end > len(row.cells) for row in table.rows):
            raise ValueError(f"column range selector out of range: {start}:{end}")
        return
    if target.selector_kind == "column_by_header":
        _resolve_header_named_column_index(table, target.header_name or "")
        return
    if target.selector_kind == "column":
        column_index = target.indices[0]
        if not table.rows:
            raise ValueError("table has no rows")
        if any(column_index >= len(row.cells) for row in table.rows):
            raise ValueError(f"column selector out of range: {column_index}")
        return
    if target.selector_kind == "cell":
        row_index, column_index = target.indices
        if row_index >= len(table.rows):
            raise ValueError(f"cell row selector out of range: {row_index}")
        if column_index >= len(table.rows[row_index].cells):
            raise ValueError(
                f"cell column selector out of range: row={row_index}, column={column_index}"
            )
        return
    raise ValueError(f"unsupported table selector kind: {target.selector_kind}")


def _describe_textual_change(row: dict[str, str]) -> str:
    property_name = row["property"]
    before = row["before"]
    after = row["after"]
    table_target = _parse_table_property_target(property_name)
    if table_target is not None:
        if table_target.selector_kind == "cell_range":
            position_label = (
                f"\u7b2c{table_target.indices[0] + 1}\u884c\u5230\u7b2c{table_target.indices[1]}\u884c\u3001"
                f"\u7b2c{table_target.indices[2] + 1}\u5217\u5230\u7b2c{table_target.indices[3]}\u5217"
            )
        elif table_target.selector_kind == "row":
            position_label = f"\u7b2c{table_target.indices[0] + 1}\u884c"
        elif table_target.selector_kind == "row_range":
            position_label = (
                f"\u7b2c{table_target.indices[0] + 1}\u884c\u5230\u7b2c{table_target.indices[1]}\u884c"
            )
        elif table_target.selector_kind == "column_by_header":
            position_label = f"\u8868\u5934\u201c{table_target.header_name}\u201d\u5217"
        elif table_target.selector_kind == "column_range":
            position_label = (
                f"\u7b2c{table_target.indices[0] + 1}\u5217\u5230\u7b2c{table_target.indices[1]}\u5217"
            )
        elif table_target.selector_kind == "column":
            position_label = f"\u7b2c{table_target.indices[0] + 1}\u5217"
        else:
            position_label = (
                f"\u7b2c{table_target.indices[0] + 1}\u884c\u7b2c{table_target.indices[1] + 1}\u5217"
            )
        if table_target.base_property == "font_name":
            return (
                f"{position_label}\u5b57\u4f53\u539f\u4e3a {_prettify_font_name_value(before)}"
                f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_name_value(after)}"
            )
        if table_target.base_property == "font_size":
            return (
                f"{position_label}\u5b57\u53f7\u539f\u4e3a {_prettify_font_size_value(before)}"
                f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_size_value(after)}"
            )
        if table_target.base_property == "bold":
            return (
                f"{position_label}\u52a0\u7c97\u539f\u4e3a {_format_bool_display(before)}"
                f"\uff0c\u8c03\u6574\u4e3a {_format_bool_display(after)}"
            )
        return (
            f"{position_label}\u5bf9\u9f50\u539f\u4e3a {_format_alignment(before, '')}"
            f"\uff0c\u8c03\u6574\u4e3a {_format_alignment(after, '')}"
        )

    if property_name == "font_name":
        return (
            f"\u5b57\u4f53\u539f\u4e3a {_prettify_font_name_value(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_name_value(after)}"
        )
    if property_name == "font_size":
        return (
            f"\u5b57\u53f7\u539f\u4e3a {_prettify_font_size_value(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_size_value(after)}"
        )
    if property_name == "header_row_font_name":
        return (
            f"\u8868\u5934\u5b57\u4f53\u539f\u4e3a {_prettify_font_name_value(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_name_value(after)}"
        )
    if property_name == "header_row_font_size":
        return (
            f"\u8868\u5934\u5b57\u53f7\u539f\u4e3a {_prettify_font_size_value(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_size_value(after)}"
        )
    if property_name == "body_rows_font_name":
        return (
            f"\u6570\u636e\u884c\u5b57\u4f53\u539f\u4e3a {_prettify_font_name_value(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_name_value(after)}"
        )
    if property_name == "body_rows_font_size":
        return (
            f"\u6570\u636e\u884c\u5b57\u53f7\u539f\u4e3a {_prettify_font_size_value(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_size_value(after)}"
        )
    if property_name == "header_row_bold":
        return (
            f"\u8868\u5934\u52a0\u7c97\u539f\u4e3a {_format_bool_display(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_format_bool_display(after)}"
        )
    if property_name == "body_rows_bold":
        return (
            f"\u6570\u636e\u884c\u52a0\u7c97\u539f\u4e3a {_format_bool_display(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_format_bool_display(after)}"
        )
    if property_name == "header_row_alignment":
        return (
            f"\u8868\u5934\u5bf9\u9f50\u539f\u4e3a {_format_alignment(before, '')}"
            f"\uff0c\u8c03\u6574\u4e3a {_format_alignment(after, '')}"
        )
    if property_name == "body_rows_alignment":
        return (
            f"\u6570\u636e\u884c\u5bf9\u9f50\u539f\u4e3a {_format_alignment(before, '')}"
            f"\uff0c\u8c03\u6574\u4e3a {_format_alignment(after, '')}"
        )
    if property_name == "label_font_name":
        return (
            f"\u6807\u7b7e\u5b57\u4f53\u539f\u4e3a {_prettify_font_name_value(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_name_value(after)}"
        )
    if property_name == "label_font_size":
        return (
            f"\u6807\u7b7e\u5b57\u53f7\u539f\u4e3a {_prettify_font_size_value(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_size_value(after)}"
        )
    if property_name == "content_font_name":
        return (
            f"\u5185\u5bb9\u5b57\u4f53\u539f\u4e3a {_prettify_font_name_value(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_name_value(after)}"
        )
    if property_name == "content_font_size":
        return (
            f"\u5185\u5bb9\u5b57\u53f7\u539f\u4e3a {_prettify_font_size_value(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_prettify_font_size_value(after)}"
        )
    if property_name == "label_bold":
        return (
            f"\u6807\u7b7e\u52a0\u7c97\u539f\u4e3a {_format_bool_display(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_format_bool_display(after)}"
        )
    if property_name == "content_bold":
        return (
            f"\u5185\u5bb9\u52a0\u7c97\u539f\u4e3a {_format_bool_display(before)}"
            f"\uff0c\u8c03\u6574\u4e3a {_format_bool_display(after)}"
        )
    return (
        f"{_PROPERTY_LABELS.get(property_name, property_name)}"
        f"\u539f\u4e3a {_prettify_annotation_value(before)}"
        f"\uff0c\u8c03\u6574\u4e3a {_prettify_annotation_value(after)}"
    )


def _special_object_target_object_type(
    rules: list[SpecialObjectRule],
    rule_id: str | None,
) -> str | None:
    if rule_id is None:
        return None
    for rule in rules:
        if rule.rule_id == rule_id:
            return rule.target_object_type
    return None


def _report_row(
    *,
    object_id: str,
    object_type_before: str,
    object_type_after: str,
    location: str,
    text_preview: str,
    property_name: str,
    before: str,
    after: str,
    rule_id: str,
    status: ReportStatus,
    reason: str,
) -> dict[str, str]:
    return {
        "object_id": object_id,
        "object_type_before": object_type_before,
        "object_type_after": object_type_after,
        "location": location,
        "text_preview": text_preview,
        "property": property_name,
        "before": before,
        "after": after,
        "rule_id": rule_id,
        "status": status,
        "reason": reason,
    }
