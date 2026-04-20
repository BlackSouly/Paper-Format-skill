from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Cm, Pt


def build_sample_docx(path: Path) -> Path:
    document = Document()
    document.settings.odd_and_even_pages_header_footer = True

    section = document.sections[0]
    section.different_first_page_header_footer = True
    section.header.is_linked_to_previous = False
    section.header.paragraphs[0].text = "Shared header"
    section.first_page_header.is_linked_to_previous = False
    section.first_page_header.paragraphs[0].text = "Section 1 first-page header"
    section.even_page_header.is_linked_to_previous = False
    section.even_page_header.paragraphs[0].text = "Section 1 even-page header"

    document.add_paragraph("First body paragraph")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Cell 1"
    table.cell(0, 1).text = "Cell 2"
    table.cell(1, 0).text = "Cell 3"
    table.cell(1, 1).text = "Cell 4"

    document.add_paragraph("Second body paragraph")

    document.add_section(WD_SECTION.NEW_PAGE)

    document.save(path)
    return path


def build_header_footer_variant_docx(path: Path) -> Path:
    document = Document()
    document.settings.odd_and_even_pages_header_footer = True

    section = document.sections[0]
    section.different_first_page_header_footer = True

    section.header.is_linked_to_previous = False
    section.header.paragraphs[0].text = "Shared header"
    section.first_page_header.is_linked_to_previous = False
    section.first_page_header.paragraphs[0].text = "Section 1 first-page header"
    section.even_page_header.is_linked_to_previous = False
    section.even_page_header.paragraphs[0].text = "Section 1 even-page header"

    section.footer.is_linked_to_previous = False
    section.footer.paragraphs[0].text = "Shared footer"
    section.first_page_footer.is_linked_to_previous = False
    section.first_page_footer.paragraphs[0].text = "Section 1 first-page footer"
    section.even_page_footer.is_linked_to_previous = False
    section.even_page_footer.paragraphs[0].text = "Section 1 even-page footer"

    document.add_paragraph("Body paragraph")

    document.save(path)
    return path


def build_normalization_sample_docx(
    path: Path,
    *,
    include_blank_paragraph_before_body: bool = False,
) -> Path:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(5.0)
    section.bottom_margin = Cm(4.0)

    heading = document.add_paragraph()
    heading.style = "Heading 1"
    heading_run = heading.add_run("第一章 绪论")
    heading_run.font.name = "Calibri"
    heading_run.font.size = Pt(16)

    if include_blank_paragraph_before_body:
        document.add_paragraph("")

    body = document.add_paragraph()
    body.style = "Body Text"
    body.paragraph_format.first_line_indent = Cm(0)
    body.paragraph_format.space_after = Pt(18)
    body_run = body.add_run("Body paragraph text")
    body_run.font.name = "Arial"
    body_run.font.size = Pt(11)

    unresolved = document.add_paragraph()
    unresolved.style = "Normal"
    unresolved.paragraph_format.space_after = Pt(6)
    unresolved_run = unresolved.add_run("Unresolved paragraph text")
    unresolved_run.font.name = "Courier New"
    unresolved_run.font.size = Pt(10)

    document.save(path)
    return path


def build_mixed_run_font_size_docx(path: Path) -> Path:
    document = Document()

    paragraph = document.add_paragraph()
    paragraph.style = "Body Text"

    first_run = paragraph.add_run("Mixed ")
    first_run.font.name = "Arial"
    first_run.font.size = Pt(10)

    second_run = paragraph.add_run("run sizes")
    second_run.font.name = "Arial"
    second_run.font.size = Pt(14)

    document.save(path)
    return path


def build_header_and_table_normalization_docx(path: Path) -> Path:
    document = Document()
    section = document.sections[0]
    section.header.is_linked_to_previous = False

    header = section.header.paragraphs[0]
    header_run = header.add_run("页眉标题")
    header_run.font.name = "Calibri"
    header_run.font.size = Pt(9)

    document.add_paragraph("Body paragraph")

    table = document.add_table(rows=2, cols=2)
    values = ("项目一", "项目二", "项目三", "项目四")
    for cell, value in zip((cell for row in table.rows for cell in row.cells), values, strict=True):
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(value)
        run.font.name = "Calibri"
        run.font.size = Pt(9)

    document.save(path)
    return path


def build_sparse_header_normalization_docx(path: Path) -> Path:
    document = Document()
    document.settings.odd_and_even_pages_header_footer = True

    first_section = document.sections[0]
    first_section.header.is_linked_to_previous = False
    first_section.first_page_header.is_linked_to_previous = False
    first_section.even_page_header.is_linked_to_previous = False

    document.add_paragraph("Section zero body")
    second_section = document.add_section(WD_SECTION.NEW_PAGE)
    second_section.header.is_linked_to_previous = False
    header = second_section.header.paragraphs[0]
    header_run = header.add_run("后置页眉")
    header_run.font.name = "Calibri"
    header_run.font.size = Pt(9)
    document.add_paragraph("Section one body")

    document.save(path)
    return path


def build_mixed_script_font_name_docx(path: Path) -> Path:
    document = Document()

    paragraph = document.add_paragraph()
    paragraph.style = "Body Text"

    run = paragraph.add_run("中文ABC123，测试")
    run.font.name = "Arial"
    run.font.size = Pt(11)

    document.save(path)
    return path


def build_abstract_and_reference_docx(path: Path) -> Path:
    document = Document()

    abstract_paragraph = document.add_paragraph()
    abstract_paragraph.style = "Body Text"
    abstract_run = abstract_paragraph.add_run("【Abstract】Graduation thesis typesetting improves review efficiency.")
    abstract_run.font.size = Pt(11)

    reference_paragraph = document.add_paragraph()
    reference_paragraph.style = "Body Text"
    reference_run = reference_paragraph.add_run("明日科技.Java从入门到精通[M].北京：清华大学出版社,2016:10-564")
    reference_run.font.name = "Calibri"
    reference_run.font.size = Pt(11)

    document.save(path)
    return path


def build_inline_tag_docx(path: Path) -> Path:
    document = Document()

    paragraph = document.add_paragraph()
    paragraph.style = "Body Text"
    run = paragraph.add_run("【摘要】中文ABC")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = False

    document.save(path)
    return path


def build_section_page_number_docx(path: Path) -> Path:
    document = Document()

    first_section = document.sections[0]
    first_section.footer.is_linked_to_previous = False
    document.add_paragraph("Section one body")

    second_section = document.add_section(WD_SECTION.NEW_PAGE)
    second_section.footer.is_linked_to_previous = False
    document.add_paragraph("Section two body")

    document.save(path)
    return path
