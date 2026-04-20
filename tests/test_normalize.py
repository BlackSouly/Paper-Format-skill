from __future__ import annotations

import csv
import importlib.util
from pathlib import Path

import pytest
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from paper_format_normalizer.model import (
    DocumentRule,
    ParagraphRule,
    ReportSchemaField,
    RuleSet,
    SpecialObjectRule,
    TableRule,
)
from paper_format_normalizer.normalize import normalize_document
import paper_format_normalizer.normalize as normalize_module

_BUILDER_PATH = Path(__file__).resolve().parent / "fixtures" / "sample_docx_builder.py"
_BUILDER_SPEC = importlib.util.spec_from_file_location(
    "sample_docx_builder",
    _BUILDER_PATH,
)
if _BUILDER_SPEC is None or _BUILDER_SPEC.loader is None:
    raise RuntimeError(f"Unable to load fixture builder from {_BUILDER_PATH}")
_BUILDER_MODULE = importlib.util.module_from_spec(_BUILDER_SPEC)
_BUILDER_SPEC.loader.exec_module(_BUILDER_MODULE)
build_sample_docx = _BUILDER_MODULE.build_sample_docx
build_header_footer_variant_docx = _BUILDER_MODULE.build_header_footer_variant_docx
build_normalization_sample_docx = _BUILDER_MODULE.build_normalization_sample_docx
build_mixed_run_font_size_docx = _BUILDER_MODULE.build_mixed_run_font_size_docx
build_header_and_table_normalization_docx = _BUILDER_MODULE.build_header_and_table_normalization_docx
build_sparse_header_normalization_docx = _BUILDER_MODULE.build_sparse_header_normalization_docx
build_mixed_script_font_name_docx = _BUILDER_MODULE.build_mixed_script_font_name_docx
build_abstract_and_reference_docx = _BUILDER_MODULE.build_abstract_and_reference_docx
build_inline_tag_docx = _BUILDER_MODULE.build_inline_tag_docx
build_section_page_number_docx = _BUILDER_MODULE.build_section_page_number_docx


def _cell_border_values(cell) -> set[str]:
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return set()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        return set()

    values: set[str] = set()
    for edge_name in ("top", "left", "bottom", "right"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            values.add("")
            continue
        value = (edge.get(qn("w:val")) or "").lower()
        values.add("none" if value in {"nil", "none"} else value)
    return values


def _section_pg_num_type(section):
    return section._sectPr.find(qn("w:pgNumType"))


def _footer_has_page_field(section) -> bool:
    return " PAGE " in section.footer.paragraphs[0]._p.xml or "PAGE" in section.footer.paragraphs[0]._p.xml


def test_normalize_document_writes_non_destructive_copy(tmp_path: Path) -> None:
    input_path = build_normalization_sample_docx(tmp_path / "paper.docx")
    output_dir = tmp_path / "normalized"

    output_path, report_path, annotated_path = normalize_document(input_path, _rule_set(), output_dir)

    assert output_path == output_dir / "paper_规范化.docx"
    assert report_path == output_dir / "paper_规范化_修改报告.csv"
    assert input_path.exists()
    assert output_path.exists()
    assert report_path.exists()
    assert annotated_path.exists()
    assert annotated_path.name.endswith(".docx")

    source_document = Document(input_path)
    normalized_document = Document(output_path)

    assert source_document.sections[0].top_margin.cm == pytest.approx(5.0, abs=0.01)
    assert normalized_document.sections[0].top_margin.cm == pytest.approx(2.54, abs=0.01)


def test_normalize_document_can_apply_page_number_rules_to_all_sections(tmp_path: Path) -> None:
    input_path = build_section_page_number_docx(tmp_path / "sectioned.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_page_number_rules(),
        tmp_path / "normalized",
    )

    document = Document(output_path)
    first_section = document.sections[0]
    second_section = document.sections[1]

    for section in (first_section, second_section):
        pg_num_type = _section_pg_num_type(section)
        assert pg_num_type is not None
        assert pg_num_type.get(qn("w:fmt")) == "lowerRoman"
        assert pg_num_type.get(qn("w:start")) == "1"
        assert section.footer.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert _footer_has_page_field(section)

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    assert any(row["property"] == "page_number_format" and row["after"] == "lowerRoman" for row in rows)
    assert any(row["property"] == "page_number_start" and row["after"] == "1" for row in rows)
    assert any(
        row["property"] == "footer_page_number_alignment" and row["after"] == "center"
        for row in rows
    )


def test_normalize_document_reuses_body_mapping_per_document_instance(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    input_path = build_normalization_sample_docx(tmp_path / "paper.docx")
    call_count = 0
    original_iter = normalize_module._iter_body_elements

    def counting_iter(document):
        nonlocal call_count
        call_count += 1
        yield from original_iter(document)

    monkeypatch.setattr(normalize_module, "_iter_body_elements", counting_iter)

    normalize_document(input_path, _rule_set(), tmp_path / "normalized")

    assert call_count == 2


def test_normalize_document_does_not_reparse_header_or_footer_locations(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    input_path = build_header_footer_variant_docx(tmp_path / "header-footer.docx")

    def fail_header(*_args, **_kwargs):
        raise AssertionError("header location reparsing should not run")

    def fail_footer(*_args, **_kwargs):
        raise AssertionError("footer location reparsing should not run")

    monkeypatch.setattr(normalize_module, "_header_location_indexes", fail_header)
    monkeypatch.setattr(normalize_module, "_footer_location_indexes", fail_footer)

    normalize_document(
        input_path,
        _rule_set_with_variant_specific_header_rules(),
        tmp_path / "normalized",
    )


def test_build_annotation_plan_groups_modified_rows_once() -> None:
    report_rows = [
        {
            "location": "sections[0]",
            "status": "modified",
            "property": "page_margin_top",
            "rule_id": "DOC-1",
        },
        {
            "location": "body_items[0]",
            "status": "modified",
            "property": "font_name",
            "rule_id": "PAR-1",
        },
        {
            "location": "body_items[0]",
            "status": "modified",
            "property": "space_after",
            "rule_id": "PAR-2",
        },
        {
            "location": "body_items[1]",
            "status": "unchanged",
            "property": "font_name",
            "rule_id": "PAR-3",
        },
        {
            "location": "document_settings",
            "status": "modified",
            "property": "odd_and_even_pages_header_footer",
            "rule_id": "DOC-2",
        },
    ]

    plan = normalize_module._build_annotation_plan(report_rows)

    assert [row["rule_id"] for row in plan.section_rows] == ["DOC-1", "DOC-2"]
    assert list(plan.rows_by_location) == ["body_items[0]"]
    assert [row["rule_id"] for row in plan.rows_by_location["body_items[0]"]] == ["PAR-1", "PAR-2"]


def test_normalize_document_reuses_in_memory_document_for_annotated_output(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    input_path = build_normalization_sample_docx(tmp_path / "paper.docx")
    original_document = normalize_module.Document
    call_count = 0

    def counting_document(*args, **kwargs):
        nonlocal call_count
        call_count += 1
        return original_document(*args, **kwargs)

    monkeypatch.setattr(normalize_module, "Document", counting_document)

    normalize_document(input_path, _rule_set(), tmp_path / "normalized")

    assert call_count == 1


def test_normalize_document_can_apply_section_start_type_rule(tmp_path: Path) -> None:
    input_path = build_section_page_number_docx(tmp_path / "sectioned.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_section_start_type_rule(),
        tmp_path / "normalized",
    )

    document = Document(output_path)

    assert document.sections[0].start_type == WD_SECTION_START.ODD_PAGE
    assert document.sections[1].start_type == WD_SECTION_START.ODD_PAGE

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    assert any(
        row["property"] == "section_start_type" and row["after"] == "odd_page"
        for row in rows
    )


def test_annotated_document_adds_section_start_type_note(tmp_path: Path) -> None:
    input_path = build_section_page_number_docx(tmp_path / "sectioned.docx")

    _, _, annotated_path = normalize_document(
        input_path,
        _rule_set_with_section_start_type_rule(),
        tmp_path / "normalized",
    )

    document = Document(annotated_path)
    first_paragraph = document.paragraphs[0]

    assert "规范化批注" in first_paragraph.text
    assert "分节起始方式已按规范设为 odd_page" in first_paragraph.text


def test_normalize_document_can_apply_different_first_page_header_footer_rule(tmp_path: Path) -> None:
    input_path = build_section_page_number_docx(tmp_path / "sectioned.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_different_first_page_header_footer_rule(),
        tmp_path / "normalized",
    )

    document = Document(output_path)

    assert document.sections[0].different_first_page_header_footer is True
    assert document.sections[1].different_first_page_header_footer is True

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    assert any(
        row["property"] == "different_first_page_header_footer" and row["after"] == "\u662f"
        for row in rows
    )


def test_annotated_document_adds_different_first_page_header_footer_note(tmp_path: Path) -> None:
    input_path = build_section_page_number_docx(tmp_path / "sectioned.docx")

    _, _, annotated_path = normalize_document(
        input_path,
        _rule_set_with_different_first_page_header_footer_rule(),
        tmp_path / "normalized",
    )

    document = Document(annotated_path)
    first_paragraph = document.paragraphs[0]

    assert "[\u89c4\u8303\u5316\u6279\u6ce8]" in first_paragraph.text
    assert "\u9996\u9875\u4e0d\u540c\u9875\u7709\u9875\u811a\u5df2\u6309\u89c4\u8303\u8bbe\u4e3a \u662f" in first_paragraph.text


def test_normalize_document_can_apply_odd_and_even_pages_header_footer_rule(tmp_path: Path) -> None:
    input_path = build_section_page_number_docx(tmp_path / "sectioned.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_odd_and_even_pages_header_footer_rule(),
        tmp_path / "normalized",
    )

    document = Document(output_path)

    assert document.settings.odd_and_even_pages_header_footer is True

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    assert any(
        row["property"] == "odd_and_even_pages_header_footer" and row["after"] == "\u662f"
        for row in rows
    )


def test_normalize_document_can_apply_variant_specific_header_rules(tmp_path: Path) -> None:
    input_path = build_sample_docx(tmp_path / "header-variants.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_variant_specific_header_rules(),
        tmp_path / "normalized",
    )

    document = Document(output_path)
    section = document.sections[0]

    assert section.header.paragraphs[0].runs[0]._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "SimHei"
    assert section.first_page_header.paragraphs[0].runs[0]._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "KaiTi"
    assert section.even_page_header.paragraphs[0].runs[0]._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "FangSong"

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    assert any(row["location"] == "headers[0].items[0]" and row["status"] == "modified" for row in rows)
    assert any(row["location"] == "headers[1].items[0]" and row["status"] == "modified" for row in rows)
    assert any(row["location"] == "headers[2].items[0]" and row["status"] == "modified" for row in rows)


def test_normalize_document_can_apply_variant_specific_footer_rules(tmp_path: Path) -> None:
    input_path = build_header_footer_variant_docx(tmp_path / "footer-variants.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_variant_specific_footer_rules(),
        tmp_path / "normalized",
    )

    document = Document(output_path)
    section = document.sections[0]

    assert section.footer.paragraphs[0].runs[0]._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "SimHei"
    assert section.first_page_footer.paragraphs[0].runs[0]._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "KaiTi"
    assert section.even_page_footer.paragraphs[0].runs[0]._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "FangSong"

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    assert any(row["location"] == "footers[0].items[0]" and row["status"] == "modified" for row in rows)
    assert any(row["location"] == "footers[1].items[0]" and row["status"] == "modified" for row in rows)
    assert any(row["location"] == "footers[2].items[0]" and row["status"] == "modified" for row in rows)


def test_annotated_document_adds_odd_and_even_pages_header_footer_note(tmp_path: Path) -> None:
    input_path = build_section_page_number_docx(tmp_path / "sectioned.docx")

    _, _, annotated_path = normalize_document(
        input_path,
        _rule_set_with_odd_and_even_pages_header_footer_rule(),
        tmp_path / "normalized",
    )

    document = Document(annotated_path)
    first_paragraph = document.paragraphs[0]

    assert "[\u89c4\u8303\u5316\u6279\u6ce8]" in first_paragraph.text
    assert "\u5947\u5076\u9875\u4e0d\u540c\u9875\u7709\u9875\u811a\u5df2\u6309\u89c4\u8303\u8bbe\u4e3a \u662f" in first_paragraph.text


def test_annotated_document_adds_section_page_number_note(tmp_path: Path) -> None:
    input_path = build_section_page_number_docx(tmp_path / "sectioned.docx")

    _, _, annotated_path = normalize_document(
        input_path,
        _rule_set_with_page_number_rules(),
        tmp_path / "normalized",
    )

    document = Document(annotated_path)
    first_paragraph = document.paragraphs[0]

    assert "[\u89c4\u8303\u5316\u6279\u6ce8]" in first_paragraph.text
    assert "\u9875\u7801\u683c\u5f0f\u5df2\u6309\u89c4\u8303\u8bbe\u4e3a lowerRoman" in first_paragraph.text
    assert "\u9875\u7801\u8d77\u59cb\u503c\u5df2\u6309\u89c4\u8303\u8bbe\u4e3a 1" in first_paragraph.text
    assert "\u9875\u811a\u9875\u7801\u5bf9\u9f50\u5df2\u6309\u89c4\u8303\u8bbe\u4e3a center" in first_paragraph.text


def test_normalize_document_resets_heading_and_body_paragraph_formatting(
    tmp_path: Path,
) -> None:
    input_path = build_normalization_sample_docx(tmp_path / "paper.docx")

    output_path, _, _ = normalize_document(input_path, _rule_set(), tmp_path / "normalized")

    document = Document(output_path)
    heading = next(paragraph for paragraph in document.paragraphs if paragraph.text == "第一章 绪论")
    body = next(paragraph for paragraph in document.paragraphs if paragraph.text == "Body paragraph text")

    assert heading.runs[0]._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "SimHei"
    assert body.runs[0].font.name == "Times New Roman"
    assert body.paragraph_format.first_line_indent.cm == pytest.approx(0.74, abs=0.01)
    assert body.paragraph_format.space_after.pt == pytest.approx(0.0, abs=0.1)


def test_normalize_document_writes_report_with_schema_columns(tmp_path: Path) -> None:
    input_path = build_normalization_sample_docx(tmp_path / "paper.docx")

    _, report_path, _ = normalize_document(input_path, _rule_set(), tmp_path / "normalized")

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        rows = list(reader)

    assert reader.fieldnames == [field.column_name for field in _rule_set().report_schema]
    assert rows


def test_normalize_document_reports_unresolved_items(tmp_path: Path) -> None:
    input_path = build_normalization_sample_docx(tmp_path / "paper.docx")

    _, report_path, _ = normalize_document(input_path, _rule_set(), tmp_path / "normalized")

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    unresolved_row = next(row for row in rows if row["text_preview"] == "Unresolved paragraph text")

    assert unresolved_row["status"] == "unresolved"
    assert unresolved_row["reason"] == "no matching classification rule"
    assert unresolved_row["after"] == ""


def test_unresolved_paragraph_stays_unchanged_in_output_docx(tmp_path: Path) -> None:
    input_path = build_normalization_sample_docx(tmp_path / "paper.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_unresolved_body_conflict(),
        tmp_path / "normalized",
    )

    source_document = Document(input_path)
    normalized_document = Document(output_path)
    source_paragraph = next(
        paragraph for paragraph in source_document.paragraphs if paragraph.text == "Unresolved paragraph text"
    )
    normalized_paragraph = next(
        paragraph for paragraph in normalized_document.paragraphs if paragraph.text == "Unresolved paragraph text"
    )

    assert normalized_paragraph.runs[0].font.name == source_paragraph.runs[0].font.name == "Courier New"
    assert normalized_paragraph.paragraph_format.space_after.pt == pytest.approx(
        source_paragraph.paragraph_format.space_after.pt,
        abs=0.1,
    )

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    unresolved_rows = [
        row for row in rows if row["text_preview"] == "Unresolved paragraph text" and row["status"] == "unresolved"
    ]
    assert unresolved_rows


def test_blank_paragraph_before_target_does_not_shift_normalization_target(
    tmp_path: Path,
) -> None:
    input_path = build_normalization_sample_docx(
        tmp_path / "paper.docx",
        include_blank_paragraph_before_body=True,
    )

    output_path, report_path, _ = normalize_document(input_path, _rule_set(), tmp_path / "normalized")

    document = Document(output_path)
    body = next(paragraph for paragraph in document.paragraphs if paragraph.text == "Body paragraph text")
    unresolved = next(paragraph for paragraph in document.paragraphs if paragraph.text == "Unresolved paragraph text")
    blank_paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text == ""]

    assert blank_paragraphs
    assert body.runs[0].font.name == "Times New Roman"
    assert body.paragraph_format.first_line_indent.cm == pytest.approx(0.74, abs=0.01)
    assert unresolved.runs[0].font.name == "Courier New"

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    modified_body_rows = [row for row in rows if row["text_preview"] == "Body paragraph text" and row["status"] == "modified"]
    unresolved_rows = [row for row in rows if row["text_preview"] == "Unresolved paragraph text" and row["status"] == "unresolved"]

    assert modified_body_rows
    assert unresolved_rows


def test_heading_is_not_mutated_by_default_body_rule_once_classified_as_heading(
    tmp_path: Path,
) -> None:
    input_path = build_normalization_sample_docx(tmp_path / "paper.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_default_body_heading_conflict(),
        tmp_path / "normalized",
    )

    document = Document(output_path)
    heading = next(paragraph for paragraph in document.paragraphs if paragraph.text == "第一章 绪论")

    assert heading.runs[0]._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "SimHei"
    assert heading.paragraph_format.space_after is None

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    assert not any(
        row["text_preview"] == "第一章 绪论" and row["property"] == "space_after"
        for row in rows
    )


def test_mixed_run_font_size_normalization_reports_modified(tmp_path: Path) -> None:
    input_path = build_mixed_run_font_size_docx(tmp_path / "mixed.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_font_size_rule(),
        tmp_path / "normalized",
    )

    document = Document(output_path)
    paragraph = next(paragraph for paragraph in document.paragraphs if paragraph.text == "Mixed run sizes")

    assert [run.font.size.pt for run in paragraph.runs] == pytest.approx([12.0, 12.0], abs=0.1)

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    row = next(
        row for row in rows
        if row["text_preview"] == "Mixed run sizes" and row["property"] == "font_size"
    )

    assert row["status"] == "modified"
    assert row["before"].startswith("mixed[")
    assert "10pt" in row["before"]
    assert "14pt" in row["before"]
    assert row["after"] == "12pt"


def test_report_columns_follow_schema_order_not_list_order(tmp_path: Path) -> None:
    input_path = build_normalization_sample_docx(tmp_path / "paper.docx")
    rule_set = _rule_set()
    shuffled_schema = [
        rule_set.report_schema[5],
        rule_set.report_schema[0],
        rule_set.report_schema[10],
        rule_set.report_schema[3],
        rule_set.report_schema[7],
        rule_set.report_schema[1],
        rule_set.report_schema[9],
        rule_set.report_schema[2],
        rule_set.report_schema[8],
        rule_set.report_schema[4],
        rule_set.report_schema[6],
    ]
    shuffled_rule_set = RuleSet(
        document_rules=rule_set.document_rules,
        paragraph_rules=rule_set.paragraph_rules,
        numbering_rules=rule_set.numbering_rules,
        table_rules=rule_set.table_rules,
        special_object_rules=rule_set.special_object_rules,
        report_schema=shuffled_schema,
    )

    _, report_path, _ = normalize_document(input_path, shuffled_rule_set, tmp_path / "normalized")

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        rows = list(reader)

    assert rows
    assert reader.fieldnames == [field.column_name for field in sorted(rule_set.report_schema, key=lambda field: (field.order, field.column_name))]


def test_normalize_document_applies_header_paragraph_rules(tmp_path: Path) -> None:
    input_path = build_header_and_table_normalization_docx(tmp_path / "header-table.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_header_and_table_rules(),
        tmp_path / "normalized",
    )

    document = Document(output_path)
    header = document.sections[0].header.paragraphs[0]

    assert header.runs[0]._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "SimHei"
    assert header.runs[0].font.size.pt == pytest.approx(10.5, abs=0.1)

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    font_row = next(
        row for row in rows
        if row["location"] == "headers[0].items[0]" and row["property"] == "font_name"
    )
    assert font_row["status"] == "modified"
    assert font_row["after"] == "SimHei"


def test_normalize_document_applies_table_rules_to_all_cells(tmp_path: Path) -> None:
    input_path = build_header_and_table_normalization_docx(tmp_path / "header-table.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_header_and_table_rules(),
        tmp_path / "normalized",
    )

    document = Document(output_path)
    table = document.tables[0]

    for row in table.rows:
        for cell in row.cells:
            run = cell.paragraphs[0].runs[0]
            assert run._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "SimSun"
            assert run.font.size.pt == pytest.approx(10.5, abs=0.1)

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    font_row = next(
        row for row in rows
        if row["location"] == "body_items[1]" and row["property"] == "font_name"
    )
    assert font_row["status"] == "modified"
    assert font_row["after"] == "SimSun"


def test_normalize_document_can_apply_distinct_table_header_and_body_rules(tmp_path: Path) -> None:
    input_path = build_header_and_table_normalization_docx(tmp_path / "header-table.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_split_table_rules(),
        tmp_path / "normalized",
    )

    document = Document(output_path)
    table = document.tables[0]

    header_cell_run = table.cell(0, 0).paragraphs[0].runs[0]
    body_cell_run = table.cell(1, 0).paragraphs[0].runs[0]

    assert header_cell_run._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "SimHei"
    assert header_cell_run.font.size.pt == pytest.approx(12.0, abs=0.1)
    assert body_cell_run._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "SimSun"
    assert body_cell_run.font.size.pt == pytest.approx(10.5, abs=0.1)

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    header_font_row = next(
        row for row in rows
        if row["location"] == "body_items[1]" and row["property"] == "header_row_font_name"
    )
    body_font_row = next(
        row for row in rows
        if row["location"] == "body_items[1]" and row["property"] == "body_rows_font_name"
    )

    assert header_font_row["status"] == "modified"
    assert header_font_row["after"] == "SimHei"
    assert body_font_row["status"] == "modified"
    assert body_font_row["after"] == "SimSun"


def test_annotated_document_marks_only_targeted_table_rows_for_split_rules(tmp_path: Path) -> None:
    input_path = build_header_and_table_normalization_docx(tmp_path / "header-table.docx")

    _, _, annotated_path = normalize_document(
        input_path,
        _rule_set_with_split_table_rules(),
        tmp_path / "normalized",
    )

    document = Document(annotated_path)
    table = document.tables[0]

    header_cell_paragraph = table.cell(0, 0).paragraphs[0]
    body_cell_paragraph = table.cell(1, 0).paragraphs[0]

    assert header_cell_paragraph.runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
    assert body_cell_paragraph.runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
    assert "【规范化说明：" in header_cell_paragraph.runs[-1].text
    assert "【规范化说明：" in body_cell_paragraph.runs[-1].text


def test_normalize_document_can_apply_column_and_cell_table_rules(tmp_path: Path) -> None:
    input_path = build_header_and_table_normalization_docx(tmp_path / "header-table.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_column_and_cell_table_rules(),
        tmp_path / "normalized",
    )

    document = Document(output_path)
    table = document.tables[0]

    first_column_header = table.cell(0, 0).paragraphs[0].runs[0]
    first_column_body = table.cell(1, 0).paragraphs[0].runs[0]
    untouched_cell = table.cell(0, 1).paragraphs[0].runs[0]
    targeted_cell = table.cell(1, 1).paragraphs[0].runs[0]

    assert first_column_header._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "SimHei"
    assert first_column_body._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "SimHei"
    assert untouched_cell._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) in {None, "Calibri"}
    assert targeted_cell.font.size.pt == pytest.approx(14.0, abs=0.1)
    assert untouched_cell.font.size.pt == pytest.approx(9.0, abs=0.1)

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    column_row = next(
        row for row in rows
        if row["text_preview"] == "项目一\t项目二\n项目三\t项目四"
        and row["property"] == "column[0]_font_name"
    )
    cell_row = next(
        row for row in rows
        if row["text_preview"] == "项目一\t项目二\n项目三\t项目四"
        and row["property"] == "cell[1,1]_font_size"
    )

    assert column_row["after"] == "SimHei"
    assert cell_row["after"] == "14pt"


def test_annotated_document_marks_only_targeted_table_cells_for_column_and_cell_rules(
    tmp_path: Path,
) -> None:
    input_path = build_header_and_table_normalization_docx(tmp_path / "header-table.docx")

    _, _, annotated_path = normalize_document(
        input_path,
        _rule_set_with_column_and_cell_table_rules(),
        tmp_path / "normalized",
    )

    document = Document(annotated_path)
    table = document.tables[0]

    targeted_column_paragraph = table.cell(0, 0).paragraphs[0]
    untouched_paragraph = table.cell(0, 1).paragraphs[0]
    targeted_cell_paragraph = table.cell(1, 1).paragraphs[0]

    assert targeted_column_paragraph.runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
    assert untouched_paragraph.runs[0].font.color.rgb is None
    assert targeted_cell_paragraph.runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
    assert "第1列" in "".join(run.text for run in targeted_column_paragraph.runs[1:])
    assert "第2行第2列" in "".join(run.text for run in targeted_cell_paragraph.runs[1:])


def test_normalize_document_can_apply_table_rules_by_header_name(tmp_path: Path) -> None:
    input_path = build_header_and_table_normalization_docx(tmp_path / "header-table.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_header_named_table_rules(),
        tmp_path / "normalized",
    )

    document = Document(output_path)
    table = document.tables[0]

    target_header = table.cell(0, 1).paragraphs[0].runs[0]
    target_body = table.cell(1, 1).paragraphs[0].runs[0]
    untouched_body = table.cell(1, 0).paragraphs[0].runs[0]

    assert target_header._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "SimHei"
    assert target_body._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "SimHei"
    assert untouched_body._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) in {None, "Calibri"}

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    header_named_row = next(
        row for row in rows
        if row["text_preview"] == "项目一\t项目二\n项目三\t项目四"
        and row["property"] == "column_by_header[项目二]_font_name"
    )

    assert header_named_row["after"] == "SimHei"


def test_annotated_document_marks_only_header_named_table_column_in_red(tmp_path: Path) -> None:
    input_path = build_header_and_table_normalization_docx(tmp_path / "header-table.docx")

    _, _, annotated_path = normalize_document(
        input_path,
        _rule_set_with_header_named_table_rules(),
        tmp_path / "normalized",
    )

    document = Document(annotated_path)
    table = document.tables[0]

    targeted_header_paragraph = table.cell(0, 1).paragraphs[0]
    targeted_body_paragraph = table.cell(1, 1).paragraphs[0]
    untouched_paragraph = table.cell(1, 0).paragraphs[0]

    assert targeted_header_paragraph.runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
    assert targeted_body_paragraph.runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
    assert untouched_paragraph.runs[0].font.color.rgb is None
    assert '表头“项目二”列' in "".join(run.text for run in targeted_body_paragraph.runs[1:])


def test_normalize_document_can_apply_table_rules_by_column_range(tmp_path: Path) -> None:
    input_path = build_header_and_table_normalization_docx(tmp_path / "header-table.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_column_range_table_rules(),
        tmp_path / "normalized",
    )

    document = Document(output_path)
    table = document.tables[0]

    first_column_run = table.cell(0, 0).paragraphs[0].runs[0]
    second_column_run = table.cell(0, 1).paragraphs[0].runs[0]

    assert first_column_run.font.size.pt == pytest.approx(14.0, abs=0.1)
    assert second_column_run.font.size.pt == pytest.approx(14.0, abs=0.1)

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    range_row = next(
        row for row in rows
        if row["text_preview"] == "项目一\t项目二\n项目三\t项目四"
        and row["property"] == "column_range[0:2]_font_size"
    )

    assert range_row["after"] == "14pt"


def test_annotated_document_marks_only_targeted_table_column_range_in_red(tmp_path: Path) -> None:
    input_path = build_header_and_table_normalization_docx(tmp_path / "header-table.docx")

    _, _, annotated_path = normalize_document(
        input_path,
        _rule_set_with_column_range_table_rules(),
        tmp_path / "normalized",
    )

    document = Document(annotated_path)
    table = document.tables[0]

    first_column_paragraph = table.cell(0, 0).paragraphs[0]
    second_column_paragraph = table.cell(0, 1).paragraphs[0]

    assert first_column_paragraph.runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
    assert second_column_paragraph.runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
    assert "第1列到第2列" in "".join(run.text for run in second_column_paragraph.runs[1:])


def test_normalize_document_can_apply_table_rules_by_row_index(tmp_path: Path) -> None:
    input_path = build_header_and_table_normalization_docx(tmp_path / "header-table.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_row_table_rules(),
        tmp_path / "normalized",
    )

    document = Document(output_path)
    table = document.tables[0]

    target_row_first_cell = table.cell(1, 0).paragraphs[0].runs[0]
    target_row_second_cell = table.cell(1, 1).paragraphs[0].runs[0]
    untouched_row_first_cell = table.cell(0, 0).paragraphs[0].runs[0]

    assert target_row_first_cell._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "SimHei"
    assert target_row_second_cell._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "SimHei"
    assert untouched_row_first_cell._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) in {None, "Calibri"}

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    row_rule = next(
        row for row in rows
        if row["text_preview"] == "项目一\t项目二\n项目三\t项目四"
        and row["property"] == "row[1]_font_name"
    )

    assert row_rule["after"] == "SimHei"


def test_normalize_document_can_apply_table_rules_by_row_range(tmp_path: Path) -> None:
    input_path = build_header_and_table_normalization_docx(tmp_path / "header-table.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_row_range_table_rules(),
        tmp_path / "normalized",
    )

    document = Document(output_path)
    table = document.tables[0]

    header_row_cell = table.cell(0, 0).paragraphs[0].runs[0]
    body_row_cell = table.cell(1, 0).paragraphs[0].runs[0]

    assert header_row_cell.font.size.pt == pytest.approx(14.0, abs=0.1)
    assert body_row_cell.font.size.pt == pytest.approx(14.0, abs=0.1)

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    row_range_rule = next(
        row for row in rows
        if row["text_preview"] == "项目一\t项目二\n项目三\t项目四"
        and row["property"] == "row_range[0:2]_font_size"
    )

    assert row_range_rule["after"] == "14pt"


def test_annotated_document_marks_only_targeted_rows_in_red(tmp_path: Path) -> None:
    input_path = build_header_and_table_normalization_docx(tmp_path / "header-table.docx")

    _, _, annotated_path = normalize_document(
        input_path,
        _rule_set_with_row_table_rules(),
        tmp_path / "normalized",
    )

    document = Document(annotated_path)
    table = document.tables[0]

    targeted_row_paragraph = table.cell(1, 0).paragraphs[0]
    untouched_row_paragraph = table.cell(0, 0).paragraphs[0]

    assert targeted_row_paragraph.runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
    assert untouched_row_paragraph.runs[0].font.color.rgb is None
    assert "第2行" in "".join(run.text for run in targeted_row_paragraph.runs[1:])


def test_annotated_document_marks_only_targeted_row_range_in_red(tmp_path: Path) -> None:
    input_path = build_header_and_table_normalization_docx(tmp_path / "header-table.docx")

    _, _, annotated_path = normalize_document(
        input_path,
        _rule_set_with_row_range_table_rules(),
        tmp_path / "normalized",
    )

    document = Document(annotated_path)
    table = document.tables[0]

    header_row_paragraph = table.cell(0, 0).paragraphs[0]
    body_row_paragraph = table.cell(1, 0).paragraphs[0]

    assert header_row_paragraph.runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
    assert body_row_paragraph.runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
    assert "第1行到第2行" in "".join(run.text for run in body_row_paragraph.runs[1:])


def test_normalize_document_can_apply_table_rules_by_cell_range(tmp_path: Path) -> None:
    input_path = build_header_and_table_normalization_docx(tmp_path / "header-table.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_cell_range_table_rules(),
        tmp_path / "normalized",
    )

    document = Document(output_path)
    table = document.tables[0]

    top_left = table.cell(0, 0).paragraphs[0].runs[0]
    top_right = table.cell(0, 1).paragraphs[0].runs[0]
    bottom_left = table.cell(1, 0).paragraphs[0].runs[0]
    bottom_right = table.cell(1, 1).paragraphs[0].runs[0]

    assert top_left._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "SimHei"
    assert top_right._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "SimHei"
    assert bottom_left._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "SimHei"
    assert bottom_right._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "SimHei"

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    range_rule = next(
        row for row in rows
        if row["text_preview"] == "项目一\t项目二\n项目三\t项目四"
        and row["property"] == "cell_range[0:2,0:2]_font_name"
    )

    assert range_rule["after"] == "SimHei"


def test_annotated_document_marks_only_targeted_cell_range_in_red(tmp_path: Path) -> None:
    input_path = build_header_and_table_normalization_docx(tmp_path / "header-table.docx")

    _, _, annotated_path = normalize_document(
        input_path,
        _rule_set_with_cell_range_table_rules(),
        tmp_path / "normalized",
    )

    document = Document(annotated_path)
    table = document.tables[0]

    top_left = table.cell(0, 0).paragraphs[0]
    bottom_right = table.cell(1, 1).paragraphs[0]

    assert top_left.runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
    assert bottom_right.runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
    assert "第1行到第2行、第1列到第2列" in "".join(run.text for run in bottom_right.runs[1:])


def test_normalize_document_can_apply_table_bold_and_alignment_rules(tmp_path: Path) -> None:
    input_path = build_header_and_table_normalization_docx(tmp_path / "header-table.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_table_bold_and_alignment_rules(),
        tmp_path / "normalized",
    )

    document = Document(output_path)
    table = document.tables[0]

    first_column_header = table.cell(0, 0).paragraphs[0]
    first_column_body = table.cell(1, 0).paragraphs[0]
    second_column_header = table.cell(0, 1).paragraphs[0]

    assert first_column_header.runs[0].bold is True
    assert first_column_body.runs[0].bold is True
    assert second_column_header.runs[0].bold in {None, False}
    assert second_column_header.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert first_column_header.alignment != WD_ALIGN_PARAGRAPH.CENTER

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    bold_row = next(
        row for row in rows
        if row["text_preview"] == "项目一\t项目二\n项目三\t项目四"
        and row["property"] == "column[0]_bold"
    )
    alignment_row = next(
        row for row in rows
        if row["text_preview"] == "项目一\t项目二\n项目三\t项目四"
        and row["property"] == "cell[0,1]_alignment"
    )

    assert bold_row["after"] == "是"
    assert alignment_row["after"] == "center"


def test_annotated_document_marks_table_bold_and_alignment_changes_in_red(tmp_path: Path) -> None:
    input_path = build_header_and_table_normalization_docx(tmp_path / "header-table.docx")

    _, _, annotated_path = normalize_document(
        input_path,
        _rule_set_with_table_bold_and_alignment_rules(),
        tmp_path / "normalized",
    )

    document = Document(annotated_path)
    table = document.tables[0]

    bold_target = table.cell(0, 0).paragraphs[0]
    alignment_target = table.cell(0, 1).paragraphs[0]

    assert bold_target.runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
    assert alignment_target.runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
    assert "第1列加粗原为 否，调整为 是" in "".join(run.text for run in bold_target.runs[1:])
    assert "第1行第2列对齐原为 left，调整为 center" in "".join(
        run.text for run in alignment_target.runs[1:]
    )


def test_normalize_document_can_apply_table_vertical_alignment_and_border_rules(tmp_path: Path) -> None:
    input_path = build_header_and_table_normalization_docx(tmp_path / "header-table.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_table_vertical_alignment_and_border_rules(),
        tmp_path / "normalized",
    )

    document = Document(output_path)
    table = document.tables[0]

    header_left = table.cell(0, 0)
    header_right = table.cell(0, 1)
    body_left = table.cell(1, 0)

    assert header_left.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER
    assert header_right.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER
    assert body_left.vertical_alignment != WD_CELL_VERTICAL_ALIGNMENT.CENTER
    assert _cell_border_values(header_left) == {"single"}
    assert _cell_border_values(body_left) != {"single"}

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    vertical_row = next(row for row in rows if row["property"] == "header_row_vertical_alignment")
    border_row = next(row for row in rows if row["property"] == "cell[0,0]_border")

    assert vertical_row["after"] == "center"
    assert border_row["after"] == "single"


def test_annotated_document_adds_table_vertical_alignment_and_border_notes(tmp_path: Path) -> None:
    input_path = build_header_and_table_normalization_docx(tmp_path / "header-table.docx")

    _, _, annotated_path = normalize_document(
        input_path,
        _rule_set_with_table_vertical_alignment_and_border_rules(),
        tmp_path / "normalized",
    )

    document = Document(annotated_path)
    note_texts = [paragraph.text for paragraph in document.paragraphs if "规范" in paragraph.text]

    assert any("表头垂直对齐原为" in text and "规范为 center" in text for text in note_texts)
    assert any("第1行第1列边框原为" in text and "规范为 single" in text for text in note_texts)


def test_normalize_document_aligns_header_indexes_when_earlier_headers_are_empty(
    tmp_path: Path,
) -> None:
    input_path = build_sparse_header_normalization_docx(tmp_path / "sparse-header.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_late_header_rule(),
        tmp_path / "normalized",
    )

    document = Document(output_path)
    header = document.sections[1].header.paragraphs[0]

    assert header.runs[0]._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "SimHei"
    assert header.runs[0].font.size.pt == pytest.approx(10.5, abs=0.1)

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    font_row = next(
        row for row in rows
        if row["text_preview"] == "后置页眉" and row["property"] == "font_name"
    )
    assert font_row["status"] == "modified"


def test_annotated_document_marks_textual_changes_in_red(tmp_path: Path) -> None:
    input_path = build_normalization_sample_docx(tmp_path / "paper.docx")

    _, _, annotated_path = normalize_document(input_path, _rule_set(), tmp_path / "normalized")

    document = Document(annotated_path)
    body = next(paragraph for paragraph in document.paragraphs if paragraph.text.startswith("Body paragraph text"))

    assert body.runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
    assert body.runs[-1].text == "【规范化说明：字体原为 Arial（无衬线体），调整为 Times New Roman（新罗马体）】"
    assert body.runs[-1].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)


def test_annotated_document_inserts_red_note_for_layout_only_changes(tmp_path: Path) -> None:
    input_path = build_normalization_sample_docx(tmp_path / "paper.docx")

    _, _, annotated_path = normalize_document(input_path, _rule_set(), tmp_path / "normalized")

    document = Document(annotated_path)
    texts = [paragraph.text for paragraph in document.paragraphs]
    body_index = next(index for index, text in enumerate(texts) if text.startswith("Body paragraph text"))
    note_index = texts.index("[规范化批注] 首行缩进原为 0cm，规范为 0.74cm；段后原为 18pt，规范为 0pt")
    note_paragraph = document.paragraphs[note_index]

    assert note_index > body_index
    assert note_paragraph.runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)


def test_annotated_note_uses_consistent_review_style(tmp_path: Path) -> None:
    input_path = build_normalization_sample_docx(tmp_path / "paper.docx")

    _, _, annotated_path = normalize_document(input_path, _rule_set(), tmp_path / "normalized")

    document = Document(annotated_path)
    note = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == "[规范化批注] 首行缩进原为 0cm，规范为 0.74cm；段后原为 18pt，规范为 0pt"
    )

    assert note.runs[0].font.name == "SimSun"
    assert note.runs[0].font.size.pt == pytest.approx(10.5, abs=0.1)
    assert note.paragraph_format.first_line_indent.cm == pytest.approx(0.0, abs=0.01)


def test_annotated_document_explains_font_and_size_changes_inline(tmp_path: Path) -> None:
    input_path = build_header_and_table_normalization_docx(tmp_path / "header-table.docx")

    _, _, annotated_path = normalize_document(
        input_path,
        _rule_set_with_header_and_table_rules(),
        tmp_path / "normalized",
    )

    document = Document(annotated_path)
    header = document.sections[0].header.paragraphs[0]
    cell_paragraph = document.tables[0].cell(0, 0).paragraphs[0]

    assert header.runs[-1].text == "【规范化说明：字体原为 Calibri（默认无衬线体），调整为 黑体（SimHei）；字号原为 小五（9pt），调整为 五号（10.5pt）】"
    annotation_text = "".join(run.text for run in cell_paragraph.runs[1:])
    assert "【规范化说明：字体原为 Calibri（默认无衬线体），调整为 宋体（SimSun）】" in annotation_text
    assert "【规范化说明：字号原为 小五（9pt），调整为 五号（10.5pt）】" in annotation_text


def test_normalize_document_assigns_chinese_and_western_font_slots_for_mixed_script_text(
    tmp_path: Path,
) -> None:
    input_path = build_mixed_script_font_name_docx(tmp_path / "mixed-script.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_mixed_script_body_font_rule(),
        tmp_path / "normalized",
    )

    document = Document(output_path)
    paragraph = next(paragraph for paragraph in document.paragraphs if paragraph.text == "中文ABC123，测试")
    run = paragraph.runs[0]
    r_fonts = run._element.get_or_add_rPr().rFonts

    assert r_fonts.get(qn("w:eastAsia")) == "SimSun"
    assert r_fonts.get(qn("w:ascii")) == "Times New Roman"
    assert r_fonts.get(qn("w:hAnsi")) == "Times New Roman"
    assert r_fonts.get(qn("w:cs")) == "Times New Roman"

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    font_row = next(
        row for row in rows
        if row["text_preview"] == "中文ABC123，测试" and row["property"] == "font_name"
    )
    assert font_row["after"] == "SimSun|Times New Roman"


def test_normalize_document_can_separate_english_abstract_and_reference_entry_fonts(
    tmp_path: Path,
) -> None:
    input_path = build_abstract_and_reference_docx(tmp_path / "abstract-reference.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_abstract_and_reference_rules(),
        tmp_path / "normalized",
    )

    document = Document(output_path)
    abstract_paragraph = next(
        paragraph for paragraph in document.paragraphs if paragraph.text.startswith("【Abstract】")
    )
    abstract_fonts = abstract_paragraph.runs[0]._element.get_or_add_rPr().rFonts
    assert abstract_fonts.get(qn("w:eastAsia")) == "Times New Roman"
    assert abstract_fonts.get(qn("w:ascii")) == "Times New Roman"

    reference_paragraph = next(
        paragraph for paragraph in document.paragraphs if "[M]." in paragraph.text
    )
    reference_fonts = reference_paragraph.runs[0]._element.get_or_add_rPr().rFonts
    assert reference_fonts.get(qn("w:eastAsia")) == "SimSun"
    assert reference_fonts.get(qn("w:ascii")) == "Times New Roman"

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    abstract_row = next(
        row for row in rows
        if row["text_preview"].startswith("【Abstract】") and row["property"] == "font_name"
    )
    reference_row = next(
        row for row in rows
        if "[M]." in row["text_preview"] and row["property"] == "font_name"
    )

    assert abstract_row["after"] == "Times New Roman"
    assert reference_row["after"] == "SimSun|Times New Roman"


def test_normalize_document_can_apply_reference_hanging_indent(
    tmp_path: Path,
) -> None:
    input_path = build_abstract_and_reference_docx(tmp_path / "abstract-reference.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_abstract_and_reference_rules(),
        tmp_path / "normalized",
    )

    document = Document(output_path)
    reference_paragraph = next(
        paragraph for paragraph in document.paragraphs if "[M]." in paragraph.text
    )

    assert reference_paragraph.paragraph_format.left_indent.cm == pytest.approx(0.74, abs=0.01)
    assert reference_paragraph.paragraph_format.first_line_indent.cm == pytest.approx(-0.74, abs=0.01)

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    hanging_row = next(
        row for row in rows
        if "[M]." in row["text_preview"] and row["property"] == "hanging_indent"
    )
    assert hanging_row["after"] == "0.740833cm"


def test_annotated_document_describes_inherited_font_as_not_explicitly_set(
    tmp_path: Path,
) -> None:
    input_path = build_abstract_and_reference_docx(tmp_path / "abstract-reference.docx")

    _, _, annotated_path = normalize_document(
        input_path,
        _rule_set_with_abstract_and_reference_rules(),
        tmp_path / "normalized",
    )

    document = Document(annotated_path)
    abstract_paragraph = next(
        paragraph for paragraph in document.paragraphs if paragraph.text.startswith("【Abstract】")
    )

    assert abstract_paragraph.runs[-1].text == (
        "【规范化说明：字体原为 未显式设置（继承样式），调整为 Times New Roman（新罗马体）】"
    )


def test_annotated_document_describes_inherited_layout_values_as_not_explicitly_set(
    tmp_path: Path,
) -> None:
    input_path = build_normalization_sample_docx(tmp_path / "paper.docx")

    _, _, annotated_path = normalize_document(
        input_path,
        _rule_set_with_inherited_layout_rules(),
        tmp_path / "normalized",
    )

    document = Document(annotated_path)
    annotation = next(
        paragraph.text
        for paragraph in document.paragraphs
        if "0.74cm" in paragraph.text
        and "20pt" in paragraph.text
        and "18pt" in paragraph.text
    )

    inherited = "\u672a\u663e\u5f0f\u8bbe\u7f6e\uff08\u7ee7\u627f\u6837\u5f0f\uff09"
    assert inherited in annotation
    assert f"\u6bb5\u524d\u539f\u4e3a {inherited}" in annotation
    assert f"\u884c\u8ddd\u539f\u4e3a {inherited}" in annotation
    assert "254000" not in annotation


def test_normalize_document_can_apply_inline_label_and_content_rules(
    tmp_path: Path,
) -> None:
    input_path = build_inline_tag_docx(tmp_path / "inline-tag.docx")

    output_path, report_path, _ = normalize_document(
        input_path,
        _rule_set_with_inline_tag_rules(),
        tmp_path / "normalized",
    )

    document = Document(output_path)
    paragraph = next(paragraph for paragraph in document.paragraphs if paragraph.text == "【摘要】中文ABC")

    assert len(paragraph.runs) == 2
    assert paragraph.runs[0].text == "【摘要】"
    assert paragraph.runs[0]._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "SimHei"
    assert paragraph.runs[0].font.bold is True
    assert paragraph.runs[0].font.size.pt == pytest.approx(15.0, abs=0.1)
    assert paragraph.runs[1].text == "中文ABC"
    assert paragraph.runs[1]._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia")) == "SimSun"
    assert paragraph.runs[1]._element.get_or_add_rPr().rFonts.get(qn("w:ascii")) == "Times New Roman"
    assert paragraph.runs[1].font.size.pt == pytest.approx(12.0, abs=0.1)

    with report_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    label_font_row = next(
        row for row in rows
        if row["text_preview"] == "【摘要】中文ABC" and row["property"] == "label_font_name"
    )
    content_font_row = next(
        row for row in rows
        if row["text_preview"] == "【摘要】中文ABC" and row["property"] == "content_font_name"
    )
    label_size_row = next(
        row for row in rows
        if row["text_preview"] == "【摘要】中文ABC" and row["property"] == "label_font_size"
    )
    content_size_row = next(
        row for row in rows
        if row["text_preview"] == "【摘要】中文ABC" and row["property"] == "content_font_size"
    )

    assert label_font_row["after"] == "SimHei"
    assert content_font_row["after"] == "SimSun|Times New Roman"
    assert label_size_row["after"] == "15pt"
    assert content_size_row["after"] == "12pt"


def test_annotated_document_marks_only_changed_inline_segment_in_red(
    tmp_path: Path,
) -> None:
    input_path = build_inline_tag_docx(tmp_path / "inline-tag.docx")

    _, _, annotated_path = normalize_document(
        input_path,
        _rule_set_with_inline_label_only_rules(),
        tmp_path / "normalized",
    )

    document = Document(annotated_path)
    paragraph = next(paragraph for paragraph in document.paragraphs if paragraph.text.startswith("【摘要】中文ABC"))

    assert paragraph.runs[0].text == "【摘要】"
    assert paragraph.runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
    assert paragraph.runs[1].text == "中文ABC"
    assert paragraph.runs[1].font.color.rgb is None
    assert paragraph.runs[-1].text == "【规范化说明：标签字体原为 Calibri（默认无衬线体），调整为 黑体（SimHei）；标签字号原为 11pt，调整为 小三（15pt）；标签加粗原为 否，调整为 是】"


def _rule_set() -> RuleSet:
    return RuleSet(
        document_rules=[
            DocumentRule(
                rule_id="DOC-TOP",
                priority=10,
                property_name="page_margin_top",
                value="2.54cm",
                scope="document",
            ),
            DocumentRule(
                rule_id="DOC-BOTTOM",
                priority=11,
                property_name="page_margin_bottom",
                value="2.54cm",
                scope="document",
            ),
        ],
        paragraph_rules=[
            ParagraphRule(
                rule_id="PAR-HEADING-FONT",
                priority=5,
                match_type="style",
                match_value="Heading 1",
                target_property="font_name",
                target_value="SimHei",
            ),
            ParagraphRule(
                rule_id="PAR-BODY-FONT",
                priority=10,
                match_type="style",
                match_value="Body Text",
                target_property="font_name",
                target_value="Times New Roman",
            ),
            ParagraphRule(
                rule_id="PAR-BODY-INDENT",
                priority=11,
                match_type="style",
                match_value="Body Text",
                target_property="first_line_indent",
                target_value="0.74cm",
            ),
            ParagraphRule(
                rule_id="PAR-BODY-SPACE-AFTER",
                priority=12,
                match_type="style",
                match_value="Body Text",
                target_property="space_after",
                target_value="0pt",
            ),
        ],
        numbering_rules=[],
        table_rules=[],
        special_object_rules=[],
        report_schema=[
            ReportSchemaField(column_name="object_id", order=1, description="Stable object identifier"),
            ReportSchemaField(column_name="object_type_before", order=2, description="Object type before normalization"),
            ReportSchemaField(column_name="object_type_after", order=3, description="Object type after normalization"),
            ReportSchemaField(column_name="location", order=4, description="Document location for review"),
            ReportSchemaField(column_name="text_preview", order=5, description="Preview of matched text"),
            ReportSchemaField(column_name="property", order=6, description="Property name"),
            ReportSchemaField(column_name="before", order=7, description="Original value"),
            ReportSchemaField(column_name="after", order=8, description="Normalized value"),
            ReportSchemaField(column_name="rule_id", order=9, description="Applied rule identifier"),
            ReportSchemaField(column_name="status", order=10, description="Change status"),
            ReportSchemaField(column_name="reason", order=11, description="Status explanation"),
        ],
    )


def _rule_set_with_unresolved_body_conflict() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=base.document_rules,
        paragraph_rules=[
            *base.paragraph_rules,
            ParagraphRule(
                rule_id="PAR-UNRESOLVED-FONT",
                priority=30,
                match_type="text",
                match_value="Unresolved paragraph text",
                target_property="font_name",
                target_value="SimSun",
            ),
            ParagraphRule(
                rule_id="PAR-UNRESOLVED-SPACE",
                priority=30,
                match_type="text",
                match_value="Unresolved paragraph text",
                target_property="space_after",
                target_value="24pt",
            ),
        ],
        numbering_rules=base.numbering_rules,
        table_rules=base.table_rules,
        special_object_rules=base.special_object_rules,
        report_schema=base.report_schema,
    )


def _rule_set_with_default_body_heading_conflict() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=base.document_rules,
        paragraph_rules=[
            *base.paragraph_rules,
            ParagraphRule(
                rule_id="PAR-DEFAULT-BODY-SPACE",
                priority=100,
                match_type="default",
                match_value="body",
                target_property="space_after",
                target_value="24pt",
            ),
        ],
        numbering_rules=base.numbering_rules,
        table_rules=base.table_rules,
        special_object_rules=base.special_object_rules,
        report_schema=base.report_schema,
    )


def _rule_set_with_font_size_rule() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=[],
        paragraph_rules=[
            ParagraphRule(
                rule_id="PAR-BODY-FONT-SIZE",
                priority=10,
                match_type="style",
                match_value="Body Text",
                target_property="font_size",
                target_value="12pt",
            ),
        ],
        numbering_rules=[],
        table_rules=[],
        special_object_rules=[],
        report_schema=base.report_schema,
    )


def _rule_set_with_page_number_rules() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=[
            DocumentRule(
                rule_id="DOC-PAGE-FORMAT",
                priority=10,
                property_name="page_number_format",
                value="lowerRoman",
                scope="document",
            ),
            DocumentRule(
                rule_id="DOC-PAGE-START",
                priority=11,
                property_name="page_number_start",
                value="1",
                scope="document",
            ),
            DocumentRule(
                rule_id="DOC-PAGE-ALIGN",
                priority=12,
                property_name="footer_page_number_alignment",
                value="center",
                scope="document",
            ),
        ],
        paragraph_rules=[],
        numbering_rules=[],
        table_rules=[],
        special_object_rules=[],
        report_schema=base.report_schema,
    )


def _rule_set_with_section_start_type_rule() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=[
            DocumentRule(
                rule_id="DOC-SECTION-START",
                priority=10,
                property_name="section_start_type",
                value="odd_page",
                scope="document",
            ),
        ],
        paragraph_rules=[],
        numbering_rules=[],
        table_rules=[],
        special_object_rules=[],
        report_schema=base.report_schema,
    )


def _rule_set_with_different_first_page_header_footer_rule() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=[
            DocumentRule(
                rule_id="DOC-DIFFERENT-FIRST-PAGE",
                priority=10,
                property_name="different_first_page_header_footer",
                value="true",
                scope="document",
            ),
        ],
        paragraph_rules=[],
        numbering_rules=[],
        table_rules=[],
        special_object_rules=[],
        report_schema=base.report_schema,
    )


def _rule_set_with_odd_and_even_pages_header_footer_rule() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=[
            DocumentRule(
                rule_id="DOC-ODD-EVEN-HEADERS",
                priority=10,
                property_name="odd_and_even_pages_header_footer",
                value="true",
                scope="document",
            ),
        ],
        paragraph_rules=[],
        numbering_rules=[],
        table_rules=[],
        special_object_rules=[],
        report_schema=base.report_schema,
    )


def _rule_set_with_variant_specific_header_rules() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=[],
        paragraph_rules=[
            ParagraphRule(
                rule_id="HDR-DEFAULT-FONT",
                priority=10,
                match_type="class",
                match_value="default_running_header",
                target_property="font_name",
                target_value="SimHei",
            ),
            ParagraphRule(
                rule_id="HDR-FIRST-FONT",
                priority=10,
                match_type="class",
                match_value="first_page_running_header",
                target_property="font_name",
                target_value="KaiTi",
            ),
            ParagraphRule(
                rule_id="HDR-EVEN-FONT",
                priority=10,
                match_type="class",
                match_value="even_page_running_header",
                target_property="font_name",
                target_value="FangSong",
            ),
        ],
        numbering_rules=[],
        table_rules=[],
        special_object_rules=[
            SpecialObjectRule(
                rule_id="HDR-DEFAULT-CLASS",
                priority=10,
                object_type="default_header",
                match_type="text",
                match_value="Shared header",
                target_object_type="default_running_header",
            ),
            SpecialObjectRule(
                rule_id="HDR-FIRST-CLASS",
                priority=10,
                object_type="first_page_header",
                match_type="text",
                match_value="Section 1 first-page header",
                target_object_type="first_page_running_header",
            ),
            SpecialObjectRule(
                rule_id="HDR-EVEN-CLASS",
                priority=10,
                object_type="even_page_header",
                match_type="text",
                match_value="Section 1 even-page header",
                target_object_type="even_page_running_header",
            ),
        ],
        report_schema=base.report_schema,
    )


def _rule_set_with_variant_specific_footer_rules() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=[],
        paragraph_rules=[
            ParagraphRule(
                rule_id="FTR-DEFAULT-FONT",
                priority=10,
                match_type="class",
                match_value="default_running_footer",
                target_property="font_name",
                target_value="SimHei",
            ),
            ParagraphRule(
                rule_id="FTR-FIRST-FONT",
                priority=10,
                match_type="class",
                match_value="first_page_running_footer",
                target_property="font_name",
                target_value="KaiTi",
            ),
            ParagraphRule(
                rule_id="FTR-EVEN-FONT",
                priority=10,
                match_type="class",
                match_value="even_page_running_footer",
                target_property="font_name",
                target_value="FangSong",
            ),
        ],
        numbering_rules=[],
        table_rules=[],
        special_object_rules=[
            SpecialObjectRule(
                rule_id="FTR-DEFAULT-CLASS",
                priority=10,
                object_type="default_footer",
                match_type="text",
                match_value="Shared footer",
                target_object_type="default_running_footer",
            ),
            SpecialObjectRule(
                rule_id="FTR-FIRST-CLASS",
                priority=10,
                object_type="first_page_footer",
                match_type="text",
                match_value="Section 1 first-page footer",
                target_object_type="first_page_running_footer",
            ),
            SpecialObjectRule(
                rule_id="FTR-EVEN-CLASS",
                priority=10,
                object_type="even_page_footer",
                match_type="text",
                match_value="Section 1 even-page footer",
                target_object_type="even_page_running_footer",
            ),
        ],
        report_schema=base.report_schema,
    )


def _rule_set_with_header_and_table_rules() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=[],
        paragraph_rules=[
            ParagraphRule(
                rule_id="HDR-FONT",
                priority=10,
                match_type="class",
                match_value="running_header",
                target_property="font_name",
                target_value="SimHei",
            ),
            ParagraphRule(
                rule_id="HDR-SIZE",
                priority=11,
                match_type="class",
                match_value="running_header",
                target_property="font_size",
                target_value="10.5pt",
            ),
        ],
        numbering_rules=[],
        table_rules=[
            TableRule(
                rule_id="TBL-FONT",
                priority=10,
                match_type="regex",
                match_value="项目一",
                target_property="font_name",
                target_value="SimSun",
            ),
            TableRule(
                rule_id="TBL-SIZE",
                priority=11,
                match_type="regex",
                match_value="项目一",
                target_property="font_size",
                target_value="10.5pt",
            ),
        ],
        special_object_rules=[
            SpecialObjectRule(
                rule_id="HDR-CLASS",
                priority=10,
                object_type="header",
                match_type="text",
                match_value="页眉标题",
                target_object_type="running_header",
            ),
        ],
        report_schema=base.report_schema,
    )


def _rule_set_with_split_table_rules() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=[],
        paragraph_rules=[],
        numbering_rules=[],
        table_rules=[
            TableRule(
                rule_id="TBL-HEADER-FONT",
                priority=10,
                match_type="regex",
                match_value=".*",
                target_property="header_row_font_name",
                target_value="SimHei",
            ),
            TableRule(
                rule_id="TBL-HEADER-SIZE",
                priority=11,
                match_type="regex",
                match_value=".*",
                target_property="header_row_font_size",
                target_value="12pt",
            ),
            TableRule(
                rule_id="TBL-BODY-FONT",
                priority=12,
                match_type="regex",
                match_value=".*",
                target_property="body_rows_font_name",
                target_value="SimSun",
            ),
            TableRule(
                rule_id="TBL-BODY-SIZE",
                priority=13,
                match_type="regex",
                match_value=".*",
                target_property="body_rows_font_size",
                target_value="10.5pt",
            ),
        ],
        special_object_rules=[],
        report_schema=base.report_schema,
    )


def _rule_set_with_column_and_cell_table_rules() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=base.document_rules,
        paragraph_rules=[],
        numbering_rules=[],
        table_rules=[
            TableRule(
                rule_id="TBL-COL0-FONT",
                priority=10,
                match_type="regex",
                match_value=".*",
                target_property="column[0]_font_name",
                target_value="SimHei",
            ),
            TableRule(
                rule_id="TBL-CELL11-SIZE",
                priority=11,
                match_type="regex",
                match_value=".*",
                target_property="cell[1,1]_font_size",
                target_value="14pt",
            ),
        ],
        special_object_rules=[],
        report_schema=base.report_schema,
    )


def _rule_set_with_header_named_table_rules() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=base.document_rules,
        paragraph_rules=[],
        numbering_rules=[],
        table_rules=[
            TableRule(
                rule_id="TBL-HEADER-NAME-FONT",
                priority=10,
                match_type="regex",
                match_value=".*",
                target_property="column_by_header[项目二]_font_name",
                target_value="SimHei",
            ),
        ],
        special_object_rules=[],
        report_schema=base.report_schema,
    )


def _rule_set_with_column_range_table_rules() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=base.document_rules,
        paragraph_rules=[],
        numbering_rules=[],
        table_rules=[
            TableRule(
                rule_id="TBL-COL-RANGE-SIZE",
                priority=10,
                match_type="regex",
                match_value=".*",
                target_property="column_range[0:2]_font_size",
                target_value="14pt",
            ),
        ],
        special_object_rules=[],
        report_schema=base.report_schema,
    )


def _rule_set_with_row_table_rules() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=base.document_rules,
        paragraph_rules=[],
        numbering_rules=[],
        table_rules=[
            TableRule(
                rule_id="TBL-ROW1-FONT",
                priority=10,
                match_type="regex",
                match_value=".*",
                target_property="row[1]_font_name",
                target_value="SimHei",
            ),
        ],
        special_object_rules=[],
        report_schema=base.report_schema,
    )


def _rule_set_with_row_range_table_rules() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=base.document_rules,
        paragraph_rules=[],
        numbering_rules=[],
        table_rules=[
            TableRule(
                rule_id="TBL-ROW-RANGE-SIZE",
                priority=10,
                match_type="regex",
                match_value=".*",
                target_property="row_range[0:2]_font_size",
                target_value="14pt",
            ),
        ],
        special_object_rules=[],
        report_schema=base.report_schema,
    )


def _rule_set_with_cell_range_table_rules() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=base.document_rules,
        paragraph_rules=[],
        numbering_rules=[],
        table_rules=[
            TableRule(
                rule_id="TBL-CELL-RANGE-FONT",
                priority=10,
                match_type="regex",
                match_value=".*",
                target_property="cell_range[0:2,0:2]_font_name",
                target_value="SimHei",
            ),
        ],
        special_object_rules=[],
        report_schema=base.report_schema,
    )


def _rule_set_with_table_bold_and_alignment_rules() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=base.document_rules,
        paragraph_rules=[],
        numbering_rules=[],
        table_rules=[
            TableRule(
                rule_id="TBL-COL0-BOLD",
                priority=10,
                match_type="regex",
                match_value=".*",
                target_property="column[0]_bold",
                target_value="true",
            ),
            TableRule(
                rule_id="TBL-CELL01-ALIGN",
                priority=11,
                match_type="regex",
                match_value=".*",
                target_property="cell[0,1]_alignment",
                target_value="center",
            ),
        ],
        special_object_rules=[],
        report_schema=base.report_schema,
    )


def _rule_set_with_table_vertical_alignment_and_border_rules() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=base.document_rules,
        paragraph_rules=[],
        numbering_rules=[],
        table_rules=[
            TableRule(
                rule_id="TBL-HEADER-VALIGN",
                priority=10,
                match_type="regex",
                match_value=".*",
                target_property="header_row_vertical_alignment",
                target_value="center",
            ),
            TableRule(
                rule_id="TBL-CELL00-BORDER",
                priority=11,
                match_type="regex",
                match_value=".*",
                target_property="cell[0,0]_border",
                target_value="single",
            ),
        ],
        special_object_rules=[],
        report_schema=base.report_schema,
    )


def _rule_set_with_late_header_rule() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=[],
        paragraph_rules=[
            ParagraphRule(
                rule_id="HDR-LATE-FONT",
                priority=10,
                match_type="class",
                match_value="late_header",
                target_property="font_name",
                target_value="SimHei",
            ),
            ParagraphRule(
                rule_id="HDR-LATE-SIZE",
                priority=11,
                match_type="class",
                match_value="late_header",
                target_property="font_size",
                target_value="10.5pt",
            ),
        ],
        numbering_rules=[],
        table_rules=[],
        special_object_rules=[
            SpecialObjectRule(
                rule_id="HDR-LATE-CLASS",
                priority=10,
                object_type="header",
                match_type="text",
                match_value="后置页眉",
                target_object_type="late_header",
            ),
        ],
        report_schema=base.report_schema,
    )


def _rule_set_with_mixed_script_body_font_rule() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=[],
        paragraph_rules=[
            ParagraphRule(
                rule_id="BODY-MIXED-FONT",
                priority=10,
                match_type="style",
                match_value="Body Text",
                target_property="font_name",
                target_value="SimSun",
            ),
        ],
        numbering_rules=[],
        table_rules=[],
        special_object_rules=[],
        report_schema=base.report_schema,
    )


def _rule_set_with_abstract_and_reference_rules() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=[],
        paragraph_rules=[
            ParagraphRule(
                rule_id="ABS-EN-FONT",
                priority=10,
                match_type="regex",
                match_value=r"^【Abstract】",
                target_property="font_name",
                target_value="Times New Roman",
            ),
            ParagraphRule(
                rule_id="REF-ENTRY-FONT",
                priority=20,
                match_type="regex",
                match_value=r"^.+\[[A-Z]\]\s*\.",
                target_property="font_name",
                target_value="SimSun",
            ),
            ParagraphRule(
                rule_id="REF-ENTRY-HANGING",
                priority=21,
                match_type="regex",
                match_value=r"^.+\[[A-Z]\]\s*\.",
                target_property="hanging_indent",
                target_value="0.74cm",
            ),
        ],
        numbering_rules=[],
        table_rules=[],
        special_object_rules=[],
        report_schema=base.report_schema,
    )


def _rule_set_with_inherited_layout_rules() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=base.document_rules,
        paragraph_rules=[
            *base.paragraph_rules,
            ParagraphRule(
                rule_id="PAR-BODY-SPACE-BEFORE",
                priority=13,
                match_type="style",
                match_value="Body Text",
                target_property="space_before",
                target_value="0pt",
            ),
            ParagraphRule(
                rule_id="PAR-BODY-LINE-SPACING",
                priority=14,
                match_type="style",
                match_value="Body Text",
                target_property="line_spacing",
                target_value="20pt",
            ),
        ],
        numbering_rules=base.numbering_rules,
        table_rules=base.table_rules,
        special_object_rules=base.special_object_rules,
        report_schema=base.report_schema,
    )


def _rule_set_with_inline_tag_rules() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=[],
        paragraph_rules=[
            ParagraphRule(
                rule_id="INLINE-LABEL-FONT",
                priority=10,
                match_type="regex",
                match_value=r"^【摘要】",
                target_property="label_font_name",
                target_value="SimHei",
            ),
            ParagraphRule(
                rule_id="INLINE-LABEL-BOLD",
                priority=11,
                match_type="regex",
                match_value=r"^【摘要】",
                target_property="label_bold",
                target_value="true",
            ),
            ParagraphRule(
                rule_id="INLINE-LABEL-SIZE",
                priority=12,
                match_type="regex",
                match_value=r"^【摘要】",
                target_property="label_font_size",
                target_value="15pt",
            ),
            ParagraphRule(
                rule_id="INLINE-CONTENT-FONT",
                priority=13,
                match_type="regex",
                match_value=r"^【摘要】",
                target_property="content_font_name",
                target_value="SimSun",
            ),
            ParagraphRule(
                rule_id="INLINE-CONTENT-SIZE",
                priority=14,
                match_type="regex",
                match_value=r"^【摘要】",
                target_property="content_font_size",
                target_value="12pt",
            ),
        ],
        numbering_rules=[],
        table_rules=[],
        special_object_rules=[],
        report_schema=base.report_schema,
    )


def _rule_set_with_inline_label_only_rules() -> RuleSet:
    base = _rule_set()
    return RuleSet(
        document_rules=[],
        paragraph_rules=[
            ParagraphRule(
                rule_id="INLINE-LABEL-ONLY-FONT",
                priority=10,
                match_type="regex",
                match_value=r"^【摘要】",
                target_property="label_font_name",
                target_value="SimHei",
            ),
            ParagraphRule(
                rule_id="INLINE-LABEL-ONLY-BOLD",
                priority=11,
                match_type="regex",
                match_value=r"^【摘要】",
                target_property="label_bold",
                target_value="true",
            ),
            ParagraphRule(
                rule_id="INLINE-LABEL-ONLY-SIZE",
                priority=12,
                match_type="regex",
                match_value=r"^【摘要】",
                target_property="label_font_size",
                target_value="15pt",
            ),
        ],
        numbering_rules=[],
        table_rules=[],
        special_object_rules=[],
        report_schema=base.report_schema,
    )
